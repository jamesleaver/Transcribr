#!/usr/bin/env python3
"""
Transcribr - GUI for transcribing audio/video files with Whisper,
grouping the result into paragraphs, and reviewing/labelling speakers.

(c) James Leaver, 2026. This software is experimental. Everything runs
locally on your computer: a Whisper engine (faster-whisper by default;
openai-whisper and mlx-whisper are used when installed) transcribes
the audio, an optional speaker-detection pass works out who is
speaking, the result is grouped into paragraphs, and a review pane
lets you check the speaker labels, edit text, and listen back to the
source audio before saving as .docx, .pdf, or .txt. Several files can
be queued and transcribed in one unattended batch. When a particular
model is run for the first time it is downloaded and stored locally;
'large-v3-turbo' (the Standard choice) is recommended. Use at your own
risk.
Questions: jleaver@sgchambers.com.au.

Run with:
    python3 transcribr.py
"""

__version__ = "0.9.0"

ABOUT_TEXT = (
    f"Version {__version__}\n"
    "(c) James Leaver, 2026.\n\n"
    "Transcribr transcribes audio and video with Whisper, can work out "
    "who is speaking, groups the result into paragraphs, and gives you "
    "a review pane to check speaker labels, edit text, and listen back "
    "to the source audio before saving as Word (.docx), PDF, or plain "
    "text. Several files can be queued and transcribed in one "
    "unattended batch. Everything runs locally on your computer - "
    "nothing is uploaded.\n\n"
    "When a particular model is run for the first time, that model will "
    "be downloaded to your computer and stored locally. The Standard "
    "model (large-v3-turbo) is recommended.\n\n"
    "This software is experimental. Use at your own risk.\n\n"
    "Questions: jleaver@sgchambers.com.au"
)

import contextlib
import os
import queue
import re
import shutil
import subprocess
import sys
import threading
import time
import traceback
from pathlib import Path

# The colour palettes below are the single source served to the web
# front-end as CSS variables. (The Tk-era theme packages sv-ttk and
# darkdetect are no longer used - the front-end resolves light/dark
# itself, including "auto" via prefers-color-scheme.)


# =====================================================================
# Colour palettes (light / dark)
# =====================================================================
#
# ttk widgets are restyled wholesale by sv-ttk, but plain tk widgets (the
# transcript Text, the prompt/log text boxes, the batch Listbox, speaker
# badges, the drop zone) need explicit colours. Everything below reads
# the active palette via _palette() so a theme switch can re-skin live.

_PALETTES = {
    "light": {
        "text_bg": "#ffffff", "text_fg": "#111111", "insert": "#111111",
        "speaker_fg": "#222222", "timestamp_fg": "#777777",
        "selected_bg": "#d6d6d6", "editing_bg": "#ffcc80",
        "search_bg": "#ffe066",
        "conf_low": "#ffd2d2", "conf_med": "#ffe9c7",
        "badge_fg": "#111111",
        "speaker_colours": {
            "1": "#fff4d6",   # warm yellow
            "2": "#dfeeff",   # soft blue
            "3": "#e3f4d8",   # soft green
            "4": "#f9d9e7",   # soft pink
            "5": "#e7ddf6",   # lavender
            "6": "#ffe0c2",   # peach
            "7": "#d3f0ee",   # teal
            "8": "#f0e4c8",   # tan
            "9": "#d9e8c0",   # olive
        },
        "drop_bg": "#f5f6f8", "drop_hover": "#e8f0fe",
        "drop_border": "#9aa0a6", "drop_fg": "#5f6368",
    },
    "dark": {
        "text_bg": "#1e1e20", "text_fg": "#e6e6e6", "insert": "#e6e6e6",
        "speaker_fg": "#dddddd", "timestamp_fg": "#9a9a9a",
        "selected_bg": "#46484c", "editing_bg": "#6b4e1e",
        "search_bg": "#7a6a14",
        "conf_low": "#5c2b2b", "conf_med": "#5c4a26",
        "badge_fg": "#e6e6e6",
        "speaker_colours": {
            "1": "#4d4426",   # warm yellow, dimmed
            "2": "#283d52",   # soft blue, dimmed
            "3": "#2f4528",   # soft green, dimmed
            "4": "#4d2c3c",   # soft pink, dimmed
            "5": "#3b3050",   # lavender, dimmed
            "6": "#52391f",   # peach, dimmed
            "7": "#1f4441",   # teal, dimmed
            "8": "#46402a",   # tan, dimmed
            "9": "#3a4426",   # olive, dimmed
        },
        "drop_bg": "#2a2a2c", "drop_hover": "#1f3a5f",
        "drop_border": "#5f6368", "drop_fg": "#9aa0a6",
    },
}

# The currently-active palette key. Set by _apply_theme(); defaults to
# light so module-level consumers work before the GUI configures a theme.
_ACTIVE_THEME = "light"


def _resolve_theme(setting):
    """Map a theme setting ("auto"/"light"/"dark") to a palette key. The
    running app resolves "auto" in the front-end (prefers-color-scheme);
    server-side it falls back to light."""
    if setting in ("light", "dark"):
        return setting
    return "light"


def _palette():
    return _PALETTES[_ACTIVE_THEME]


# =====================================================================
# Cross-platform helpers
# =====================================================================

if sys.platform == "darwin":
    _REVEAL_LABEL = "Reveal in Finder"
elif sys.platform == "win32":
    _REVEAL_LABEL = "Show in Explorer"
else:
    _REVEAL_LABEL = "Show in Folder"


def _open_path(path):
    """Open a file with the OS's default handler."""
    path = str(path)
    try:
        if sys.platform == "darwin":
            subprocess.run(["open", path], check=False)
        elif sys.platform == "win32":
            os.startfile(path)  # type: ignore[attr-defined]
        else:
            subprocess.run(["xdg-open", path], check=False)
    except Exception:
        # Best-effort; failing to open shouldn't break anything else.
        pass


def _reveal_path(path):
    """Show the file in the OS's file manager, with the file selected."""
    path = str(path)
    try:
        if sys.platform == "darwin":
            subprocess.run(["open", "-R", path], check=False)
        elif sys.platform == "win32":
            # explorer /select, requires no space after the comma; it also
            # exits non-zero even on success, so we don't check the return.
            subprocess.run(["explorer", f"/select,{path}"], check=False)
        else:
            # Linux: open the parent folder; selecting a specific file isn't
            # universally supported across file managers.
            subprocess.run(["xdg-open", str(Path(path).parent)], check=False)
    except Exception:
        pass


def _find_readme():
    """Return a path to README.md / README.txt if one exists near the script."""
    here = Path(__file__).resolve().parent
    candidates = [
        here / "README.md",
        here / "README.txt",
        here.parent / "README.md",
        here.parent / "README.txt",
    ]
    for c in candidates:
        if c.exists():
            return c
    return None


def _log_dir():
    """Return a writable per-user directory for log files."""
    if sys.platform == "darwin":
        base = Path.home() / "Library" / "Logs" / "Transcribr"
    elif sys.platform == "win32":
        base = Path(os.environ.get("LOCALAPPDATA",
                                   Path.home() / "AppData" / "Local"))
        base = base / "Transcribr" / "Logs"
    else:
        base = Path(os.environ.get("XDG_STATE_HOME",
                                   Path.home() / ".local" / "state"))
        base = base / "Transcribr"
    base.mkdir(parents=True, exist_ok=True)
    return base


def _log_file_path():
    return _log_dir() / "transcribr.log"


def _log(msg, *, exc_info=None):
    """Append a timestamped line to the log file. Best-effort; never raises."""
    try:
        line = f"[{time.strftime('%Y-%m-%d %H:%M:%S')}] {msg}\n"
        with open(_log_file_path(), "a", encoding="utf-8") as fh:
            fh.write(line)
            if exc_info is not None:
                traceback.print_exception(*exc_info, file=fh)
                fh.write("\n")
    except Exception:
        pass


def _install_crash_logging(root=None):
    """Route uncaught Python and Tk callback exceptions to the log file."""
    def py_hook(exc_type, exc, tb):
        _log("UNCAUGHT EXCEPTION", exc_info=(exc_type, exc, tb))
        sys.__excepthook__(exc_type, exc, tb)
    sys.excepthook = py_hook
    if root is not None:
        def tk_hook(exc_type, exc, tb):
            _log("TK CALLBACK EXCEPTION", exc_info=(exc_type, exc, tb))
            traceback.print_exception(exc_type, exc, tb)
        root.report_callback_exception = tk_hook


def _next_revision_path(original):
    """Return a sibling path with a .revN suffix that doesn't yet exist.

    e.g. /a/b.docx -> /a/b.rev1.docx, then .rev2.docx, etc. Strips an
    existing .revN suffix on the input so revisions of revisions don't
    pile up nested suffixes."""
    p = Path(original)
    ext = p.suffix
    stem = p.stem
    m = re.match(r"^(.*)\.rev\d+$", stem)
    if m:
        stem = m.group(1)
    n = 1
    while True:
        candidate = p.with_name(f"{stem}.rev{n}{ext}")
        if not candidate.exists():
            return candidate
        n += 1


_AUDIO_EXTENSIONS = (".mp3", ".wav", ".m4a", ".aac", ".flac", ".ogg",
                     ".opus", ".mp4", ".mov", ".mkv", ".avi", ".webm")


def _guess_audio_for_transcript(transcript_path):
    """Best-effort: find the source audio/video file next to a saved
    transcript, so playback works when re-opening it for review.

    Transcripts are written as <media stem>.transcript.<ext> (possibly
    with a .revN inserted), so we strip those suffixes and probe the
    known media extensions. Returns a str path or None."""
    p = Path(transcript_path)
    stem = p.stem  # e.g. "interview.transcript" or "interview.transcript.rev2"
    m = re.match(r"^(.*)\.rev\d+$", stem)
    if m:
        stem = m.group(1)
    if stem.endswith(".transcript"):
        stem = stem[: -len(".transcript")]
    for ext in _AUDIO_EXTENSIONS:
        candidate = p.with_name(stem + ext)
        if candidate.exists():
            return str(candidate)
    return None


# ---- Recent transcripts persistence -------------------------------------

_RECENT_MAX = 10


def _config_dir():
    if sys.platform == "darwin":
        base = Path.home() / "Library" / "Application Support" / "Transcribr"
    elif sys.platform == "win32":
        base = Path(os.environ.get("APPDATA",
                                   Path.home() / "AppData" / "Roaming"))
        base = base / "Transcribr"
    else:
        base = Path(os.environ.get("XDG_CONFIG_HOME",
                                   Path.home() / ".config"))
        base = base / "Transcribr"
    base.mkdir(parents=True, exist_ok=True)
    return base


def _recent_file():
    return _config_dir() / "recent.json"


def _recent_load():
    """Return the list of recent transcript paths, most-recent first.
    Best-effort; returns [] on any read/parse error."""
    try:
        import json
        path = _recent_file()
        if not path.exists():
            return []
        with open(path, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        if isinstance(data, list):
            return [str(p) for p in data if isinstance(p, str)]
    except Exception:
        pass
    return []


def _recent_save(paths):
    try:
        import json
        with open(_recent_file(), "w", encoding="utf-8") as fh:
            json.dump(list(paths)[:_RECENT_MAX], fh, indent=2)
    except Exception as e:
        _log(f"Could not write recent.json: {e}")


def _recent_add(path):
    """Move `path` to the front of the recent list (deduping). No-op on error."""
    try:
        path = str(Path(path).resolve())
    except Exception:
        path = str(path)
    items = _recent_load()
    items = [p for p in items if p != path]
    items.insert(0, path)
    _recent_save(items)


def _settings_file():
    return _config_dir() / "settings.json"


def _settings_load():
    """Return the saved UI settings dict, or {} on any read/parse error."""
    try:
        import json
        path = _settings_file()
        if not path.exists():
            return {}
        with open(path, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        if isinstance(data, dict):
            return data
    except Exception:
        pass
    return {}


def _settings_save(settings):
    """Persist the UI settings dict. Best-effort; logs on failure."""
    try:
        import json
        with open(_settings_file(), "w", encoding="utf-8") as fh:
            json.dump(dict(settings), fh, indent=2)
    except Exception as e:
        _log(f"Could not write settings.json: {e}")


# ---- Developer annotations ------------------------------------------------
#
# A development aid ported from the Chambers Availability tool: with
# `--annotate` (or TRANSCRIBR_ANNOTATE=1) a pencil overlay appears in
# the UI; clicking any element captures where it is (selector path,
# text, markup, geometry) and pins a note to it. Notes land in
# annotations.json under the config dir, ready to be pasted into a
# development session. The overlay is invisible to normal users.

ANNOTATE_MODE = False


def annotate_enabled():
    """Whether the annotation overlay should show. Checked live so the
    marker file works without a restart: the --annotate flag, the
    TRANSCRIBR_ANNOTATE=1 environment variable, or an `annotate.on`
    file in the config dir (`touch` it to enable the overlay for the
    double-clicked app; delete it to hide again)."""
    if ANNOTATE_MODE:
        return True
    if os.environ.get("TRANSCRIBR_ANNOTATE", "").strip() == "1":
        return True
    try:
        return (_config_dir() / "annotate.on").exists()
    except OSError:
        return False


_annotations_lock = threading.Lock()

# Field -> stored size cap. Keeps a runaway page from bloating the file.
_ANNOTATION_FIELDS = {
    "view": 100, "selector": 500, "element_text": 500,
    "html": 4000, "rect": 300, "note": 4000,
}


def _annotations_file():
    return _config_dir() / "annotations.json"


def _annotations_load():
    try:
        import json
        with open(_annotations_file(), encoding="utf-8") as fh:
            data = json.load(fh)
        if isinstance(data, list):
            return data
    except Exception:
        pass
    return []


def _annotations_save(items):
    import json
    with open(_annotations_file(), "w", encoding="utf-8") as fh:
        json.dump(items, fh, indent=2, ensure_ascii=False)


def annotation_add(body):
    """Validate + append one annotation; returns the stored record."""
    record = {}
    for key, cap in _ANNOTATION_FIELDS.items():
        record[key] = str(body.get(key) or "")[:cap]
    if not record["note"].strip():
        raise ApiFail(400, "empty_note", "The note is empty.")
    with _annotations_lock:
        items = _annotations_load()
        record["id"] = (max((a.get("id", 0) for a in items), default=0)
                        + 1)
        record["created"] = time.strftime("%Y-%m-%d %H:%M:%S")
        record["app_version"] = __version__
        items.append(record)
        _annotations_save(items)
    return record


def annotation_delete(ann_id):
    with _annotations_lock:
        items = _annotations_load()
        kept = [a for a in items if a.get("id") != ann_id]
        if len(kept) == len(items):
            raise ApiFail(404, "not_found", f"No annotation #{ann_id}.")
        _annotations_save(kept)


def annotations_clear():
    with _annotations_lock:
        _annotations_save([])


# ---- Review auto-save (crash recovery) -----------------------------------
#
# While the review pane is open, the labelled-but-unsaved state is
# periodically written here. It is deleted on any clean exit from review;
# if it survives to the next launch, the app offers to restore it.

def _autosave_file():
    return _config_dir() / "autosave.json"


def _autosave_load():
    """Return the saved review session dict, or None."""
    try:
        import json
        path = _autosave_file()
        if not path.exists():
            return None
        with open(path, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        if isinstance(data, dict) and data.get("paragraphs"):
            return data
    except Exception:
        pass
    return None


def _autosave_save(data):
    try:
        import json
        with open(_autosave_file(), "w", encoding="utf-8") as fh:
            json.dump(data, fh)
    except Exception as e:
        _log(f"Could not write autosave.json: {e}")


def _autosave_clear():
    try:
        _autosave_file().unlink(missing_ok=True)
    except Exception:
        pass


# =====================================================================
# Paragraphify (identical to the standalone script)
# =====================================================================

# Sentinel speaker label used in saved transcripts to mark a paragraph
# that explicitly has no speaker assigned, so that loading the file back
# distinguishes "unlabelled by intent" from "continuing the previous
# speaker". Only emitted on a transition from a named speaker to an
# unattributed paragraph; runs of unattributed paragraphs at the start of
# a transcript don't need the marker.
UNATTRIBUTED_LABEL = "[Unattributed]"


SHORT_RESPONSES = {
    "yes", "yeah", "yep", "yup",
    "no", "nope", "nah",
    "okay", "ok", "alright", "all right",
    "mm", "mhm", "mmhm", "uh huh", "uh-huh",
    "right", "sure", "correct", "true", "exactly",
    "thanks", "thank you",
}


def format_timestamp(seconds: float) -> str:
    total = int(seconds)
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"[{h:d}:{m:02d}:{s:02d}]"
    return f"[{m:02d}:{s:02d}]"


def _format_duration(seconds: float) -> str:
    """Human-readable duration: '12s', '3m 04s', '1h 23m'."""
    seconds = max(0, int(seconds))
    if seconds < 60:
        return f"{seconds}s"
    if seconds < 3600:
        m, s = divmod(seconds, 60)
        return f"{m}m {s:02d}s"
    h, rem = divmod(seconds, 3600)
    m = rem // 60
    return f"{h}h {m:02d}m"


# ---- Audio decoding (PyAV first, ffmpeg binary as fallback) ----------
#
# PyAV ships FFmpeg's libraries inside its wheel (it is already a hard
# dependency of faster-whisper), so with it installed Transcribr needs
# no ffmpeg/ffprobe binaries at all: duration probing, decoding for the
# engines and the speaker detector, and playback extraction all go
# through it. The subprocess fallbacks remain for environments that
# only have the binaries (e.g. an openai-whisper-only install).

def _have_pyav():
    import importlib.util
    return importlib.util.find_spec("av") is not None


AUDIO_SAMPLE_RATE = 16000


def decode_audio_16k(path):
    """Decode any audio/video file to mono float32 at 16 kHz via PyAV -
    the input format Whisper and the speaker detector both expect.
    Returns a numpy array, or None when PyAV is unavailable or decoding
    fails (callers then fall back to handing the engine the file
    path)."""
    try:
        import av
        import numpy as np
    except ImportError:
        return None
    try:
        frames = []
        with av.open(str(path)) as container:
            if not container.streams.audio:
                return None
            stream = container.streams.audio[0]
            resampler = av.AudioResampler(
                format="s16", layout="mono", rate=AUDIO_SAMPLE_RATE)
            for frame in container.decode(stream):
                for rf in resampler.resample(frame):
                    frames.append(rf.to_ndarray())
            for rf in resampler.resample(None):     # flush
                frames.append(rf.to_ndarray())
        if not frames:
            return None
        pcm = np.concatenate([f.reshape(-1) for f in frames])
        return pcm.astype(np.float32) / 32768.0
    except Exception:
        _log(f"PyAV decode failed for {path}:\n{traceback.format_exc()}")
        return None


# File signatures that regularly turn up wearing an audio extension
# (exports from evidence/records systems are the usual culprits). Used
# to explain a decode failure in plain language instead of dumping an
# ffmpeg traceback on the user.
_NON_AUDIO_SIGNATURES = (
    (b"PK\x03\x04", "a zip archive or Microsoft Office document"),
    (b"%PDF", "a PDF document"),
    (b"{\\rtf", "an RTF document"),
    (b"<!DO", "a web page"),
    (b"<htm", "a web page"),
    (b"\xd0\xcf\x11\xe0", "a legacy Microsoft Office document"),
)


def describe_non_audio_file(path):
    """Why `path` can't be decoded as audio, in plain language, or None
    when it plausibly is audio (only consulted after a decode failed)."""
    try:
        with open(path, "rb") as fh:
            head = fh.read(8)
    except OSError as e:
        return f"it could not be read ({e.strerror or e})"
    if not head:
        return "it is empty (0 bytes)"
    for magic, kind in _NON_AUDIO_SIGNATURES:
        if head.startswith(magic):
            return (f"it appears to be {kind} that has been given an "
                    "audio file's name, not an actual recording")
    try:
        import av
        with av.open(str(path)) as container:
            if not container.streams.audio:
                return "it contains no audio stream"
        return None
    except ImportError:
        return None
    except Exception as e:
        detail = str(e).strip() or type(e).__name__
        return f"it could not be read as audio ({detail})"


def _pyav_duration(path):
    """Duration in seconds via PyAV container metadata, or None."""
    try:
        import av
    except ImportError:
        return None
    try:
        with av.open(str(path)) as container:
            # container.duration is in AV_TIME_BASE units (microseconds).
            if container.duration:
                return float(container.duration) / 1_000_000
            for s in container.streams.audio:
                if s.duration and s.time_base:
                    return float(s.duration * s.time_base)
    except Exception:
        return None
    return None


def _ffprobe_duration(path):
    """Duration in seconds via the ffprobe binary, or None.

    On Windows we hide the console window that subprocess.run would
    otherwise briefly flash for the ffprobe process.
    """
    import subprocess
    creationflags = 0
    if sys.platform == "win32":
        creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    try:
        result = subprocess.run(
            ["ffprobe", "-v", "error",
             "-show_entries", "format=duration",
             "-of", "default=noprint_wrappers=1:nokey=1",
             str(path)],
            capture_output=True, text=True, timeout=15,
            creationflags=creationflags,
        )
    except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
        return None
    if result.returncode != 0:
        return None
    try:
        return float(result.stdout.strip())
    except ValueError:
        return None


def get_audio_duration(path):
    """Return the duration of an audio/video file in seconds, or None on
    failure. PyAV first; ffprobe when PyAV is missing or stumped."""
    duration = _pyav_duration(path)
    if duration is not None:
        return duration
    return _ffprobe_duration(path)


def _is_short_response(text: str) -> bool:
    cleaned = re.sub(r"[^\w\s'-]", "", text).strip().lower()
    if not cleaned:
        return False
    if cleaned in SHORT_RESPONSES:
        return True
    words = cleaned.split()
    return len(words) <= 3 and all(w in SHORT_RESPONSES for w in words)


# A sentence ending (". !") only breaks a paragraph when the speaker
# also paused: this fraction of the user's gap threshold, floored so a
# very low threshold doesn't turn every full stop into a paragraph.
_SOFT_BREAK_FRACTION = 0.4
_SOFT_BREAK_MIN_SECONDS = 0.35

# Trailing marks Whisper emits when speech is cut off or trails away -
# in dialogue that's a turn change almost every time.
_INTERRUPTION_ENDINGS = ("--", "—", "–", "...", "…")


def _leading_acknowledgment(text: str) -> bool:
    """True when the segment OPENS with a yes/no/okay-style token
    ("Yeah, I did") - a strong cue that a different person is now
    answering, even mid-flow."""
    cleaned = re.sub(r"[^\w\s'-]", " ", text or "").strip().lower()
    first = cleaned.split(" ", 1)[0] if cleaned else ""
    return first in SHORT_RESPONSES


def _gaps_are_informative(segments) -> bool:
    """Whether the segment times carry real silence information.

    faster-whisper without word timestamps emits wall-to-wall segments
    (each end equals the next start), hiding every actual pause. When
    nearly every boundary is contiguous like that, gap-based rules
    would never fire, so the soft cues must stand on their own."""
    if len(segments) < 2:
        return True
    contiguous = sum(1 for a, b in zip(segments, segments[1:])
                     if b[0] - a[1] < 0.05)
    return contiguous / (len(segments) - 1) < 0.7


def _should_break(prev, curr, gap_threshold: float,
                  gaps_informative: bool = True) -> bool:
    """The paragraph-boundary heuristics, graded by strength:

    - a silence at least the user's threshold always breaks;
    - a question mark, an interruption/trail-off ending, or a short
      conversational response always breaks (dialogue turn cues);
    - a sentence ending (". !") or an opening acknowledgment ("Yeah,
      ...") breaks only when paired with a smaller "soft" pause - so
      a monologue read at speed stays one paragraph instead of
      splintering at every full stop, while real turn-taking (which
      nearly always carries a beat of silence) still splits.

    With `gaps_informative` False (the engine recorded no real pauses)
    the soft cues break unconditionally - punctuation is the only
    signal left."""
    if prev is None:
        return True
    _, prev_end, prev_text = prev
    curr_start, _, curr_text = curr
    prev_stripped = prev_text.rstrip()
    gap = curr_start - prev_end
    soft_gap = max(_SOFT_BREAK_MIN_SECONDS,
                   gap_threshold * _SOFT_BREAK_FRACTION)

    # Hard silence beyond the user's pause threshold.
    if gap >= gap_threshold:
        return True
    # A question hands the floor over; the answer is its own paragraph.
    if prev_stripped.endswith("?"):
        return True
    # Cut-off / trailing-away speech is a turn change in dialogue.
    if prev_stripped.endswith(_INTERRUPTION_ENDINGS):
        return True
    # Short conversational responses (yes/no/etc.) on either side.
    if _is_short_response(prev_text) or _is_short_response(curr_text):
        return True
    # Weaker cues need a beat of silence to confirm them - unless no
    # real silence was ever recorded.
    if gap >= soft_gap or not gaps_informative:
        if prev_stripped.endswith((".", "!")):
            return True
        if _leading_acknowledgment(curr_text):
            return True
    return False


# Safety cap: even if no break signal fires, a paragraph shouldn't grow
# beyond this many seconds of audio. Without the cap, an uninterrupted
# monologue can collapse into a single unreadable paragraph in the
# review pane.
_PARAGRAPH_SECONDS_CAP = 60.0


def _refine_segment_times(segments, words):
    """Snap each segment's start/end inward to its first/last word.

    Whisper pads segment boundaries with the surrounding silence, so
    the gap between two segments understates the actual pause between
    the spoken words. When word timestamps were recorded, trimming each
    segment to its words makes every gap-based paragraph decision (and
    the printed timestamps) measure real silence. Inward-only, so
    segment order and spans stay consistent; segments whose words can't
    be found pass through untouched."""
    if not words:
        return segments
    refined = []
    wi, n_words = 0, len(words)
    for seg in segments:
        s_start, s_end, text = seg[0], seg[1], seg[2]
        while wi < n_words and words[wi][1] <= s_start:
            wi += 1
        first_start, last_end = None, None
        wj = wi
        while wj < n_words and words[wj][0] < s_end:
            if first_start is None:
                first_start = words[wj][0]
            last_end = words[wj][1]
            wj += 1
        wi = wj
        new_start = (first_start
                     if first_start is not None
                     and s_start <= first_start < s_end else s_start)
        new_end = (last_end
                   if last_end is not None
                   and s_start < last_end <= s_end else s_end)
        if new_start < new_end:
            refined.append((new_start, new_end, text))
        else:
            refined.append(seg)
    return refined


def paragraphify(segments, gap_threshold: float, words=None):
    """Group segments into paragraphs on purely programmatic cues
    (see _should_break). `words`, when word timestamps were recorded,
    sharpens the gap measurements via _refine_segment_times."""
    segments = _refine_segment_times(segments, words or [])
    gaps_informative = _gaps_are_informative(segments)
    paragraphs, current, prev = [], [], None
    para_start = None
    for seg in segments:
        if current:
            forced_by_cap = (
                para_start is not None
                and seg[1] - para_start >= _PARAGRAPH_SECONDS_CAP)
            if forced_by_cap or _should_break(prev, seg, gap_threshold,
                                              gaps_informative):
                paragraphs.append(current)
                current = []
                para_start = None
        if para_start is None:
            para_start = seg[0]
        current.append(seg)
        prev = seg
    if current:
        paragraphs.append(current)
    return paragraphs


def render(paragraphs, *, show_timestamp=True, title=None, speakers=None) -> str:
    """Render paragraphs to plain text.

    `speakers`, if given, is a list parallel to `paragraphs` where each
    entry is either a speaker label string (e.g. "CONSTABLE MACKLEBUM")
    or None for paragraphs with no assigned speaker. We only emit the
    speaker label when it differs from the previous paragraph's, which
    matches conventional transcript style.
    """
    out_lines = []
    if title:
        # Plain text has no bold; we just put the title at the top.
        # The "\n\n".join below produces a blank line between this title
        # and the first paragraph.
        out_lines.append(title)

    last_speaker = None
    for i, para in enumerate(paragraphs):
        start = para[0][0]
        body = " ".join(seg[2] for seg in para)
        speaker = speakers[i] if speakers and i < len(speakers) else None
        block_lines = []
        if speaker and speaker != last_speaker:
            block_lines.append(f"{speaker}:")
        elif speaker is None and last_speaker is not None:
            # Transition from a named speaker to an unattributed paragraph.
            # Emit an explicit marker so re-parsing this file doesn't carry
            # the previous speaker over.
            block_lines.append(f"{UNATTRIBUTED_LABEL}:")
        if show_timestamp:
            block_lines.append(f"{format_timestamp(start)}  {body}")
        else:
            block_lines.append(body)
        out_lines.append("\n".join(block_lines))
        last_speaker = speaker

    return "\n\n".join(out_lines) + "\n"


def write_paragraphs_to_file(paragraphs, out_path, *, show_timestamp=True,
                              title=None, output_format="txt", speakers=None):
    """Single entry point for writing transcript output.

    Centralises the txt-vs-docx-vs-pdf switch so the worker (direct-write
    path) and the GUI (review-screen path) can both call the same function.

    Raises ImportError with a friendly message if .docx or .pdf is
    requested but the needed package isn't installed.
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if output_format == "docx":
        try:
            _write_docx(paragraphs, out_path,
                        show_timestamp=show_timestamp, title=title,
                        speakers=speakers)
        except ImportError:
            raise ImportError(
                "Cannot write .docx: the 'python-docx' package is not "
                "installed. Install it with:\n  pip install python-docx\n"
                "Or pick the .txt format instead.")
    elif output_format == "pdf":
        try:
            _write_pdf(paragraphs, out_path,
                       show_timestamp=show_timestamp, title=title,
                       speakers=speakers)
        except ImportError:
            raise ImportError(
                "Cannot write .pdf: the 'reportlab' package is not "
                "installed. Install it with:\n  pip install reportlab\n"
                "Or pick the .txt or .docx format instead.")
    else:
        out_path.write_text(
            render(paragraphs, show_timestamp=show_timestamp,
                   title=title, speakers=speakers),
            encoding="utf-8",
        )


# =====================================================================
# Reading transcripts back in (the inverse of writing)
# =====================================================================

# Matches "[MM:SS]" or "[H:MM:SS]" at the start of a line (paragraph body).
_TS_PREFIX_RE = re.compile(r"^\s*\[(\d+:)?(\d+):(\d+)\]\s*")

# Matches a Transcribr disclaimer line. We strip these so they don't
# come back as paragraphs.
_DISCLAIMER_RE = re.compile(
    r"^Transcribed using Transcribr",
    re.IGNORECASE,
)


def _parse_timestamp_prefix(line):
    """Return (start_seconds, body_text) if `line` begins with [MM:SS] or
    [H:MM:SS], else (None, line). Used by both .txt and .docx parsers."""
    m = _TS_PREFIX_RE.match(line)
    if not m:
        return None, line
    h = int(m.group(1)[:-1]) if m.group(1) else 0
    mm = int(m.group(2))
    ss = int(m.group(3))
    return h * 3600 + mm * 60 + ss, line[m.end():]


class TranscriptParseError(Exception):
    """Raised when a file can't be parsed as a Transcribr transcript."""


def read_paragraphs_from_file(path):
    """Read a Transcribr-written .txt or .docx and return a dict with:
        paragraphs    - list of paragraphs in the format used by the
                        review pane (each paragraph is a list of
                        (start, end, text) tuples; we synthesise one
                        segment per paragraph since the original
                        segmentation isn't recoverable)
        speakers      - list of speaker name strings (or None) parallel
                        to paragraphs
        title         - the title at the top of the document, or None
        show_timestamp - True if the file contained timestamps, False
                        if not (informs how it should be re-rendered)

    Raises TranscriptParseError on any failure with a message suitable
    for showing the user.
    """
    path = Path(path)
    if not path.exists():
        raise TranscriptParseError(f"File not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".txt":
        parsed = _parse_txt_transcript(path)
    elif suffix == ".docx":
        parsed = _parse_docx_transcript(path)
    else:
        raise TranscriptParseError(
            f"Unsupported file extension: {suffix}. Only .txt and .docx are supported."
        )
    _infer_paragraph_end_times(parsed["paragraphs"])
    return parsed


def _infer_paragraph_end_times(paragraphs):
    """Saved transcripts only record each paragraph's start time, so the
    parsers synthesise a placeholder 1-second span. Stretch each
    paragraph's end out to the next paragraph's start so the span covers
    the actual speech - audio playback depends on this. The last
    paragraph's true end is unknowable from the file; it keeps the
    placeholder span and playback treats it as open-ended."""
    for i in range(len(paragraphs) - 1):
        para = paragraphs[i]
        nxt = paragraphs[i + 1]
        if not para or not nxt:
            continue
        next_start = nxt[0][0]
        start, end, text = para[-1]
        if next_start > end:
            para[-1] = (start, next_start, text)


def _parse_txt_transcript(path):
    """Parse a .txt file written by render(). Format:

        <title (optional)>

        SPEAKER NAME:
        [MM:SS]  body text

        [MM:SS]  body text from same speaker (no label since unchanged)

        OTHER SPEAKER:
        [MM:SS]  ...

        Transcribed using Transcribr ... (disclaimer line - skipped)

    Blocks are separated by blank lines. We tolerate transcripts that
    omit timestamps (text-only output) and transcripts that omit
    speaker labels (just text and timestamps)."""
    text = path.read_text(encoding="utf-8")
    # Split into blocks at blank lines. Strip trailing whitespace per block
    # but preserve internal newlines (a block is "SPEAKER:\n[MM:SS]  body").
    blocks = [b.strip("\n") for b in re.split(r"\n[ \t]*\n", text) if b.strip()]
    if not blocks:
        raise TranscriptParseError("File is empty.")

    title = None
    paragraphs = []
    speakers = []
    current_speaker = None
    saw_timestamp = False  # tracks whether we ever saw [MM:SS]

    # First block could be the title. A title looks like a single line
    # that isn't a speaker label and doesn't start with a timestamp.
    first = blocks[0]
    first_lines = first.splitlines()
    looks_like_speaker = (len(first_lines) >= 1
                          and first_lines[0].rstrip().endswith(":")
                          and not _TS_PREFIX_RE.match(first_lines[0]))
    has_timestamp = bool(_TS_PREFIX_RE.match(first_lines[0]))
    if not looks_like_speaker and not has_timestamp and len(first_lines) <= 3:
        # Treat as title and consume.
        title = first
        blocks = blocks[1:]

    for block in blocks:
        # Skip disclaimer
        if _DISCLAIMER_RE.match(block):
            continue
        lines = block.splitlines()
        # Speaker label is the first line iff it ends with ":" and isn't
        # a timestamp.
        if lines and lines[0].rstrip().endswith(":") and not _TS_PREFIX_RE.match(lines[0]):
            label = lines[0].rstrip().rstrip(":").strip()
            # Recognise the explicit "no speaker" marker we emit on
            # transitions from a named speaker back to unattributed.
            current_speaker = None if label == UNATTRIBUTED_LABEL else label
            lines = lines[1:]
        if not lines:
            # Speaker line on its own with no body. Skip; we'll attach
            # the speaker to the next block's body.
            continue
        # The remaining lines are paragraph body. Look for timestamp prefix
        # on the first content line.
        first_content = lines[0]
        ts, after_ts = _parse_timestamp_prefix(first_content)
        if ts is not None:
            saw_timestamp = True
            start_t = ts
            lines[0] = after_ts.lstrip()
        else:
            start_t = 0.0
        body = " ".join(line.strip() for line in lines if line.strip())
        if not body:
            continue
        # Synthesise a single segment for this paragraph.
        paragraphs.append([(start_t, start_t + 1.0, body)])
        speakers.append(current_speaker)

    if not paragraphs:
        raise TranscriptParseError(
            "No transcript content found. The file may not be a "
            "Transcribr-produced transcript."
        )
    return {
        "paragraphs": paragraphs,
        "speakers": speakers,
        "title": title,
        "show_timestamp": saw_timestamp,
    }


def _parse_docx_transcript(path):
    """Parse a .docx file written by _write_docx().

    We rely on paragraph styles to identify what each paragraph is:
        - "Title" / first bold paragraph -> title
        - "SpeakerLabel" -> speaker line
        - "Transcript" -> body paragraph (timestamp + tab + body)
        - everything else -> ignored (disclaimer, blank spacers, etc.)

    Word's "Clear Formatting" or paste operations can lose the styles.
    For robustness, we also fall back to text-based heuristics if a
    paragraph has no recognised style: lines ending with ':' look like
    speakers, lines starting with [MM:SS] look like body."""
    try:
        from docx import Document
    except ImportError:
        raise TranscriptParseError(
            "Cannot read .docx files: the 'python-docx' package is not "
            "installed."
        )
    try:
        doc = Document(str(path))
    except Exception as e:
        raise TranscriptParseError(f"Could not open .docx file: {e}")

    title = None
    paragraphs = []
    speakers = []
    current_speaker = None
    saw_timestamp = False

    docx_paragraphs = list(doc.paragraphs)
    if not docx_paragraphs:
        raise TranscriptParseError("Document is empty.")

    # First paragraph is treated as the title if it's bold and doesn't
    # look like a speaker label or body line. We also accept the
    # explicit pattern we wrote: bold Courier New 12pt without a colon
    # at the end.
    first = docx_paragraphs[0]
    first_text = first.text.strip()
    if first_text:
        runs_bold = all(
            (r.bold or False) for r in first.runs if r.text.strip()
        ) if first.runs else False
        looks_speaker = first_text.endswith(":")
        has_ts = bool(_TS_PREFIX_RE.match(first_text))
        if runs_bold and not looks_speaker and not has_ts:
            title = first_text
            docx_paragraphs = docx_paragraphs[1:]

    for p in docx_paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if _DISCLAIMER_RE.match(text):
            continue
        style_name = ""
        try:
            style_name = (p.style.name or "").strip()
        except Exception:
            pass

        # Speaker label: explicit style, or bold + ends with ':'.
        if style_name == "SpeakerLabel" or (
            text.endswith(":") and not _TS_PREFIX_RE.match(text)
        ):
            label = text.rstrip(":").strip()
            current_speaker = None if label == UNATTRIBUTED_LABEL else label
            continue

        # Body paragraph: explicit style, or first chars are timestamp.
        # In python-docx, tabs in run text appear as "\t" characters.
        body_text = text.replace("\t", " ")
        ts, after_ts = _parse_timestamp_prefix(body_text)
        if ts is not None:
            saw_timestamp = True
            start_t = ts
            body = after_ts.strip()
        else:
            start_t = 0.0
            body = body_text.strip()
        if not body:
            continue
        paragraphs.append([(start_t, start_t + 1.0, body)])
        speakers.append(current_speaker)

    if not paragraphs:
        raise TranscriptParseError(
            "No transcript content found. The file may not be a "
            "Transcribr-produced transcript."
        )
    return {
        "paragraphs": paragraphs,
        "speakers": speakers,
        "title": title,
        "show_timestamp": saw_timestamp,
    }


# =====================================================================
# Whisper option metadata
# =====================================================================

# The canonical model choices shown in the UI. "large" and "turbo" are
# omitted: they are just aliases of "large-v3" and "large-v3-turbo"
# (identical weights), so listing them was a confusing duplicate. The
# engines still accept the alias names, and _MODEL_ALIASES normalises any
# stored/legacy value to its canonical form.
WHISPER_MODELS = [
    "tiny.en", "base.en", "small.en", "medium.en",
    "tiny", "base", "small", "medium",
    "large-v1", "large-v2", "large-v3", "large-v3-turbo",
]

_MODEL_ALIASES = {"large": "large-v3", "turbo": "large-v3-turbo"}


def _canonical_model(name):
    """Map an alias model name to its canonical equivalent (identical
    weights), leaving anything else untouched."""
    return _MODEL_ALIASES.get(name, name)


# The three choices most users should pick between, presented instead of
# the raw twelve-model list (which stays available behind a "show all
# models" toggle, and in full detail in the Models view). `model_en` is
# substituted when the language is English - the .en variants are
# slightly more accurate there. Sizes are the one-time download on
# first use.
MODEL_TIERS = [
    {"id": "draft", "label": "Quick draft",
     "model": "small", "model_en": "small.en", "size": "~500 MB",
     "note": "Fast first pass. Fine for clear speech; expect mistakes."},
    {"id": "standard", "label": "Standard", "recommended": True,
     "model": "large-v3-turbo", "model_en": "large-v3-turbo",
     "size": "~1.6 GB",
     "note": "Accurate and reasonably quick — right for most work."},
    {"id": "max", "label": "Maximum accuracy",
     "model": "large-v3", "model_en": "large-v3", "size": "~3 GB",
     "note": "Slowest. For difficult audio where every word matters."},
]


# ---- Engine detection ---------------------------------------------------
#
# Transcribr can drive three Whisper implementations. We probe each one's
# import at startup and only offer installed ones to the user. The engine
# key (e.g. "faster") is what flows through params; the display name is
# what appears in the dropdown.

def _macos_supports_mlx():
    """mlx (Apple's framework, which mlx-whisper depends on) refuses to
    import on macOS < 13.5. Return False on older versions so we don't
    offer the engine in the dropdown - find_spec sees the package but
    the runtime import would raise."""
    if sys.platform != "darwin":
        return False
    try:
        import platform
        ver = platform.mac_ver()[0]
        if not ver:
            return False
        parts = ver.split(".")
        major = int(parts[0])
        minor = int(parts[1]) if len(parts) > 1 else 0
        return (major, minor) >= (13, 5)
    except (ValueError, IndexError):
        return False


def _detect_engines():
    """Return a list of (key, display_name) for installed engines."""
    import importlib.util as _ilu
    engines = []
    if _ilu.find_spec("whisper") is not None:
        engines.append(("whisper", "OpenAI Whisper (reference)"))
    if _ilu.find_spec("faster_whisper") is not None:
        engines.append(("faster", "faster-whisper (CTranslate2)"))
    # mlx-whisper is Apple Silicon-only and needs macOS 13.5+.
    if (_macos_supports_mlx()
            and _ilu.find_spec("mlx_whisper") is not None):
        engines.append(("mlx", "mlx-whisper (Apple Silicon)"))
    return engines


AVAILABLE_ENGINES = _detect_engines()

# The "just do the right thing" engine choice, shown first in the
# dropdown and used as the default. It resolves at run time so the same
# settings file does the right thing when copied between machines.
ENGINE_AUTO_NAME = "Automatic (recommended)"


def resolve_engine_key(engine_name):
    """Map a settings engine display name to a runner key. The Automatic
    entry (and any stale/unknown name, e.g. after the engine it named
    was uninstalled from the Models view) resolves to the best installed
    engine: mlx-whisper on Apple Silicon, then faster-whisper, then the
    reference implementation."""
    for key, name in AVAILABLE_ENGINES:
        if name == engine_name:
            return key
    keys = [key for key, _ in AVAILABLE_ENGINES]
    for preferred in ("mlx", "faster", "whisper"):
        if preferred in keys:
            return preferred
    return keys[0] if keys else "whisper"


# mlx-whisper takes a HuggingFace repo path (since the runtime weights are
# converted to MLX format). Map our common short names to mlx-community
# repos. Fall back to a heuristic for anything not in the table.
_MLX_REPO_MAP = {
    "tiny": "mlx-community/whisper-tiny-mlx",
    "tiny.en": "mlx-community/whisper-tiny.en-mlx",
    "base": "mlx-community/whisper-base-mlx",
    "base.en": "mlx-community/whisper-base.en-mlx",
    "small": "mlx-community/whisper-small-mlx",
    "small.en": "mlx-community/whisper-small.en-mlx",
    "medium": "mlx-community/whisper-medium-mlx",
    "medium.en": "mlx-community/whisper-medium.en-mlx",
    "large": "mlx-community/whisper-large-v3-mlx",
    "large-v1": "mlx-community/whisper-large-v1-mlx",
    "large-v2": "mlx-community/whisper-large-v2-mlx",
    "large-v3": "mlx-community/whisper-large-v3-mlx",
    "turbo": "mlx-community/whisper-large-v3-turbo",
    "large-v3-turbo": "mlx-community/whisper-large-v3-turbo",
}


def _mlx_repo_for(model_name):
    return _MLX_REPO_MAP.get(
        model_name, f"mlx-community/whisper-{model_name}-mlx")


class _EngineNotAvailable(Exception):
    """Raised when the worker can't import the requested engine."""
    pass

# (Display name, code passed to whisper). None = auto-detect.
LANGUAGES = [
    ("Auto-detect", None),
    ("English", "en"),
    ("Chinese (Mandarin)", "zh"),
    ("Spanish", "es"),
    ("French", "fr"),
    ("German", "de"),
    ("Italian", "it"),
    ("Japanese", "ja"),
    ("Korean", "ko"),
    ("Portuguese", "pt"),
    ("Russian", "ru"),
    ("Arabic", "ar"),
    ("Hindi", "hi"),
    ("Vietnamese", "vi"),
    ("Indonesian", "id"),
    ("Dutch", "nl"),
    ("Turkish", "tr"),
    ("Polish", "pl"),
]


# =====================================================================
# Speaker detection (diarization)
# =====================================================================
#
# An optional post-pass that works out who is speaking: sherpa-onnx
# runs a small segmentation model (pyannote segmentation-3.0) plus a
# speaker-embedding model (WeSpeaker ResNet34-LM - the same embedding
# family pyannote's own 3.1 pipeline uses) and clusters the voices.
# Both models are ONNX and total ~33 MB, downloaded once on first use;
# onnxruntime is already in the dependency tree via faster-whisper.
#
# EXPERIMENTAL, and deliberately decoupled from paragraphify: the
# paragraph boundaries always come from the programmatic heuristics
# alone, and diarization only SUGGESTS a speaker label for each
# finished paragraph (majority voice by time overlap). Any paragraph
# whose attribution is doubtful is left unlabelled for the reviewer
# rather than guessed.

_DIARIZE_SEGMENTATION = {
    "url": ("https://github.com/k2-fsa/sherpa-onnx/releases/download/"
            "speaker-segmentation-models/"
            "sherpa-onnx-pyannote-segmentation-3-0.tar.bz2"),
    "sha256": ("24615ee884c897d9d2ba09bb4d30da6b"
               "b1b15e685065962db5b02e76e4996488"),
    "archive_member": "model.onnx",
    "target": "pyannote-segmentation-3.0.onnx",
    "label": "speaker segmentation model (~7 MB)",
}

# The selectable voice-matching (speaker embedding) models. All run
# through the identical sherpa-onnx pipeline; they differ in how they
# "hear" voices, so when one merges two similar speakers another may
# separate them. Each downloads once, on first use.
DIARIZE_VOICE_MODELS = [
    {"id": "wespeaker",
     "label": "WeSpeaker ResNet34",
     "note": "The default — the same voiceprint family pyannote's own "
             "pipeline uses.",
     "size": "~27 MB",
     "url": ("https://github.com/k2-fsa/sherpa-onnx/releases/download/"
             "speaker-recongition-models/"      # (sic - upstream tag)
             "wespeaker_en_voxceleb_resnet34_LM.onnx"),
     "sha256": ("e9848563da86f263117134dfd7ad63c9"
                "2355b37de492b55e325400c9d9c39012"),
     "archive_member": None,
     "target": "wespeaker-voxceleb-resnet34-LM.onnx"},
    {"id": "campplus",
     "label": "CAM++ (3D-Speaker)",
     "note": "A different architecture that is markedly faster and "
             "sometimes separates similar voices the default merges.",
     "size": "~30 MB",
     "url": ("https://github.com/k2-fsa/sherpa-onnx/releases/download/"
             "speaker-recongition-models/"
             "3dspeaker_speech_campplus_sv_en_voxceleb_16k.onnx"),
     "sha256": ("357a834f702b80161e5b981182c038e1"
                "8553c1f2ca752ed6cec2052365d4129b"),
     "archive_member": None,
     "target": "campplus-voxceleb.onnx"},
    {"id": "titanet",
     "label": "TitaNet-S (NVIDIA)",
     "note": "A third opinion — worth trying when the other two "
             "disagree with what you hear.",
     "size": "~38 MB",
     "url": ("https://github.com/k2-fsa/sherpa-onnx/releases/download/"
             "speaker-recongition-models/nemo_en_titanet_small.onnx"),
     "sha256": ("ad4a1802485d8b34c722d2a9d0424966"
                "2f2ece5d28a7a039063ca22f515a789e"),
     "archive_member": None,
     "target": "titanet-small.onnx"},
]

DIARIZE_DEFAULT_VOICE_MODEL = "wespeaker"


def _voice_model_spec(voice_id):
    for spec in DIARIZE_VOICE_MODELS:
        if spec["id"] == voice_id:
            return spec
    return next(spec for spec in DIARIZE_VOICE_MODELS
                if spec["id"] == DIARIZE_DEFAULT_VOICE_MODEL)

# A paragraph is labelled only when its majority voice wins clearly
# among the time the diarizer attributed (MIN_MAJORITY of it), AND the
# diarizer attributed a meaningful share of the paragraph's span at all
# (MIN_COVERAGE - deliberately low, because engines pad segment spans
# with silence). Wrong labels are worse than missing ones.
_DIARIZE_MIN_MAJORITY = 0.6
_DIARIZE_MIN_COVERAGE = 0.3
# Clustering distance threshold when the speaker count is unknown
# (sherpa-onnx's documented default for these models).
_DIARIZE_CLUSTER_THRESHOLD = 0.5


class DiarizationUnavailable(Exception):
    """Speaker detection can't run; the message says why and the run
    continues without labels."""
    pass


def _diarize_models_dir():
    return _config_dir() / "models" / "diarization"


def diarize_models_ready(voice_id=DIARIZE_DEFAULT_VOICE_MODEL):
    d = _diarize_models_dir()
    return ((d / _DIARIZE_SEGMENTATION["target"]).exists()
            and (d / _voice_model_spec(voice_id)["target"]).exists())


def _download_one_model(spec, q, cancel_event):
    """Fetch, verify (sha256) and install one model file atomically.
    Progress goes to the run log."""
    import hashlib
    import tarfile
    import tempfile
    import urllib.request

    target = _diarize_models_dir() / spec["target"]
    if target.exists():
        return
    target.parent.mkdir(parents=True, exist_ok=True)
    q.put(("log", f"Downloading {spec['label']}...\n"))

    with tempfile.TemporaryDirectory(dir=str(target.parent)) as tmpdir:
        fetched = Path(tmpdir) / "download"
        digest = hashlib.sha256()
        last_decile = -1
        req = urllib.request.Request(
            spec["url"], headers={"User-Agent": f"Transcribr/{__version__}"})
        with urllib.request.urlopen(req, timeout=60) as resp, \
                open(fetched, "wb") as out:
            total = int(resp.headers.get("Content-Length") or 0)
            done = 0
            while True:
                if cancel_event is not None and cancel_event.is_set():
                    raise _CancelledByUser()
                chunk = resp.read(65536)
                if not chunk:
                    break
                out.write(chunk)
                digest.update(chunk)
                done += len(chunk)
                if total:
                    decile = int(done * 10 / total)
                    if decile > last_decile:
                        last_decile = decile
                        q.put(("log", f"  ... {min(100, decile * 10)}%\n"))

        if digest.hexdigest() != spec["sha256"]:
            raise DiarizationUnavailable(
                f"The downloaded {spec['label']} failed its integrity "
                "check. Please try again; if it keeps failing, the "
                "upstream file may have changed.")

        if spec["archive_member"]:
            with tarfile.open(fetched, "r:bz2") as tar:
                member = next(
                    (m for m in tar.getmembers()
                     if m.isfile() and Path(m.name).name
                     == spec["archive_member"]),
                    None)
                if member is None:
                    raise DiarizationUnavailable(
                        f"{spec['archive_member']} missing from the "
                        "downloaded archive.")
                src = tar.extractfile(member)
                extracted = Path(tmpdir) / "extracted"
                with open(extracted, "wb") as out:
                    shutil.copyfileobj(src, out)
                os.replace(extracted, target)
        else:
            os.replace(fetched, target)


def ensure_diarize_models(q, cancel_event=None,
                          voice_id=DIARIZE_DEFAULT_VOICE_MODEL):
    """Download any missing diarization models for `voice_id`. Raises
    DiarizationUnavailable (network trouble etc.) or _CancelledByUser."""
    import urllib.error
    voice = _voice_model_spec(voice_id)
    for spec in (_DIARIZE_SEGMENTATION,
                 {**voice,
                  "label": f"voice-matching model "
                           f"'{voice['label']}' ({voice['size']})"}):
        try:
            _download_one_model(spec, q, cancel_event)
        except (urllib.error.URLError, OSError, TimeoutError) as e:
            raise DiarizationUnavailable(
                f"Could not download the {spec['label']}:\n  {e}\n"
                "Check your internet connection and run again - the "
                "download is only needed once.")


def _run_diarization(audio, num_speakers, q, cancel_event=None,
                     voice_id=DIARIZE_DEFAULT_VOICE_MODEL,
                     threshold=None):
    """Run sherpa-onnx speaker diarization over mono 16 kHz float32
    samples. Returns [(start_seconds, end_seconds, speaker_int)] sorted
    by start time. `num_speakers` <= 0 means detect automatically using
    `threshold` (smaller = readier to hear two similar voices as
    different people)."""
    try:
        import sherpa_onnx
    except ImportError:
        raise DiarizationUnavailable(
            "Speaker detection needs the 'sherpa-onnx' package, which "
            "is not installed in this Python.\nInstall it with:\n"
            "  pip install sherpa-onnx")

    if threshold is None:
        threshold = _DIARIZE_CLUSTER_THRESHOLD
    threshold = min(0.9, max(0.2, float(threshold)))
    d = _diarize_models_dir()
    voice = _voice_model_spec(voice_id)
    q.put(("log", f"  voice-matching model: {voice['label']}\n"))
    config = sherpa_onnx.OfflineSpeakerDiarizationConfig(
        segmentation=sherpa_onnx.OfflineSpeakerSegmentationModelConfig(
            pyannote=sherpa_onnx.OfflineSpeakerSegmentationPyannoteModelConfig(
                model=str(d / _DIARIZE_SEGMENTATION["target"])),
        ),
        embedding=sherpa_onnx.SpeakerEmbeddingExtractorConfig(
            model=str(d / voice["target"])),
        clustering=sherpa_onnx.FastClusteringConfig(
            num_clusters=(int(num_speakers) if num_speakers
                          and num_speakers > 0 else -1),
            threshold=threshold),
        min_duration_on=0.3,
        min_duration_off=0.5,
    )
    sd = sherpa_onnx.OfflineSpeakerDiarization(config)

    last_decile = [-1]

    def progress(num_processed_chunks, num_total_chunks):
        if num_total_chunks:
            decile = int(num_processed_chunks * 10 / num_total_chunks)
            if decile > last_decile[0]:
                last_decile[0] = decile
                q.put(("log",
                       f"  ... listening for voices "
                       f"{min(100, decile * 10)}%\n"))
        return 0

    t0 = time.time()
    result = sd.process(audio, callback=progress).sort_by_start_time()
    turns = [(float(r.start), float(r.end), int(r.speaker))
             for r in result]
    n = len({spk for _, _, spk in turns})
    q.put(("log",
           f"  found {n} speaker(s) in {time.time() - t0:.1f}s\n"))
    return turns


# ---- Suggesting labels for finished paragraphs ------------------------

def label_paragraphs(paragraphs, turns):
    """Suggest a review-slot letter ("1".."9" or None) for each
    already-formed paragraph from the diarized turns - the paragraphs
    themselves are never reshaped.

    Per paragraph, each speaker is credited with the time their turns
    overlap the paragraph's spoken segments. The majority speaker gets
    the label only when it clearly wins among the attributed time
    (_DIARIZE_MIN_MAJORITY) and the diarizer attributed a meaningful
    share of the paragraph at all (_DIARIZE_MIN_COVERAGE); anything
    doubtful stays unlabelled for the reviewer. Surviving speakers are
    ranked by credited time, the busiest nine keep slots, and
    numbering follows order of first appearance - the first voice
    heard becomes Speaker 1."""
    eligible = []       # per paragraph: speaker id or None
    for para in paragraphs:
        times, span = {}, 0.0
        for seg in para:
            s_start, s_end = seg[0], seg[1]
            span += max(0.0, s_end - s_start)
            for t_start, t_end, spk in turns:
                overlap = min(s_end, t_end) - max(s_start, t_start)
                if overlap > 0:
                    times[spk] = times.get(spk, 0.0) + overlap
        attributed = sum(times.values())
        if not times or span <= 0 \
                or attributed / span < _DIARIZE_MIN_COVERAGE:
            eligible.append(None)
            continue
        best = max(times, key=times.get)
        if times[best] / attributed < _DIARIZE_MIN_MAJORITY:
            eligible.append(None)
            continue
        eligible.append(best)

    totals = {}
    for spk, para in zip(eligible, paragraphs):
        if spk is None:
            continue
        totals[spk] = (totals.get(spk, 0.0)
                       + max(0.0, para[-1][1] - para[0][0]))
    keep = set(sorted(totals, key=totals.get, reverse=True)
               [:TranscriptModel.MAX_SPEAKERS])
    letters = {}
    for spk in (s for s in eligible if s is not None and s in keep):
        if spk not in letters:
            letters[spk] = str(len(letters) + 1)

    return [letters.get(spk) if spk is not None else None
            for spk in eligible]


# =====================================================================
# Background worker
# =====================================================================

class _QueueWriter:
    """File-like that forwards writes to a queue as ('log', text) messages."""
    def __init__(self, q):
        self.q = q

    def write(self, text):
        if text:
            self.q.put(("log", text))

    def flush(self):
        pass


class _CancelledByUser(Exception):
    """Raised inside whisper to abort transcription mid-run."""
    pass


# Approximate on-disk download sizes per model, used only for first-run
# messaging before the real byte totals arrive from the downloader.
_MODEL_APPROX_MB = {
    "tiny.en": 75, "tiny": 75,
    "base.en": 145, "base": 145,
    "small.en": 480, "small": 480,
    "medium.en": 1500, "medium": 1500,
    "large-v1": 2900, "large-v2": 2900, "large-v3": 2900, "large": 2900,
    "turbo": 1600, "large-v3-turbo": 1600,
}


def _humanize_bytes(n):
    """Compact human-readable byte count, e.g. '1.6 GB', '480 MB'."""
    n = float(n or 0)
    if n >= 1024 ** 3:
        return f"{n / 1024 ** 3:.1f} GB"
    if n >= 1024 ** 2:
        return f"{n / 1024 ** 2:.0f} MB"
    if n >= 1024:
        return f"{n / 1024:.0f} KB"
    return f"{int(n)} B"


# '[mm:ss.sss --> mm:ss.sss] text' - the line openai-whisper and
# mlx-whisper both print when verbose=True. Parsed to drive the ETA/
# progress readout.
_SEGMENT_LINE_RE = re.compile(
    r"\[(\d+):(\d+(?:\.\d+)?)\s*-->\s*(\d+):(\d+(?:\.\d+)?)\]\s*(.*)")


class _ProgressWriter:
    """File-like that mirrors an engine's verbose stdout/stderr to the log
    queue while parsing segment lines to emit 'eta' progress events.

    Both openai-whisper and mlx-whisper print '[mm:ss --> mm:ss] text'
    lines as they decode (with verbose=True), so a single parser drives
    the progress bar for both. If `captured` is a list, parsed segments
    are appended to it - openai-whisper relies on that to salvage a
    partial transcript when the user hits Stop mid-run."""

    def __init__(self, q, audio_duration, transcribe_start, captured=None):
        self.q = q
        self._buf = ""
        self.audio_duration = audio_duration
        self.transcribe_start = transcribe_start
        self.captured = captured

    def write(self, text):
        if not text:
            return
        self.q.put(("log", text))
        self._buf += text
        while "\n" in self._buf:
            line, self._buf = self._buf.split("\n", 1)
            m = _SEGMENT_LINE_RE.search(line)
            if not m:
                continue
            sm, ss, em, es, body = m.groups()
            start = int(sm) * 60 + float(ss)
            end = int(em) * 60 + float(es)
            body = body.strip()
            if not body:
                continue
            if self.captured is not None:
                self.captured.append((start, end, body))
            if self.audio_duration and end >= 5.0:
                wall = time.time() - self.transcribe_start
                if wall > 0:
                    speed = end / wall
                    remaining = max(0.0, self.audio_duration - end)
                    eta = remaining / speed if speed > 0 else 0
                    self.q.put(("eta", {
                        "audio_done": end,
                        "audio_total": self.audio_duration,
                        "wall_elapsed": wall,
                        "eta_seconds": eta,
                        "speed": speed,
                    }))

    def flush(self):
        pass


class _DownloadMonitor:
    """Surfaces model downloads that would otherwise look like a frozen
    'Loading model...'.

    Every engine's downloader (openai-whisper's own fetch, and the
    huggingface_hub fetch behind faster-whisper and mlx-whisper) reports
    progress through a tqdm byte-scale bar. This context manager patches
    tqdm so those bars are forwarded to `on_progress(info)` - a dict of
    {model, downloaded, total, speed}. The transcription worker routes
    that onto its queue as ('download', ...) messages; the standalone
    model manager republishes it as a `model_progress` SSE event. A
    one-line announcement (and completion note) goes to `on_log(text)`
    when supplied. Non-download bars (e.g. tiny config files) are ignored
    via a minimum-size filter.

    It also honours cancel_event: cancelling during a download raises
    _CancelledByUser, aborting the fetch instead of forcing the user to
    wait for gigabytes."""

    _MIN_TOTAL = 1024 * 1024  # ignore sub-1 MB bars (configs, tokenizers)

    def __init__(self, cancel_event, model_name, *,
                 on_progress, on_log=None):
        self.cancel_event = cancel_event
        self.model_name = model_name
        self.on_progress = on_progress
        self.on_log = on_log or (lambda _text: None)
        self._bars = {}            # id(bar) -> (downloaded, total)
        self._tqdm = None
        self._orig_update = None
        self._orig_close = None
        self._last_emit = 0.0
        self._last_done = 0
        self._last_speed_t = 0.0
        self._announced = False

    def _is_byte_bar(self, bar):
        unit = getattr(bar, "unit", "") or ""
        total = getattr(bar, "total", None) or 0
        return "B" in unit and total >= self._MIN_TOTAL

    def _snapshot(self):
        downloaded = sum(d for d, _ in self._bars.values())
        total = sum(t for _, t in self._bars.values())
        return downloaded, total

    def _emit(self, force=False):
        now = time.time()
        if not force and now - self._last_emit < 0.25:
            return
        downloaded, total = self._snapshot()
        if total <= 0:
            return
        speed = 0.0
        if self._last_speed_t:
            dt = now - self._last_speed_t
            if dt > 0:
                speed = max(0.0, (downloaded - self._last_done) / dt)
        self._last_speed_t = now
        self._last_done = downloaded
        self._last_emit = now
        self.on_progress({
            "model": self.model_name,
            "downloaded": downloaded,
            "total": total,
            "speed": speed,
        })

    def __enter__(self):
        try:
            import tqdm
        except ImportError:
            return self          # no tqdm - degrade to no-op
        self._tqdm = tqdm
        monitor = self
        self._orig_update = tqdm.tqdm.update
        self._orig_close = tqdm.tqdm.close

        def patched_update(bar, n=1):
            result = monitor._orig_update(bar, n)
            if monitor.cancel_event.is_set():
                raise _CancelledByUser()
            if monitor._is_byte_bar(bar):
                monitor._bars[id(bar)] = (
                    getattr(bar, "n", 0) or 0, getattr(bar, "total", 0) or 0)
                if not monitor._announced:
                    monitor._announced = True
                    approx = _MODEL_APPROX_MB.get(monitor.model_name)
                    size = f" (~{approx / 1000:.1f} GB)" if approx and \
                        approx >= 1000 else (f" (~{approx} MB)" if approx
                                             else "")
                    monitor.on_log(
                        f"Downloading model '{monitor.model_name}'{size} - "
                        "first use only, then it's cached locally...\n")
                monitor._emit()
            return result

        def patched_close(bar):
            if monitor._is_byte_bar(bar):
                total = getattr(bar, "total", 0) or 0
                monitor._bars[id(bar)] = (total, total)
                monitor._emit(force=True)
            return monitor._orig_close(bar)

        tqdm.tqdm.update = patched_update
        tqdm.tqdm.close = patched_close
        return self

    def __exit__(self, exc_type, exc, tb):
        if self._tqdm is not None:
            self._tqdm.tqdm.update = self._orig_update
            self._tqdm.tqdm.close = self._orig_close
        if self._announced and exc_type is None:
            _, total = self._snapshot()
            self.on_log(f"  download complete ({_humanize_bytes(total)})\n")
        return False


def _extract_word_conf(result):
    """Flatten per-word (start, end, word, probability) tuples from a
    Whisper-style result dict. Returns [] when no word-level data is present
    (i.e. word timestamps weren't requested). Tolerant of missing keys so it
    works across the openai-whisper, faster-whisper and mlx-whisper result
    shapes (all of which expose segment['words'] as dicts after we
    normalise faster-whisper's objects)."""
    if not result:
        return []
    out = []
    for seg in result.get("segments", []):
        if not isinstance(seg, dict):
            continue
        for w in seg.get("words") or []:
            try:
                start = float(w.get("start"))
                end = float(w.get("end"))
            except (TypeError, ValueError, AttributeError):
                continue
            text = (w.get("word") or "") if isinstance(w, dict) else ""
            prob = w.get("probability") if isinstance(w, dict) else None
            try:
                prob = float(prob) if prob is not None else None
            except (TypeError, ValueError):
                prob = None
            if text.strip():
                out.append((start, end, text, prob))
    return out


def transcribe_worker(params, q, cancel_event):
    """Background-thread entry point. Dispatches to the chosen engine's
    runner, then handles the common post-processing: speaker detection,
    paragraphify, review-or-direct-save, extra output formats."""
    engine = params.get("engine", "whisper")
    runner = {
        "whisper": _run_openai_whisper,
        "faster": _run_faster_whisper,
        "mlx": _run_mlx_whisper,
    }.get(engine, _run_openai_whisper)

    # Decode once up front (PyAV). Every engine accepts the raw samples,
    # which removes their need for an ffmpeg binary, and the speaker
    # detector reuses the same array. On failure the runners receive the
    # file path and decode with whatever their engine ships.
    audio = decode_audio_16k(params["input"])
    if audio is not None:
        params["_audio"] = audio
        duration = len(audio) / AUDIO_SAMPLE_RATE
        q.put(("log", f"Decoded audio: {_format_duration(duration)}\n"))
        if not params.get("audio_duration"):
            params["audio_duration"] = duration
    elif _have_pyav():
        # PyAV is present but couldn't decode the file. The engine's own
        # loader would fail the same way with a raw ffmpeg traceback, so
        # explain the actual problem instead - the usual cause is a
        # non-audio file wearing an audio extension.
        reason = describe_non_audio_file(params["input"])
        if reason:
            q.put(("error",
                   f"Cannot transcribe {Path(params['input']).name}: "
                   f"{reason}.\n\n"
                   "Check that the file really is an audio or video "
                   "recording - if it came out of an export or download, "
                   "the export may have produced the wrong file. Opening "
                   "it in a media player is a quick way to confirm."))
            return

    try:
        segments, result, used_partial = runner(params, q, cancel_event)
    except _EngineNotAvailable as e:
        q.put(("error", str(e)))
        return
    except _CancelledByUser:
        q.put(("cancelled", None))
        return
    except Exception as e:
        q.put(("error",
               f"{type(e).__name__}: {e}\n\n"
               + traceback.format_exc()))
        return

    if not segments:
        if used_partial:
            q.put(("cancelled",
                   "Stopped before any segments were transcribed - "
                   "nothing to save."))
        else:
            q.put(("error", "No speech was detected in the file."))
        return

    paragraphs, para_letters = _paragraphs_with_speakers(
        segments, result, params, q, cancel_event)
    out_path = Path(params["output"])
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if params.get("review_before_save"):
        word_conf = (_extract_word_conf(result)
                     if params.get("highlight_confidence") else None)
        q.put((
            "paragraphs_ready",
            {
                "paragraphs": paragraphs,
                "out_path": str(out_path),
                "show_timestamp": params.get("show_timestamp", True),
                "title": params.get("title"),
                "output_format": params.get("output_format", "txt"),
                "used_partial": used_partial,
                "result": result if not used_partial else None,
                "extra_formats": params.get("extra_formats") or [],
                "word_conf": word_conf,
                "audio_path": params["input"],
                "preset_speakers": para_letters,
                "diarized": para_letters is not None,
            },
        ))
        return

    # Unattended save: detected speakers go in as generic names the
    # user can refine later from the Library's review pane.
    speaker_names = None
    if para_letters is not None:
        speaker_names = [f"Speaker {L}" if L else None
                         for L in para_letters]
    try:
        write_paragraphs_to_file(
            paragraphs, out_path,
            show_timestamp=params.get("show_timestamp", True),
            title=params.get("title"),
            output_format=params.get("output_format", "txt"),
            speakers=speaker_names,
        )
    except ImportError as e:
        q.put(("error", str(e)))
        return

    prefix = "Partial: w" if used_partial else "W"
    q.put(("log",
           f"{prefix}rote {len(paragraphs)} paragraphs "
           f"(from {len(segments)} segments)\n"
           f"to: {out_path}\n"))

    extra_formats = params.get("extra_formats") or []
    if result and extra_formats:
        _write_extra_formats(result, out_path, extra_formats, q)

    q.put(("done", str(out_path)))


def _paragraphs_with_speakers(segments, result, params, q, cancel_event):
    """Group segments into paragraphs with the programmatic heuristics,
    then - when the experimental speaker detection is on - suggest a
    label for each paragraph. The paragraphs are identical either way:
    detection only labels them, never reshapes them, and any failure
    just leaves the labels off with a logged warning."""
    words = _extract_word_conf(result) if result else []
    paragraphs = paragraphify(segments, params["gap"], words=words)
    if not params.get("diarize"):
        return paragraphs, None
    try:
        audio = params.get("_audio")
        if audio is None:
            audio = decode_audio_16k(params["input"])
        if audio is None:
            raise DiarizationUnavailable(
                "the audio could not be decoded for speaker detection "
                "(the 'av' package is required).")
        voice_id = params.get("diarize_model") or DIARIZE_DEFAULT_VOICE_MODEL
        if not diarize_models_ready(voice_id):
            ensure_diarize_models(q, cancel_event, voice_id)
        q.put(("log", "\nIdentifying speakers (experimental)...\n"))
        turns = _run_diarization(
            audio, params.get("num_speakers") or 0, q, cancel_event,
            voice_id=voice_id,
            threshold=params.get("diarize_threshold"))
        para_letters = label_paragraphs(paragraphs, turns)
        labelled = sum(1 for L in para_letters if L)
        q.put(("log",
               f"  labelled {labelled} of {len(paragraphs)} "
               "paragraphs (uncertain ones are left for review)\n"))
        return paragraphs, para_letters
    except _CancelledByUser:
        q.put(("log", "\n[Speaker detection skipped - stopped by "
                      "user]\n"))
    except DiarizationUnavailable as e:
        q.put(("log", f"\nSpeaker detection skipped: {e}\n"))
    except Exception as e:
        _log("diarization failed:\n" + traceback.format_exc())
        q.put(("log",
               f"\nSpeaker detection failed - continuing without it.\n"
               f"  {type(e).__name__}: {e}\n"))
    return paragraphs, None


def _run_openai_whisper(params, q, cancel_event):
    """Run the reference OpenAI Whisper implementation. Returns
    (segments_list, result_dict, used_partial)."""
    captured_segments = []
    used_partial = False

    q.put(("log", "Importing whisper...\n"))
    try:
        import whisper
        import tqdm
    except ImportError as e:
        raise _EngineNotAvailable(
            "Could not load openai-whisper:\n"
            f"  {type(e).__name__}: {e}\n\n"
            "Install or repair it with:\n  pip install --upgrade "
            "openai-whisper\n\n"
            "If you are using a virtual environment, make sure this app "
            "is launched from within that environment.")

    q.put(("log", f"Loading model '{params['model']}'...\n"))
    q.put(("status", {"stage": "loading",
                      "text": f"Loading model '{params['model']}'..."}))
    t0 = time.time()
    with _DownloadMonitor(
            cancel_event, params["model"],
            on_progress=lambda d: q.put(("download", d)),
            on_log=lambda t: q.put(("log", t))):
        model = whisper.load_model(params["model"])
    q.put(("log", f"  loaded in {time.time() - t0:.1f}s\n\n"))

    if cancel_event.is_set():
        return [], None, True

    q.put(("log", f"Transcribing {Path(params['input']).name}...\n"))
    q.put(("status", {"stage": "transcribing",
                      "text": "Transcribing..."}))
    t0 = time.time()

    kwargs = dict(
        language=params["language"],
        task=params["task"],
        temperature=params["temperature"],
        compression_ratio_threshold=params["compression_ratio_threshold"],
        logprob_threshold=params["logprob_threshold"],
        no_speech_threshold=params["no_speech_threshold"],
        condition_on_previous_text=params["condition_on_previous_text"],
        word_timestamps=params["word_timestamps"],
        verbose=True,
    )
    if params.get("beam_size") and params["beam_size"] > 1:
        kwargs["beam_size"] = params["beam_size"]
    if params.get("best_of"):
        kwargs["best_of"] = params["best_of"]
    if params.get("initial_prompt"):
        kwargs["initial_prompt"] = params["initial_prompt"]

    audio_duration = params.get("audio_duration")
    writer = _ProgressWriter(q, audio_duration, t0,
                             captured=captured_segments)

    original_update = tqdm.tqdm.update

    def cancelling_update(self_, n=1):
        if cancel_event.is_set():
            raise _CancelledByUser()
        return original_update(self_, n)

    audio_input = params.get("_audio")
    if audio_input is None:
        audio_input = params["input"]

    tqdm.tqdm.update = cancelling_update
    try:
        with contextlib.redirect_stdout(writer), \
             contextlib.redirect_stderr(writer):
            try:
                result = model.transcribe(audio_input, **kwargs)
            except _CancelledByUser:
                used_partial = True
                result = {"segments": []}
    finally:
        tqdm.tqdm.update = original_update

    if used_partial:
        q.put(("log",
               "\n[Stopped by user - saving what was transcribed so far]\n"))
        segments = list(captured_segments)
    else:
        q.put(("log", f"\n  transcribed in {time.time() - t0:.1f}s\n\n"))
        segments = [
            (float(s["start"]), float(s["end"]),
             (s.get("text") or "").strip())
            for s in result.get("segments", [])
            if (s.get("text") or "").strip()
        ]

    return segments, result, used_partial


def _run_faster_whisper(params, q, cancel_event):
    """Run faster-whisper (CTranslate2-based). Same model names as the
    reference engine. Cancellation between segments via cancel_event."""
    q.put(("log", "Importing faster-whisper...\n"))
    try:
        from faster_whisper import WhisperModel
    except ImportError as e:
        raise _EngineNotAvailable(
            "Could not load faster-whisper:\n"
            f"  {type(e).__name__}: {e}\n\n"
            "Install or repair it with:\n  pip install --upgrade "
            "faster-whisper")

    q.put(("log", f"Loading model '{params['model']}'...\n"))
    q.put(("status", {"stage": "loading",
                      "text": f"Loading model '{params['model']}'..."}))
    t0 = time.time()
    # device="auto" picks cuda if available, else cpu.
    # compute_type="auto" picks float16 on cuda, int8 on cpu - both fast,
    # both essentially indistinguishable from float32 for our use.
    # The download monitor surfaces the huggingface_hub fetch on first use.
    with _DownloadMonitor(
            cancel_event, params["model"],
            on_progress=lambda d: q.put(("download", d)),
            on_log=lambda t: q.put(("log", t))):
        model = WhisperModel(
            params["model"], device="auto", compute_type="auto")
    q.put(("log", f"  loaded in {time.time() - t0:.1f}s\n\n"))

    if cancel_event.is_set():
        return [], None, True

    q.put(("log", f"Transcribing {Path(params['input']).name}...\n"))
    q.put(("status", {"stage": "transcribing",
                      "text": "Transcribing..."}))
    t0 = time.time()

    kwargs = dict(
        language=params.get("language"),
        task=params.get("task", "transcribe"),
        temperature=params.get("temperature", 0.0),
        compression_ratio_threshold=params.get(
            "compression_ratio_threshold"),
        log_prob_threshold=params.get("logprob_threshold"),
        no_speech_threshold=params.get("no_speech_threshold"),
        condition_on_previous_text=params.get(
            "condition_on_previous_text", True),
        word_timestamps=params.get("word_timestamps", False),
    )
    if params.get("beam_size") and params["beam_size"] > 1:
        kwargs["beam_size"] = params["beam_size"]
    if params.get("best_of"):
        kwargs["best_of"] = params["best_of"]
    if params.get("initial_prompt"):
        kwargs["initial_prompt"] = params["initial_prompt"]

    audio_duration = params.get("audio_duration")
    captured = []
    seg_dicts = []
    used_partial = False

    audio_input = params.get("_audio")
    if audio_input is None:
        audio_input = params["input"]
    segments_iter, info = model.transcribe(audio_input, **kwargs)
    if audio_duration is None and getattr(info, "duration", None):
        audio_duration = float(info.duration)

    for seg in segments_iter:
        if cancel_event.is_set():
            used_partial = True
            q.put(("log",
                   "\n[Stopped by user - saving what was transcribed so "
                   "far]\n"))
            break
        text = (seg.text or "").strip()
        if not text:
            continue
        start = float(seg.start)
        end = float(seg.end)
        captured.append((start, end, text))
        # Normalise faster-whisper's Word objects to the dict shape the rest
        # of the code (extra-format writers, _extract_word_conf) expects.
        seg_words = [
            {"word": w.word, "start": w.start, "end": w.end,
             "probability": getattr(w, "probability", None)}
            for w in (getattr(seg, "words", None) or [])
        ]
        seg_dicts.append(
            {"start": start, "end": end, "text": text, "words": seg_words})
        q.put(("log",
               f"[{format_timestamp(start)} --> {format_timestamp(end)}]  "
               f"{text}\n"))
        if audio_duration and end >= 5.0:
            wall = time.time() - t0
            if wall > 0:
                speed = end / wall
                remaining = max(0.0, audio_duration - end)
                eta = remaining / speed if speed > 0 else 0
                q.put(("eta", {
                    "audio_done": end,
                    "audio_total": audio_duration,
                    "wall_elapsed": wall,
                    "eta_seconds": eta,
                    "speed": speed,
                }))

    if not used_partial:
        q.put(("log", f"\n  transcribed in {time.time() - t0:.1f}s\n\n"))

    # Build a whisper-compatible result for the extra-format writers and the
    # word-confidence extractor.
    result = {
        "segments": seg_dicts,
        "language": getattr(info, "language", None)
                    or params.get("language") or "",
    }
    return captured, result, used_partial


def _run_mlx_whisper(params, q, cancel_event):
    """Run mlx-whisper (Apple Silicon). No mid-run cancellation - the
    runtime doesn't expose a hook for it - so Stop only applies after the
    transcription finishes."""
    q.put(("log", "Importing mlx-whisper...\n"))
    try:
        import mlx_whisper
    except ImportError as e:
        # find_spec sees the package's metadata but the actual import
        # may fail because a dependency (typically Apple's `mlx`
        # framework) failed to install or isn't supported on this OS
        # version. Surface the real error so the user can act on it.
        raise _EngineNotAvailable(
            "Could not load mlx-whisper:\n"
            f"  {type(e).__name__}: {e}\n\n"
            "If the installer reported a warning about mlx, that's the "
            "underlying cause. To debug or repair:\n"
            "  cd '~/Library/Application Support/Transcribr'\n"
            "  venv/bin/pip install --upgrade mlx mlx-whisper\n\n"
            "Note: mlx-whisper requires macOS 13.5+ on Apple Silicon "
            "(M1, M2, M3, M4).")

    repo = _mlx_repo_for(params["model"])
    q.put(("log",
           f"Using model '{params['model']}' (mlx repo: {repo})\n"))
    q.put(("log",
           "Note: mlx-whisper doesn't support mid-run cancellation; the "
           "Stop button only takes effect after the run completes.\n\n"))

    if cancel_event.is_set():
        return [], None, True

    kwargs = dict(
        path_or_hf_repo=repo,
        language=params.get("language"),
        task=params.get("task", "transcribe"),
        temperature=params.get("temperature", 0.0),
        condition_on_previous_text=params.get(
            "condition_on_previous_text", True),
        word_timestamps=params.get("word_timestamps", False),
        # verbose=True so mlx prints '[mm:ss --> mm:ss] text' segment lines
        # as it decodes; _ProgressWriter parses them to drive the progress
        # bar (mlx exposes no per-segment callback otherwise).
        verbose=True,
    )
    if params.get("compression_ratio_threshold") is not None:
        kwargs["compression_ratio_threshold"] = params[
            "compression_ratio_threshold"]
    if params.get("logprob_threshold") is not None:
        kwargs["logprob_threshold"] = params["logprob_threshold"]
    if params.get("no_speech_threshold") is not None:
        kwargs["no_speech_threshold"] = params["no_speech_threshold"]
    if params.get("initial_prompt"):
        kwargs["initial_prompt"] = params["initial_prompt"]

    audio_input = params.get("_audio")
    if audio_input is None:
        audio_input = params["input"]

    q.put(("log", f"Transcribing {Path(params['input']).name}...\n"))
    q.put(("status", {"stage": "loading",
                      "text": f"Loading model '{params['model']}'..."}))
    t0 = time.time()
    audio_duration = params.get("audio_duration")
    writer = _ProgressWriter(q, audio_duration, t0)
    # mlx downloads the model inside transcribe() on first use, so the
    # download monitor wraps the whole call. Only stdout is redirected -
    # the huggingface_hub download bar lives on stderr and is reported via
    # the monitor instead of spamming the log with carriage returns.
    with _DownloadMonitor(
            cancel_event, params["model"],
            on_progress=lambda d: q.put(("download", d)),
            on_log=lambda t: q.put(("log", t))):
        with contextlib.redirect_stdout(writer):
            result = mlx_whisper.transcribe(audio_input, **kwargs)
    q.put(("log", f"\n  transcribed in {time.time() - t0:.1f}s\n\n"))

    captured = [
        (float(s["start"]), float(s["end"]), (s.get("text") or "").strip())
        for s in result.get("segments", [])
        if (s.get("text") or "").strip()
    ]
    used_partial = cancel_event.is_set()  # cancel pressed mid-run; we
                                          # still finished, but mark it so
                                          # the GUI handles it like Stop.
    return captured, result, used_partial


def _write_docx(paragraphs, out_path, *, show_timestamp=True, title=None,
                speakers=None):
    """Write paragraphs to a .docx file with a page-numbered footer and an
    italic disclaimer at the end.

    If `title` is given, it appears in bold at the top of the document
    before the transcript paragraphs.

    If `speakers` is given (a list parallel to paragraphs), each paragraph
    can be prefixed by a bold speaker label on its own line. To match
    legal-transcript convention, the label is only emitted when it differs
    from the previous paragraph's speaker (so a string of paragraphs from
    one speaker reads as a single labelled block).

    When `show_timestamp` is True (default), each paragraph is rendered as
    timestamp + tab + body, with a hanging indent so the timestamp sits in
    the left margin column and the body text wraps further right.

    When `show_timestamp` is False, paragraphs are rendered as flat prose
    with no indent."""
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    from docx.enum.section import WD_ORIENT

    doc = Document()

    # Force A4 portrait. python-docx defaults to US Letter (8.5x11"), so set
    # the page dimensions explicitly. A4 is 21.0 x 29.7 cm.
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # Slightly tighter margins to leave more room for the indented body.
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title at the top, if given. Bold, slightly larger than body, and
    # using the same monospaced font so the document has a consistent
    # typographic feel. A blank paragraph after the title gives the
    # transcript visual breathing room.
    if title:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.name = "Courier New"
        title_run.font.size = Pt(12)
        doc.add_paragraph()  # blank line between title and transcript

    # Body paragraph style: hanging indent (only useful when a timestamp
    # sits in the indent column).
    body_style = doc.styles.add_style("Transcript", WD_STYLE_TYPE.PARAGRAPH)
    pf = body_style.paragraph_format
    if show_timestamp:
        pf.left_indent = Cm(3.0)
        pf.first_line_indent = -Cm(3.0)
    pf.space_after = Pt(6)
    body_style.font.name = "Courier New"
    body_style.font.size = Pt(10)

    # Speaker label style: bold, no indent, slight breathing room before so
    # successive turns are visually separated.
    label_style = doc.styles.add_style("SpeakerLabel", WD_STYLE_TYPE.PARAGRAPH)
    label_style.font.name = "Courier New"
    label_style.font.size = Pt(10)
    label_style.font.bold = True
    label_style.paragraph_format.space_before = Pt(8)
    label_style.paragraph_format.space_after = Pt(0)
    # Keep the speaker label on the same page as the paragraph it
    # introduces. Without this Word can leave an orphan label at the
    # bottom of a page with its body paragraph starting overleaf.
    label_style.paragraph_format.keep_with_next = True

    last_speaker = None
    for i, para in enumerate(paragraphs):
        start = para[0][0]
        body = " ".join(seg[2] for seg in para).lstrip()
        speaker = speakers[i] if speakers and i < len(speakers) else None

        # Emit a speaker label whenever it changes. For named speakers we
        # emit their name; for an unattributed paragraph that follows a
        # named one, emit an explicit "[Unattributed]" marker so re-parsing
        # the file doesn't carry the previous speaker over.
        if speaker and speaker != last_speaker:
            doc.add_paragraph(speaker, style="SpeakerLabel")
        elif speaker is None and last_speaker is not None:
            doc.add_paragraph(UNATTRIBUTED_LABEL, style="SpeakerLabel")
        last_speaker = speaker

        p = doc.add_paragraph(style="Transcript")
        if show_timestamp:
            p.add_run(format_timestamp(start) + "\t" + body)
        else:
            p.add_run(body)

    # Italic disclaimer at the end. Plain (non-indented) paragraph.
    doc.add_paragraph()  # blank spacer line
    disc = doc.add_paragraph()
    disc_run = disc.add_run(
        "Transcribed using Transcribr - (c) James Leaver, 2026. "
        "If this text has not been deleted by the person who prepared this "
        "document, then the accuracy of this transcript may not have been "
        "checked by a human."
    )
    disc_run.italic = True

    # Footer: "Page X of Y" right-aligned, in Courier New so it matches
    # the body font. python-docx doesn't expose Word field codes
    # directly so we drop in the raw XML.
    def _add_field(paragraph, instr):
        run = paragraph.add_run()
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
        it.text = instr
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        run._r.append(f1); run._r.append(it); run._r.append(f2)

    def _styled_run(paragraph, text):
        run = paragraph.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        return run

    fp = doc.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _styled_run(fp, "Page ")
    _add_field(fp, " PAGE ")
    _styled_run(fp, " of ")
    _add_field(fp, " NUMPAGES ")

    doc.save(str(out_path))


def _write_pdf(paragraphs, out_path, *, show_timestamp=True, title=None,
               speakers=None):
    """Write paragraphs to an A4 PDF, visually matching the .docx output:
    Courier body with a hanging timestamp indent, bold speaker labels that
    stay with their paragraph, an italic disclaimer, and a right-aligned
    'Page X of Y' footer."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas as _pdfcanvas
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from xml.sax.saxutils import escape

    page_w, _page_h = A4

    class _NumberedCanvas(_pdfcanvas.Canvas):
        """Two-pass canvas so the footer can say 'Page X of Y'."""

        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_states = []

        def showPage(self):
            self._saved_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._saved_states)
            for state in self._saved_states:
                self.__dict__.update(state)
                self.setFont("Courier", 10)
                self.drawRightString(
                    page_w - 2.5 * cm, 1.5 * cm,
                    f"Page {self._pageNumber} of {total}")
                _pdfcanvas.Canvas.showPage(self)
            _pdfcanvas.Canvas.save(self)

    body_style = ParagraphStyle(
        "Transcript",
        fontName="Courier", fontSize=10, leading=12,
        spaceAfter=6,
    )
    if show_timestamp:
        # Hanging indent: timestamp in the left column, wrapped body
        # lines aligned under the body column (mirrors the docx style).
        body_style.leftIndent = 3.0 * cm
        body_style.firstLineIndent = -3.0 * cm
    label_style = ParagraphStyle(
        "SpeakerLabel",
        fontName="Courier-Bold", fontSize=10, leading=12,
        spaceBefore=8, spaceAfter=0, keepWithNext=1,
    )
    title_style = ParagraphStyle(
        "TranscriptTitle",
        fontName="Courier-Bold", fontSize=12, leading=14,
        spaceAfter=12,
    )
    disclaimer_style = ParagraphStyle(
        "Disclaimer",
        fontName="Courier-Oblique", fontSize=10, leading=12,
        spaceBefore=12,
    )

    story = []
    if title:
        story.append(Paragraph(escape(title), title_style))

    last_speaker = None
    for i, para in enumerate(paragraphs):
        start = para[0][0]
        body = " ".join(seg[2] for seg in para).lstrip()
        speaker = speakers[i] if speakers and i < len(speakers) else None
        if speaker and speaker != last_speaker:
            story.append(Paragraph(escape(speaker), label_style))
        elif speaker is None and last_speaker is not None:
            story.append(Paragraph(escape(UNATTRIBUTED_LABEL), label_style))
        last_speaker = speaker

        if show_timestamp:
            # Non-breaking spaces keep the timestamp and the gap to the
            # body out of reportlab's line-wrapping calculations.
            text = (escape(format_timestamp(start))
                    + "   " + escape(body))
        else:
            text = escape(body)
        story.append(Paragraph(text, body_style))

    story.append(Spacer(1, 6))
    story.append(Paragraph(
        escape(
            "Transcribed using Transcribr - (c) James Leaver, 2026. "
            "If this text has not been deleted by the person who prepared "
            "this document, then the accuracy of this transcript may not "
            "have been checked by a human."
        ),
        disclaimer_style,
    ))

    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=title or Path(out_path).stem,
    )
    doc.build(story, canvasmaker=_NumberedCanvas)


def _write_extra_formats(result, txt_out_path, formats, q):
    """Write SRT/VTT/JSON/TSV next to the paragraph .txt.

    Engine-agnostic: reads from result["segments"] which every engine
    runner normalises to a list of {"start", "end", "text"} dicts."""
    segments = (result or {}).get("segments", []) if result else []
    if not segments:
        q.put(("log", "Note: no segments available for extra outputs.\n"))
        return

    output_dir = txt_out_path.parent
    stem = txt_out_path.stem
    for suffix in (".transcript", ".paragraphs"):
        if stem.endswith(suffix):
            stem = stem[: -len(suffix)]
            break

    for fmt in formats:
        out = output_dir / f"{stem}.{fmt}"
        try:
            if fmt == "json":
                _write_json_result(out, result)
            elif fmt == "srt":
                _write_srt(out, segments)
            elif fmt == "vtt":
                _write_vtt(out, segments)
            elif fmt == "tsv":
                _write_tsv(out, segments)
            else:
                q.put(("log",
                       f"WARNING: unknown extra format '{fmt}', skipped.\n"))
                continue
            q.put(("log", f"Wrote: {out.name}\n"))
        except Exception as e:
            q.put(("log",
                   f"WARNING: failed to write .{fmt}: "
                   f"{type(e).__name__}: {e}\n"))


def _seg_field(seg, key):
    """Pull start/end/text from either a dict or an object segment."""
    if isinstance(seg, dict):
        return seg.get(key)
    return getattr(seg, key, None)


def _format_srt_time(seconds):
    seconds = max(0.0, float(seconds))
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds - h * 3600 - m * 60
    return f"{h:02d}:{m:02d}:{s:06.3f}".replace(".", ",")


def _format_vtt_time(seconds):
    seconds = max(0.0, float(seconds))
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds - h * 3600 - m * 60
    return f"{h:02d}:{m:02d}:{s:06.3f}"


def _write_srt(path, segments):
    with open(path, "w", encoding="utf-8") as fh:
        for i, seg in enumerate(segments, start=1):
            start = _seg_field(seg, "start") or 0
            end = _seg_field(seg, "end") or 0
            text = (_seg_field(seg, "text") or "").strip()
            fh.write(f"{i}\n")
            fh.write(f"{_format_srt_time(start)} --> "
                     f"{_format_srt_time(end)}\n")
            fh.write(f"{text}\n\n")


def _write_vtt(path, segments):
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("WEBVTT\n\n")
        for seg in segments:
            start = _seg_field(seg, "start") or 0
            end = _seg_field(seg, "end") or 0
            text = (_seg_field(seg, "text") or "").strip()
            fh.write(f"{_format_vtt_time(start)} --> "
                     f"{_format_vtt_time(end)}\n")
            fh.write(f"{text}\n\n")


def _write_tsv(path, segments):
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("start\tend\ttext\n")
        for seg in segments:
            start = _seg_field(seg, "start") or 0
            end = _seg_field(seg, "end") or 0
            text = (_seg_field(seg, "text") or "").strip().replace("\t", " ")
            fh.write(f"{int(float(start) * 1000)}\t"
                     f"{int(float(end) * 1000)}\t{text}\n")


def _write_json_result(path, result):
    """Dump the full result dict as JSON. Coerces any non-serialisable
    attributes (e.g. faster-whisper named tuples) to plain dicts."""
    import json

    def _coerce(obj):
        if isinstance(obj, (str, int, float, bool, type(None))):
            return obj
        if isinstance(obj, dict):
            return {k: _coerce(v) for k, v in obj.items()}
        if isinstance(obj, (list, tuple)):
            return [_coerce(v) for v in obj]
        if hasattr(obj, "_asdict"):
            return _coerce(obj._asdict())
        if hasattr(obj, "__dict__"):
            return _coerce(vars(obj))
        return str(obj)

    with open(path, "w", encoding="utf-8") as fh:
        json.dump(_coerce(result), fh, indent=2)


# =====================================================================
# Web UI backend - transcript model
# =====================================================================
#
# The review pane's document logic with no widgets: paragraphs (lists
# of (start, end, text) segment tuples), speaker slots "1".."9",
# snapshot-based undo/redo, split/merge/edit/replace-all, playback
# spans, confidence spans, and the debounced autosave trigger. The web
# review surface drives one of these over the API; ReviewPaneText keeps
# its own fused copy of the same logic until the Tk UI is retired.
#
# Semantics are ported 1:1 from ReviewPaneText - where behaviour looks
# arbitrary (the +0.3s playback tail, bailing out of confidence shading
# on alignment drift, empty edit meaning cancel) it is deliberate and
# matches what the Tk pane does today.

class TranscriptModel:
    MAX_SPEAKERS = 9
    DEFAULT_VISIBLE = 4
    LETTERS = tuple(str(i) for i in range(1, 10))
    _UNDO_LIMIT = 200

    def __init__(self, paragraphs, speakers=None, speaker_names=None,
                 word_conf=None, *, on_autosave=None, timer_factory=None):
        self.paragraphs = [list(p) for p in paragraphs]
        n = len(self.paragraphs)
        self.speakers = (list(speakers) if speakers is not None
                         else [None] * n)
        self.speaker_names = {letter: f"Speaker {letter}"
                              for letter in self.LETTERS}
        if speaker_names:
            self.speaker_names.update(speaker_names)
        self.word_conf = word_conf
        self.show_confidence = False
        used = [int(s) for s in self.speakers if s]
        self.visible_speakers = max(self.DEFAULT_VISIBLE,
                                    min(max(used, default=0),
                                        self.MAX_SPEAKERS))
        self._next_id = 0
        self.ids = [self._new_id() for _ in self.paragraphs]
        self.rev = 0
        self._undo_stack = []
        self._redo_stack = []
        self.on_autosave = on_autosave
        self._timer_factory = timer_factory or (
            lambda delay, fn: threading.Timer(delay, fn))
        self._autosave_timer = None

    def _new_id(self):
        self._next_id += 1
        return self._next_id

    # ----- Queries ----------------------------------------------------

    def body(self, idx):
        """The paragraph's rendered text - exactly what the client
        displays and what split()'s character offset indexes into."""
        return " ".join(seg[2] for seg in self.paragraphs[idx]).strip()

    def can_undo(self):
        return bool(self._undo_stack)

    def can_redo(self):
        return bool(self._redo_stack)

    def label_counts(self):
        labelled = sum(1 for s in self.speakers if s is not None)
        return labelled, len(self.speakers)

    def resolved_speakers(self):
        """Slot letters -> stripped display names (None stays None)."""
        return [self.speaker_names.get(letter, "").strip() or None
                if letter else None
                for letter in self.speakers]

    def playback_span(self, idx):
        """(start_seconds, duration_or_None) covering paragraph `idx`,
        or None if there's nothing to play. Loaded transcripts only
        know start times (the parsers synthesise ~1s spans), so a span
        that still looks like a placeholder plays through to the next
        paragraph's start instead - or open-ended for the last one."""
        if not (0 <= idx < len(self.paragraphs)):
            return None
        para = self.paragraphs[idx]
        if not para:
            return None
        start = max(0.0, float(para[0][0]))
        end = float(para[-1][1])
        if end - start <= 1.0:
            if idx + 1 < len(self.paragraphs) and self.paragraphs[idx + 1]:
                next_start = float(self.paragraphs[idx + 1][0][0])
                if next_start > start:
                    return (start, max(0.5, next_start - start + 0.3))
            return (start, None)
        # A touch of tail padding so the last word isn't clipped.
        return (start, max(0.5, end - start + 0.3))

    def _bucket_words_by_paragraph(self):
        """Partition word_conf into one [(word, prob)] list per
        paragraph by end time; single forward pass (both sides are
        time-ordered)."""
        buckets = [[] for _ in self.paragraphs]
        words = self.word_conf or []
        wi = 0
        n_para = len(self.paragraphs)
        for i, para in enumerate(self.paragraphs):
            if not para:
                continue
            p_end = para[-1][1]
            is_last = (i == n_para - 1)
            while wi < len(words):
                if not is_last and words[wi][0] >= p_end:
                    break
                buckets[i].append((words[wi][2], words[wi][3]))
                wi += 1
        return buckets

    def confidence_spans(self):
        """Per paragraph: [(start_char, end_char, "low"|"med"), ...] in
        body() coordinates. A paragraph whose text no longer aligns
        with the captured words gets an empty list - never a wrong
        highlight."""
        if not self.word_conf:
            return [[] for _ in self.paragraphs]
        result = []
        for i, words_i in enumerate(self._bucket_words_by_paragraph()):
            spans = []
            body = self.body(i)
            cursor = 0
            for wtext, prob in words_i:
                token = (wtext or "").strip()
                if not token:
                    continue
                pos = body.find(token, cursor)
                if pos < 0:
                    spans = []      # alignment lost - leave unshaded
                    break
                cursor = pos + len(token)
                if prob is None:
                    continue
                if prob < 0.35:
                    spans.append((pos, cursor, "low"))
                elif prob < 0.6:
                    spans.append((pos, cursor, "med"))
            result.append(spans)
        return result

    def attention_flags(self):
        """True per paragraph when it still needs attention: no speaker,
        or (with confidence shading on) a word below 0.6."""
        buckets = (self._bucket_words_by_paragraph()
                   if (self.show_confidence and self.word_conf) else None)
        flags = []
        for i, letter in enumerate(self.speakers):
            if letter is None:
                flags.append(True)
            elif buckets is not None:
                flags.append(any(p is not None and p < 0.6
                                 for _, p in buckets[i]))
            else:
                flags.append(False)
        return flags

    def _time_at_body_offset(self, para, joined, offset):
        """Start time of the word at/after character `offset` in
        `joined`, from the engine's word timestamps; None when word
        data is missing or no longer aligns with the edited text."""
        if not self.word_conf or not para:
            return None
        p_start = para[0][0]
        p_end = para[-1][1]
        cursor = 0
        for w_start, _w_end, w_text, _prob in self.word_conf:
            if w_start < p_start - 0.001:
                continue
            if w_start > p_end + 0.001:
                break
            token = (w_text or "").strip()
            if not token:
                continue
            pos = joined.find(token, cursor)
            if pos < 0:
                return None
            if pos >= offset:
                return float(w_start)
            cursor = pos + len(token)
        return None

    # ----- Undo / redo / autosave --------------------------------------

    def _snapshot(self):
        return {
            "paragraphs": [list(p) for p in self.paragraphs],
            "speakers": list(self.speakers),
            "speaker_names": dict(self.speaker_names),
            "visible_speakers": self.visible_speakers,
            "ids": list(self.ids),
        }

    def _restore(self, snap):
        self.paragraphs = [list(p) for p in snap["paragraphs"]]
        self.speakers = list(snap["speakers"])
        self.speaker_names = dict(snap["speaker_names"])
        self.visible_speakers = snap["visible_speakers"]
        self.ids = list(snap["ids"])

    def _push_undo(self):
        self._undo_stack.append(self._snapshot())
        if len(self._undo_stack) > self._UNDO_LIMIT:
            self._undo_stack.pop(0)
        self._redo_stack.clear()
        # Every mutation passes through here, so it doubles as the
        # autosave trigger (as in the Tk pane).
        self._schedule_autosave()

    def undo(self):
        if not self._undo_stack:
            return False
        self._redo_stack.append(self._snapshot())
        self._restore(self._undo_stack.pop())
        self.rev += 1
        self._schedule_autosave()
        return True

    def redo(self):
        if not self._redo_stack:
            return False
        self._undo_stack.append(self._snapshot())
        self._restore(self._redo_stack.pop())
        self.rev += 1
        self._schedule_autosave()
        return True

    def _schedule_autosave(self):
        """Debounced crash-recovery snapshot: fires 3s after the last
        mutation rather than on every keystroke of a rename."""
        if self.on_autosave is None:
            return
        if self._autosave_timer is not None:
            self._autosave_timer.cancel()
        t = self._timer_factory(3.0, self._fire_autosave)
        if hasattr(t, "daemon"):
            t.daemon = True
        t.start()
        self._autosave_timer = t

    def _fire_autosave(self):
        self._autosave_timer = None
        if self.on_autosave is not None:
            self.on_autosave(self.paragraphs, self.speakers,
                             dict(self.speaker_names))

    def flush_autosave(self):
        """Run any pending autosave immediately (shutdown, tests)."""
        if self._autosave_timer is not None:
            self._autosave_timer.cancel()
            self._fire_autosave()

    def close(self):
        if self._autosave_timer is not None:
            self._autosave_timer.cancel()
            self._autosave_timer = None

    # ----- Mutations ----------------------------------------------------

    def set_speaker(self, idx, letter):
        """Assign slot `letter` ("1".."9") or None to paragraph idx."""
        if not (0 <= idx < len(self.paragraphs)):
            return False
        if letter is not None and letter not in self.LETTERS:
            return False
        self._push_undo()
        self.speakers[idx] = letter
        if letter is not None and int(letter) > self.visible_speakers:
            self.set_visible(int(letter))
        self.rev += 1
        return True

    def set_speaker_name(self, letter, name):
        """Rename a slot. No undo step (matches the Tk pane, where each
        keystroke of a rename would otherwise flood the stack), but it
        reschedules the autosave."""
        if letter not in self.LETTERS:
            return False
        self.speaker_names[letter] = name
        self.rev += 1
        self._schedule_autosave()
        return True

    def set_visible(self, n):
        """Show at least `n` speaker-name fields (clamped)."""
        n = max(self.DEFAULT_VISIBLE, min(int(n), self.MAX_SPEAKERS))
        if n != self.visible_speakers:
            self.visible_speakers = n
            self.rev += 1
        return self.visible_speakers

    def commit_edit(self, idx, new_text):
        """Replace paragraph idx's text, collapsing it to one segment
        that keeps the original outer time span. Empty text means
        cancel (no-op); unchanged text records no undo step."""
        if not (0 <= idx < len(self.paragraphs)):
            return False
        new_text = (new_text or "").strip()
        if not new_text:
            return False
        para = self.paragraphs[idx]
        if not para:
            return False
        if new_text != self.body(idx):
            self._push_undo()
        start_t = para[0][0]
        end_t = para[-1][1]
        self.paragraphs[idx] = [(start_t, end_t, new_text)]
        self.rev += 1
        return True

    def merge_with_previous(self, idx):
        if idx is None or idx <= 0 or idx >= len(self.paragraphs):
            return False
        self._push_undo()
        self.paragraphs[idx - 1] = (list(self.paragraphs[idx - 1])
                                    + list(self.paragraphs[idx]))
        del self.paragraphs[idx]
        del self.speakers[idx]
        del self.ids[idx]
        self.rev += 1
        return True

    def split(self, idx, offset):
        """Split paragraphs[idx] at character `offset` within body().
        Both halves get real start/end times: preferring the engine's
        word timestamps, falling back to interpolating by character
        position. The new second half keeps the speaker and gets a
        fresh id. Returns the new paragraph's index, or None."""
        if not (0 <= idx < len(self.paragraphs)):
            return None
        para = self.paragraphs[idx]
        if not para:
            return None
        joined = " ".join(seg[2] for seg in para)
        leading_strip = len(joined) - len(joined.lstrip())
        adjusted_offset = offset + leading_strip
        if adjusted_offset <= 0 or adjusted_offset >= len(joined):
            return None

        seg_starts = []
        running = ""
        for k, (_, _, t) in enumerate(para):
            if k > 0:
                running += " "
            seg_starts.append(len(running))
            running += t

        split_seg = None
        split_within = 0
        for k in range(len(para)):
            seg_start = seg_starts[k]
            seg_end = seg_start + len(para[k][2])
            if adjusted_offset <= seg_start:
                split_seg = k
                split_within = 0
                break
            if adjusted_offset <= seg_end:
                split_seg = k
                split_within = adjusted_offset - seg_start
                break
        if split_seg is None:
            return None

        if split_within == 0:
            first = list(para[:split_seg])
            second = list(para[split_seg:])
        else:
            seg_start_t, seg_end_t, seg_text = para[split_seg]
            text_before = seg_text[:split_within].rstrip()
            text_after = seg_text[split_within:].lstrip()
            split_t = self._time_at_body_offset(para, joined,
                                                adjusted_offset)
            if split_t is None and seg_text:
                split_t = (seg_start_t
                           + (seg_end_t - seg_start_t)
                           * (split_within / len(seg_text)))
            if split_t is None:
                split_t = seg_start_t
            split_t = max(seg_start_t, min(float(split_t), seg_end_t))
            first = list(para[:split_seg])
            if text_before:
                first.append((seg_start_t, split_t, text_before))
            second = []
            if text_after:
                second.append((split_t, seg_end_t, text_after))
            second.extend(para[split_seg + 1:])

        if not first or not second:
            return None

        self._push_undo()
        self.paragraphs[idx] = first
        self.paragraphs.insert(idx + 1, second)
        self.speakers.insert(idx + 1, self.speakers[idx])
        self.ids.insert(idx + 1, self._new_id())
        self.rev += 1
        return idx + 1

    def replace_all(self, term, replacement, match_case=False):
        """Replace every occurrence across the document. One undo step;
        returns the number of replacements (0 = nothing recorded)."""
        if not term:
            return 0
        flags = 0 if match_case else re.IGNORECASE
        pattern = re.compile(re.escape(term), flags)
        total = sum(pattern.subn(lambda _m: replacement, seg[2])[1]
                    for para in self.paragraphs for seg in para)
        if total == 0:
            return 0
        self._push_undo()
        for para in self.paragraphs:
            for k, seg in enumerate(para):
                new_text, n = pattern.subn(lambda _m: replacement, seg[2])
                if n:
                    para[k] = (seg[0], seg[1], new_text)
        self.rev += 1
        return total


def build_worker_params(settings, in_path, out_path, *,
                        review_before_save):
    """Assemble the transcribe_worker params dict from a validated
    settings dict for one (in_path -> out_path) job - the module-level
    twin of WhisperGUI._build_params. Raises _EngineNotAvailable when
    no Whisper engine is installed."""
    if not AVAILABLE_ENGINES:
        raise _EngineNotAvailable(
            "No Whisper engine is installed in this Python.\n\n"
            "Install at least one:\n"
            "  pip install openai-whisper\n"
            "  pip install faster-whisper\n"
            "  pip install mlx-whisper   (Apple Silicon only)\n")

    lang_code = next(
        (c for n, c in LANGUAGES if n == settings["language"]), "en")

    extra_formats = [fmt for fmt in ("json", "srt", "vtt", "tsv")
                     if settings.get(f"extra_{fmt}")]

    # Title and prompt are separate: the document title is never fed to
    # the engine, and the optional prompt (a vocabulary hint) is fed as
    # initial_prompt. Prompting is opt-in because it can backfire - it can
    # bleed into the transcript or trigger hallucinations on unclear audio.
    prompt = (settings.get("prompt") or "").strip()
    doc_title = (settings.get("title") or "").strip()

    engine_key = resolve_engine_key(settings["engine"])

    return dict(
        input=in_path,
        output=out_path,
        engine=engine_key,
        model=settings["model"],
        language=lang_code,
        task=settings["task"],
        temperature=settings["temperature"],
        beam_size=settings["beam_size"],
        best_of=settings["best_of"],
        compression_ratio_threshold=settings[
            "compression_ratio_threshold"],
        logprob_threshold=settings["logprob_threshold"],
        no_speech_threshold=settings["no_speech_threshold"],
        condition_on_previous_text=settings["condition_on_previous_text"],
        # Always on since 0.9.0: word timings let paragraph gaps
        # measure real silence instead of Whisper's padded segment
        # edges, sharpen speaker labelling and playback spans, and
        # feed confidence shading. The modest speed cost buys a
        # noticeably better transcript.
        word_timestamps=True,
        highlight_confidence=settings["highlight_confidence"],
        initial_prompt=prompt or None,
        # With no title, name the document after the source file. (The
        # filename is never fed to the engine - recorder names like
        # REC_0042 would only mislead it.)
        title=doc_title or Path(in_path).name,
        gap=settings["gap"],
        extra_formats=extra_formats,
        output_format=settings["output_format"],
        show_timestamp=settings["show_timestamp"],
        audio_duration=get_audio_duration(in_path),
        review_before_save=review_before_save,
        diarize=bool(settings.get("diarize")),
        num_speakers=int(settings.get("num_speakers") or 0),
        diarize_model=(settings.get("diarize_model")
                       or DIARIZE_DEFAULT_VOICE_MODEL),
        diarize_threshold=float(settings.get("diarize_threshold")
                                or _DIARIZE_CLUSTER_THRESHOLD),
    )


# =====================================================================
# Web UI backend - settings contract
# =====================================================================
#
# The web front-end shares settings.json with the Tk UI: same keys,
# same value conventions (engine and language are stored as display
# names, exactly as the Tk comboboxes hold them). validate_settings()
# mirrors _apply_settings(): unknown keys and junk values are dropped
# and missing ones fall back to the defaults, so a stale or corrupt
# file can't wedge the app.

# ("word_timestamps" was a stored setting until 0.9.0; it is now always
# on for the run, and any stale key in settings.json is simply dropped.)
_SETTINGS_BOOL_KEYS = (
    "show_timestamp", "review", "condition_on_previous_text",
    "extra_json", "extra_srt", "extra_vtt",
    "extra_tsv", "highlight_confidence", "show_details",
    "diarize", "show_all_models",
)

_SETTINGS_NUMBER_KEYS = (
    "gap", "temperature", "beam_size", "best_of",
    "compression_ratio_threshold", "logprob_threshold",
    "no_speech_threshold", "num_speakers", "diarize_threshold",
)


def _settings_choices():
    """Allowed values for every enumerated setting, computed live so
    the engine list reflects what is actually installed."""
    return {
        "engine": [ENGINE_AUTO_NAME] + [name for _, name in
                                        AVAILABLE_ENGINES],
        "model": list(WHISPER_MODELS),
        "diarize_model": [m["id"] for m in DIARIZE_VOICE_MODELS],
        "language": [n for n, _ in LANGUAGES],
        "task": ["transcribe", "translate"],
        "output_format": ["txt", "docx", "pdf"],
        "theme": ["auto", "light", "dark"],
    }


def default_settings():
    """The out-of-the-box settings - the same values the Tk widgets
    start with."""
    return {
        "engine": ENGINE_AUTO_NAME,
        "model": "large-v3-turbo",
        "language": "English",
        "task": "transcribe",
        "output_format": "docx",
        "title": "",
        "prompt": "",
        "gap": 1.5,
        "show_timestamp": True,
        "review": True,
        "temperature": 0.0,
        "beam_size": 5,
        "best_of": 5,
        "compression_ratio_threshold": 2.4,
        "logprob_threshold": -1.0,
        "no_speech_threshold": 0.6,
        "condition_on_previous_text": True,
        "extra_json": False,
        "extra_srt": False,
        "extra_vtt": False,
        "extra_tsv": False,
        "highlight_confidence": False,
        "theme": "auto",
        "show_details": False,
        "diarize": False,
        "num_speakers": 0,
        "diarize_model": DIARIZE_DEFAULT_VOICE_MODEL,
        "diarize_threshold": _DIARIZE_CLUSTER_THRESHOLD,
        "show_all_models": False,
    }


def validate_settings(raw, base=None):
    """Merge `raw` (a loaded settings.json, or a PUT body) over `base`
    (the defaults when None), accepting only known keys with plausible
    values. Passing the currently stored settings as `base` makes a
    partial update non-destructive."""
    import math

    merged = default_settings()
    if base:
        merged.update(base)
    if not isinstance(raw, dict):
        return merged
    raw = dict(raw)
    # One-time migration: the old single "Description of file" field was
    # both the Whisper prompt and the document title. They are separate
    # now. A pre-split settings.json has no "title" key - move its text
    # into the title (the benign use) and clear the prompt, so upgrading
    # doesn't keep silently priming the engine. Only at load time
    # (base is None), never on a partial PUT from the current UI.
    if base is None and "title" not in raw and raw.get("prompt"):
        raw["title"] = str(raw["prompt"])
        raw["prompt"] = ""
    for key, allowed in _settings_choices().items():
        v = raw.get(key)
        if key == "model" and isinstance(v, str):
            v = _canonical_model(v)   # legacy "large"/"turbo" -> canonical
        if isinstance(v, str) and v in allowed:
            merged[key] = v
    for key in ("title", "prompt"):
        v = raw.get(key)
        if isinstance(v, str):
            merged[key] = v
    for key in _SETTINGS_NUMBER_KEYS:
        v = raw.get(key)
        if (isinstance(v, (int, float)) and not isinstance(v, bool)
                and math.isfinite(v)):
            merged[key] = v
    for key in _SETTINGS_BOOL_KEYS:
        v = raw.get(key)
        if isinstance(v, bool):
            merged[key] = v
    return merged


def current_settings():
    """Defaults overlaid with whatever settings.json currently holds."""
    return validate_settings(_settings_load())


# =====================================================================
# Web UI backend - event broker (SSE fan-out)
# =====================================================================

class EventBroker:
    """Fans server-side events out to any number of SSE subscribers.

    Events carry monotonically increasing ids and are kept in a bounded
    ring so a reconnecting client (Last-Event-ID) can replay what it
    missed; when the gap is wider than the ring the subscriber is told
    to resync via GET /api/state instead."""

    _RING_SIZE = 500

    def __init__(self):
        from collections import deque
        self._lock = threading.Lock()
        self._seq = 0
        self._subscribers = set()
        self._ring = deque(maxlen=self._RING_SIZE)

    def publish(self, event, data):
        import json
        with self._lock:
            self._seq += 1
            item = (self._seq, event, json.dumps(data))
            self._ring.append(item)
            targets = list(self._subscribers)
        for q in targets:
            q.put(item)

    def subscribe(self, last_event_id=None):
        """Register a subscriber. Returns (queue, backlog): backlog is
        the list of missed events, or None when the gap can't be
        bridged and the client must do a full resync."""
        q = queue.Queue()
        with self._lock:
            self._subscribers.add(q)
            if last_event_id is None:
                return q, []
            ring = list(self._ring)
        if not ring or last_event_id < ring[0][0] - 1:
            return q, None
        return q, [item for item in ring if item[0] > last_event_id]

    def unsubscribe(self, q):
        with self._lock:
            self._subscribers.discard(q)


def _sse_format(item):
    seq, event, payload = item
    return f"id: {seq}\nevent: {event}\ndata: {payload}\n\n"


# =====================================================================
# Web UI backend - run controller
# =====================================================================

class ApiFail(Exception):
    """A request failure with an HTTP status and machine-readable code;
    the route layer turns it into a JSON error response."""

    def __init__(self, status, code, message, **extra):
        super().__init__(message)
        self.status = status
        self.code = code
        self.extra = extra


_NO_ENGINE_HINT = (
    "No Whisper engine is installed in this Python.\n\n"
    "Install at least one:\n"
    "  pip install openai-whisper\n"
    "  pip install faster-whisper\n"
    "  pip install mlx-whisper   (Apple Silicon only)")


# Extensions that mark a file as a recording. A transcript must never
# be written to a path that looks like one.
_MEDIA_EXTENSIONS = {
    ".mp3", ".wav", ".m4a", ".aac", ".flac", ".ogg", ".opus",
    ".mp4", ".mov", ".mkv", ".avi", ".webm", ".wma", ".amr", ".3gp",
}


def ensure_output_is_safe(in_path, out_path):
    """Refuse output paths that could destroy a recording: the input
    itself (however it is spelled or linked), or any path wearing an
    audio/video extension. Raises ApiFail; there is deliberately no
    force/overwrite override for these.

    Born of a real incident: a save-dialog misclick set Output to the
    source .mp3, the user confirmed the generic "already exists -
    overwrite?" prompt, and the pre-review safety copy replaced a piece
    of case audio with a Word document (recovered from Dropbox version
    history)."""
    src, dst = Path(in_path), Path(out_path)
    try:
        same = (os.path.normcase(str(src.resolve()))
                == os.path.normcase(str(dst.resolve())))
        if not same and dst.exists():
            same = dst.samefile(src)
    except OSError:
        same = False
    if same:
        raise ApiFail(
            400, "output_is_input",
            "The output path is the input recording itself - saving "
            "would destroy the recording.\n\nChoose a different output "
            "file (leave the field empty to save next to the input as "
            f"'{Path(in_path).stem}.transcript.<format>').")
    if dst.suffix.lower() in _MEDIA_EXTENSIONS:
        raise ApiFail(
            400, "output_looks_like_media",
            f"The output path ends in '{dst.suffix}', which is an "
            "audio/video extension - transcripts save as .txt, .docx "
            "or .pdf. This guard exists so a transcript can never "
            "overwrite a recording.\n\nCheck the Output field: it may "
            "be pointing at a recording by mistake.")


class RunController:
    """Single-run and batch state machine - the module-level twin of
    WhisperGUI's run/batch/stop handlers. Mutating entry points raise
    ApiFail on validation problems; progress flows out as SSE events
    via the broker. `worker_fn` is injectable for tests."""

    _LOG_CAP = 4000          # lines kept for late-joining clients

    def __init__(self, broker, *, worker_fn=None):
        self.broker = broker
        self.queue = queue.Queue()
        self.cancel_event = threading.Event()
        self.worker = None
        self.worker_fn = worker_fn or transcribe_worker
        self.lock = threading.RLock()
        self._batch = None
        self.phase = "idle"    # idle|running|stopping|done|error|cancelled
        self.current_file = None
        self.run_id = 0
        self.last_output = None
        self.log_lines = []
        self.progress = None
        # Phase 3 sets this to route paragraphs_ready into a review
        # session; until then runs force review_before_save=False.
        self.on_paragraphs_ready = None
        # Set by WebBackend to the model manager's is_busy(), so a run
        # can't start while a model download holds the weights cache.
        self.model_busy_check = None

    # ----- Snapshots ----------------------------------------------------

    def state(self):
        with self.lock:
            batch = None
            if self._batch is not None:
                batch = {"index": self._batch["index"],
                         "total": len(self._batch["items"])}
            return {
                "phase": self.phase,
                "file": self.current_file,
                "run_id": self.run_id,
                "batch": batch,
                "out_path": self.last_output,
                "progress": self.progress,
                "log": "".join(self.log_lines)[-40000:],
            }

    def _publish_run_state(self, **extra):
        self.broker.publish("run_state", {**{
            "phase": self.phase,
            "file": self.current_file,
            "run_id": self.run_id,
            "batch": (None if self._batch is None else
                      {"index": self._batch["index"],
                       "total": len(self._batch["items"])}),
            "out_path": self.last_output,
        }, **extra})

    def _append_log(self, text):
        self.log_lines.append(text)
        if len(self.log_lines) > self._LOG_CAP:
            del self.log_lines[:len(self.log_lines) - self._LOG_CAP]
        self.broker.publish("log", {"text": text})

    def _reset_log(self):
        self.log_lines = []

    # ----- Starting work -------------------------------------------------

    def _ensure_idle(self):
        if self.phase in ("running", "stopping"):
            raise ApiFail(409, "busy", "A transcription is already "
                                       "running.")
        if self.model_busy_check is not None and self.model_busy_check():
            raise ApiFail(409, "model_busy",
                          "A model is downloading. Wait for it to finish "
                          "before starting a transcription.")

    def _spawn(self, params):
        self.cancel_event.clear()
        self.worker = threading.Thread(
            target=self.worker_fn,
            args=(params, self.queue, self.cancel_event),
            daemon=True)
        self.worker.start()

    def start_single(self, in_path, out_path, settings, *, force=False):
        """Validate and launch one transcription. Returns the new
        run_id. Mirrors WhisperGUI._on_run."""
        with self.lock:
            self._ensure_idle()
            in_path = (in_path or "").strip()
            if not in_path:
                raise ApiFail(400, "missing_input",
                              "Please choose an input audio/video file, "
                              "or add files to the batch queue.")
            if not Path(in_path).exists():
                raise ApiFail(400, "input_not_found",
                              f"Input file does not exist:\n{in_path}",
                              path=in_path)
            out_path = (out_path or "").strip()
            if not out_path:
                ext = settings["output_format"]
                out_path = str(Path(in_path).with_suffix(
                    f".transcript.{ext}"))
            ensure_output_is_safe(in_path, out_path)
            if Path(out_path).exists() and not force:
                raise ApiFail(409, "output_exists",
                              f"The output file already exists:\n\n"
                              f"{out_path}", path=out_path)
            try:
                params = build_worker_params(
                    settings, in_path, out_path,
                    review_before_save=bool(settings.get("review")))
            except _EngineNotAvailable:
                raise ApiFail(409, "no_engine", _NO_ENGINE_HINT)

            self._batch = None
            self.run_id += 1
            self.phase = "running"
            self.current_file = Path(in_path).name
            self.last_output = None
            self.progress = None
            self._reset_log()
            self._publish_run_state()
            self._append_log(f"=== {Path(in_path).name} ===\n")
            self._spawn(params)
            return self.run_id

    def start_batch(self, files, settings, *, force=False):
        """Validate and launch a sequential unattended batch (never
        reviews). Mirrors WhisperGUI._start_batch."""
        with self.lock:
            self._ensure_idle()
            files = [f for f in (files or []) if (f or "").strip()]
            if not files:
                raise ApiFail(400, "missing_input",
                              "The batch queue is empty.")
            ext = settings["output_format"]
            items = []
            missing = []
            for in_path in files:
                if not Path(in_path).exists():
                    missing.append(in_path)
                    continue
                batch_out = str(Path(in_path).with_suffix(
                    f".transcript.{ext}"))
                ensure_output_is_safe(in_path, batch_out)
                items.append((in_path, batch_out))
            if missing:
                raise ApiFail(400, "missing_inputs",
                              "These queued files no longer exist.",
                              missing=missing)
            existing = [o for _, o in items if Path(o).exists()]
            if existing and not force:
                raise ApiFail(409, "outputs_exist",
                              f"{len(existing)} output file(s) already "
                              "exist and will be overwritten.",
                              existing=existing[:8], total=len(existing))
            try:
                # Probe once so an engine problem surfaces before any
                # file is committed to the run.
                build_worker_params(settings, items[0][0], items[0][1],
                                    review_before_save=False)
            except _EngineNotAvailable:
                raise ApiFail(409, "no_engine", _NO_ENGINE_HINT)

            self._batch = {"items": items, "index": 0, "settings":
                           dict(settings), "succeeded": [], "failed": [],
                           "stop": False}
            self.run_id += 1
            self.last_output = None
            self.progress = None
            self._reset_log()
            self._append_log(
                f"=== Batch: {len(items)} file(s) queued ===\n")
            self._start_batch_item()
            return self.run_id

    def _start_batch_item(self):
        b = self._batch
        if b is None:
            return
        idx = b["index"]
        in_path, out_path = b["items"][idx]
        n = len(b["items"])
        self._append_log(
            f"\n--- File {idx + 1} of {n}: {Path(in_path).name} ---\n")
        self.phase = "running"
        self.current_file = f"{Path(in_path).name}  (file {idx + 1} of {n})"
        self.progress = None
        try:
            params = build_worker_params(b["settings"], in_path, out_path,
                                         review_before_save=False)
        except _EngineNotAvailable:
            # Engine vanished mid-run; abort the batch.
            self._finish_batch(stopped=True)
            return
        self._publish_run_state()
        self._spawn(params)

    def stop(self):
        """Stop after the current segment (single) / current file
        (batch), saving the partial transcript. Mirrors _on_stop."""
        with self.lock:
            if self._batch is not None:
                self._batch["stop"] = True
            if self.worker and self.worker.is_alive():
                self.cancel_event.set()
                self.phase = "stopping"
                self._append_log(
                    "\n[Stop requested - finishing current segment "
                    "and saving partial transcript...]\n")
                self._publish_run_state()
                return True
            return False

    # ----- Worker-queue dispatch (called from the pump thread) -----------

    def handle_message(self, kind, data):
        with self.lock:
            if kind == "log":
                self._append_log(data)
            elif kind == "eta":
                self._handle_eta(data)
            elif kind == "download":
                self._handle_download(data)
            elif kind == "status":
                self._handle_status(data)
            elif kind == "paragraphs_ready":
                if self.on_paragraphs_ready is not None:
                    self.on_paragraphs_ready(data)
                else:
                    self._append_log("\n[Internal: review payload arrived "
                                     "with no review surface]\n")
            elif kind == "done":
                if self._batch is not None:
                    self._batch_item_done(data, error=None)
                else:
                    self._on_done(data)
            elif kind == "error":
                if self._batch is not None:
                    self._batch_item_done(None, error=data)
                else:
                    self._on_error(data)
            elif kind == "cancelled":
                if self._batch is not None:
                    self._batch_cancelled(data)
                else:
                    self._on_cancelled(data)

    def _handle_eta(self, info):
        done = info["audio_done"]
        total = info["audio_total"]
        pct = (done / total * 100) if total else 0
        self.progress = {
            "pct": max(0.0, min(100.0, float(pct))),
            "stage": "transcribing",
            "status_text": (
                f"{_format_duration(done)} of {_format_duration(total)}"
                f"   ·   about {_format_duration(info['eta_seconds'])} "
                f"remaining   ·   {info['speed']:.1f}x audio speed"),
            **info,
        }
        self.broker.publish("progress", self.progress)

    def _handle_download(self, info):
        """Turn a worker 'download' message (first-run model fetch) into a
        moving progress bar with a clear 'Downloading model...' status, so
        a multi-gigabyte download never looks like a frozen app."""
        downloaded = info.get("downloaded") or 0
        total = info.get("total") or 0
        speed = info.get("speed") or 0
        pct = (downloaded / total * 100) if total else 0
        parts = [f"Downloading model '{info.get('model', '')}'",
                 f"{_humanize_bytes(downloaded)} of {_humanize_bytes(total)}"]
        if speed > 0:
            parts.append(f"{_humanize_bytes(speed)}/s")
        parts.append("first use only")
        self.progress = {
            "pct": max(0.0, min(100.0, float(pct))),
            "stage": "downloading",
            "status_text": "   ·   ".join(parts),
        }
        self.broker.publish("progress", self.progress)

    def _handle_status(self, info):
        """A stage marker with no measurable percentage (loading a model
        into memory, starting transcription). Shows the user which step is
        running - as an indeterminate bar - so the flat stretches before
        the first segment don't read as a hang."""
        self.progress = {
            "pct": 0.0,
            "stage": info.get("stage", ""),
            "indeterminate": True,
            "status_text": info.get("text", ""),
        }
        self.broker.publish("progress", self.progress)

    def _on_done(self, output_path):
        self.phase = "done"
        self.last_output = output_path
        self.progress = {"pct": 100.0, "status_text": "Done"}
        self.broker.publish("progress", self.progress)
        self._append_log("\n=== Done ===\n")
        if output_path and Path(output_path).exists():
            _recent_add(output_path)
            self.broker.publish("recents", {})
        self._publish_run_state()

    def _on_error(self, message):
        self.phase = "error"
        first_line = (message.splitlines()[0] if message
                      else "Unknown error")
        self._append_log(f"\n!!! Error !!!\n{message}\n")
        self._publish_run_state(message=message, first_line=first_line)

    def _on_cancelled(self, message):
        self.phase = "cancelled"
        if message:
            self._append_log(f"\n=== Stopped: {message} ===\n")
        else:
            self._append_log("\n=== Stopped ===\n")
        self._publish_run_state()

    def _batch_item_done(self, output_path, error):
        b = self._batch
        if b is None:
            return
        idx = b["index"]
        in_path = b["items"][idx][0]
        if error is not None:
            first_line = (error.splitlines()[0] if error
                          else "Unknown error")
            b["failed"].append([in_path, first_line])
            self._append_log(
                f"FAILED: {Path(in_path).name}: {first_line}\n")
        else:
            b["succeeded"].append(output_path)
            if output_path and Path(output_path).exists():
                _recent_add(output_path)
                self.broker.publish("recents", {})
        if b["stop"]:
            self._finish_batch(stopped=True)
            return
        b["index"] += 1
        if b["index"] < len(b["items"]):
            self._start_batch_item()
        else:
            self._finish_batch(stopped=False)

    def _batch_cancelled(self, message):
        if self._batch is None:
            return
        if message:
            self._append_log(f"\n=== Stopped: {message} ===\n")
        self._finish_batch(stopped=True)

    def _finish_batch(self, *, stopped):
        b = self._batch
        self._batch = None
        succeeded = b["succeeded"] if b else []
        failed = b["failed"] if b else []
        if succeeded:
            self.last_output = succeeded[-1]
        head = "Batch stopped" if stopped else "Batch complete"
        self.phase = "cancelled" if stopped else "done"
        self.current_file = head
        lines = [f"\n=== {head} ===",
                 f"Transcribed: {len(succeeded)}",
                 f"Failed: {len(failed)}"]
        for in_path, why in failed:
            lines.append(f"  - {Path(in_path).name}: {why}")
        lines.append("Open each transcript from the Library to review "
                     "and label speakers.")
        self._append_log("\n".join(lines) + "\n")
        self.broker.publish("batch_done", {
            "stopped": stopped,
            "succeeded": succeeded,
            "failed": failed,
        })
        self._publish_run_state()


# =====================================================================
# Web UI backend - model manager (list / download / uninstall)
# =====================================================================
#
# Whisper model weights are large (75 MB - 3 GB) and each engine caches
# its own copy in a different place: openai-whisper keeps a single .pt per
# model under ~/.cache/whisper, while faster-whisper and mlx-whisper both
# pull Hugging Face repos into ~/.cache/huggingface/hub. So the same model
# name can occupy disk two or three times over. The model manager reports
# what is cached (with sizes), downloads models on demand - including
# brand-new ones released after this build, for the HF-backed engines -
# and lets the user reclaim space by uninstalling.


def _whisper_cache_dir():
    """Where openai-whisper stores its .pt files. Mirrors whisper's own
    default so we look in exactly the directory it downloads to
    ($XDG_CACHE_HOME/whisper, else ~/.cache/whisper)."""
    root = os.getenv("XDG_CACHE_HOME") or os.path.join(
        os.path.expanduser("~"), ".cache")
    return Path(root) / "whisper"


def _hf_cache_dir():
    """The huggingface_hub download cache (faster-whisper, mlx-whisper).
    Prefer the library's resolved constant so HF_HOME / HF_HUB_CACHE
    overrides are honoured; fall back to the documented default."""
    try:
        from huggingface_hub import constants as _hf_const
        return Path(_hf_const.HF_HUB_CACHE)
    except Exception:
        base = os.getenv("HF_HOME") or os.path.join(
            os.path.expanduser("~"), ".cache", "huggingface")
        return Path(base) / "hub"


# openai-whisper aliases share one .pt file (large -> large-v3,
# turbo -> large-v3-turbo). The authoritative filename is the basename of
# whisper._MODELS[model]; this static map is the fallback when whisper
# can't be imported (e.g. in tests).
_WHISPER_PT_FALLBACK = {
    "tiny.en": "tiny.en.pt", "tiny": "tiny.pt",
    "base.en": "base.en.pt", "base": "base.pt",
    "small.en": "small.en.pt", "small": "small.pt",
    "medium.en": "medium.en.pt", "medium": "medium.pt",
    "large-v1": "large-v1.pt", "large-v2": "large-v2.pt",
    "large-v3": "large-v3.pt", "large": "large-v3.pt",
    "turbo": "large-v3-turbo.pt", "large-v3-turbo": "large-v3-turbo.pt",
}


def _whisper_pt_filename(model):
    """The cache filename openai-whisper uses for `model` (basename of its
    download URL), or None if the name is unknown."""
    try:
        import whisper
        url = whisper._MODELS.get(model)
        if url:
            return os.path.basename(url)
    except Exception:
        pass
    return _WHISPER_PT_FALLBACK.get(model)


def _faster_repo_for(model):
    """The Hugging Face repo faster-whisper caches `model` under.
    Authoritative from faster_whisper.utils._MODELS; falls back to the
    Systran naming used for the standard models."""
    try:
        from faster_whisper import utils as _fw_utils
        repo = getattr(_fw_utils, "_MODELS", {}).get(model)
        if repo:
            return repo
    except Exception:
        pass
    return f"Systran/faster-whisper-{model}"


def _repo_belongs_to_engine(engine, repo_id):
    """Heuristic: does a cached HF repo look like a Whisper model for this
    engine? Used to surface user-downloaded 'new' models that aren't in
    the built-in list."""
    r = repo_id.lower()
    if engine == "faster":
        return "faster-whisper" in r or "faster-distil-whisper" in r
    if engine == "mlx":
        return "whisper" in r and "mlx" in r
    return False


def _hf_repo_sizes():
    """Map {repo_id: (size_on_disk_bytes, [revision_hashes])} for every
    repo in the huggingface_hub cache. Empty when the library or the cache
    directory is absent. Uses the official scanner, which accounts for
    shared blobs so sizes aren't double-counted."""
    out = {}
    try:
        from huggingface_hub import scan_cache_dir
        info = scan_cache_dir()
    except Exception:
        return out
    for repo in info.repos:
        hashes = [r.commit_hash for r in repo.revisions]
        out[repo.repo_id] = (int(repo.size_on_disk), hashes)
    return out


def _download_status_text(model, downloaded, total, speed):
    parts = [f"Downloading '{model}'",
             f"{_humanize_bytes(downloaded)} of {_humanize_bytes(total)}"]
    if speed > 0:
        parts.append(f"{_humanize_bytes(speed)}/s")
    return "   ·   ".join(parts)


# Engines the app can pip-install on demand from the Models view, so the
# base install stays lean. Only the reference OpenAI engine qualifies: it
# drags in PyTorch (~2 GB), so it's opt-in. faster-whisper and mlx-whisper
# ship with the installer and aren't listed here.
_INSTALLABLE_ENGINES = {
    "whisper": {
        "name": "OpenAI Whisper (reference)",
        "note": "The reference engine. Downloads PyTorch (~2 GB); "
                "slower than faster-whisper with essentially identical "
                "output.",
        "packages": ["openai-whisper>=20250625"],
        # For a clean uninstall we also drop torch, which nothing else in
        # this app uses - that's where the space actually goes.
        "uninstall": ["openai-whisper", "torch"],
        # Verified after install: a torch/numpy mismatch can leave the
        # package importable-by-metadata but broken at run time.
        "verify_import": "whisper",
        "approx_mb": 2200,
    },
}


def _engine_install_args(key):
    """pip-install arguments for an installable engine. openai-whisper
    needs numpy kept in a range its numba can import; Intel macOS needs
    the whole torch chain pinned back on top of that."""
    spec = _INSTALLABLE_ENGINES[key]
    pkgs = list(spec["packages"])
    if key == "whisper":
        import platform as _plat
        if sys.platform == "darwin" and _plat.machine() == "x86_64":
            # Intel macOS: PyTorch stopped shipping x86_64 wheels after
            # 2.2.2 (which predates numpy 2.0), so pin the whole chain.
            pkgs += ["torch==2.2.2", "numpy<2", "numba<0.60"]
        else:
            # openai-whisper pulls in numba, which lags numpy and refuses
            # to import under a too-new numpy ("Numba needs NumPy 2.4 or
            # less. Got NumPy 2.5."). faster-whisper may already have
            # pulled a newer numpy into the shared venv, so cap it to a
            # numba-compatible version (pip will downgrade if needed).
            # Bump this ceiling when numba gains newer-numpy support.
            pkgs += ["numpy<2.5"]
    return ["--upgrade", "--prefer-binary", *pkgs]


def _engine_uninstall_pkgs(key):
    spec = _INSTALLABLE_ENGINES[key]
    return list(spec.get("uninstall") or spec["packages"])


def _redetect_engines():
    """Re-probe which engines are importable and refresh the module-level
    AVAILABLE_ENGINES (read live by /api/meta and the controllers), so a
    just-installed engine appears without a restart."""
    import importlib
    importlib.invalidate_caches()
    global AVAILABLE_ENGINES
    AVAILABLE_ENGINES = _detect_engines()
    return AVAILABLE_ENGINES


def _pip_stream(args, on_line, should_cancel):
    """Run `python -m pip <args>` in a subprocess, forwarding each output
    line to on_line and honouring should_cancel(). Raises _CancelledByUser
    on cancel, RuntimeError on a non-zero exit."""
    import subprocess
    import queue as _queue
    env = dict(os.environ)
    if sys.platform == "darwin":
        # Match the installer: a well-formed deployment target keeps
        # Apple's clang from rejecting the build and nudges pip toward
        # pre-built wheels rather than source builds.
        env.setdefault("MACOSX_DEPLOYMENT_TARGET", "11.0")
    proc = subprocess.Popen(
        [sys.executable, "-m", "pip", *args],
        stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True,
        bufsize=1, env=env)
    lines = _queue.Queue()

    def _reader():
        try:
            for line in proc.stdout:
                lines.put(line.rstrip("\n"))
        finally:
            proc.stdout.close()

    reader = threading.Thread(target=_reader, daemon=True)
    reader.start()
    while True:
        try:
            on_line(lines.get(timeout=0.3))
        except _queue.Empty:
            pass
        if should_cancel() and proc.poll() is None:
            proc.terminate()
            try:
                proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                proc.kill()
            raise _CancelledByUser()
        if proc.poll() is not None and lines.empty():
            break
    rc = proc.wait()
    if rc != 0:
        raise RuntimeError(f"pip exited with code {rc}")


def _run_engine_op(action, key, on_line, should_cancel):
    """Install or uninstall an engine's packages with pip, then (for an
    install) verify the engine actually imports. A torch/numpy mismatch
    can leave the package importable-by-metadata but broken at run time -
    verifying here turns that into a clear failure, and we roll the broken
    packages back so the app doesn't offer an engine that crashes."""
    if action == "install":
        _pip_stream(["install", *_engine_install_args(key)],
                    on_line, should_cancel)
        mod = _INSTALLABLE_ENGINES[key].get("verify_import")
        if mod:
            import subprocess
            on_line(f"Verifying {mod} imports cleanly...")
            check = subprocess.run([sys.executable, "-c", f"import {mod}"],
                                   capture_output=True, text=True)
            if check.returncode != 0:
                tail = (check.stderr or check.stdout or "").strip()
                last = tail.splitlines()[-1] if tail else "import failed"
                on_line("Import check failed; rolling back...")
                try:
                    _pip_stream(["uninstall", "-y",
                                 *_engine_uninstall_pkgs(key)],
                                lambda _l: None, lambda: False)
                except Exception:
                    pass
                raise RuntimeError(
                    f"{key} was installed but fails to import: {last}")
    else:
        _pip_stream(["uninstall", "-y", *_engine_uninstall_pkgs(key)],
                    on_line, should_cancel)


def _dedupe_alias_models(raw):
    """Collapse models that resolve to the same storage into one entry.

    `raw` is a list of (model, storage_key, installed, size) in preference
    order. Aliases share weights - openai-whisper's 'large' is the same
    .pt as 'large-v3', 'turbo' the same as 'large-v3-turbo', and the
    HF-backed engines map those pairs to one repo too - so showing both
    is redundant and would double-count. Each unique storage_key becomes
    one entry: the most descriptive (longest) name is canonical, the
    rest are listed as `aliases`."""
    groups, order = {}, []
    for model, skey, installed, size in raw:
        if skey not in groups:
            groups[skey] = []
            order.append(skey)
        groups[skey].append((model, installed, size))
    out = []
    for skey in order:
        members = groups[skey]
        names = [m for m, _i, _s in members]
        canonical = max(names, key=len)
        out.append({
            "model": canonical,
            "aliases": [n for n in names if n != canonical],
            "installed": any(i for _m, i, _s in members),
            "size": max((s for _m, _i, s in members), default=0),
            "storage_key": skey,
            "custom": False,
        })
    return out


class ModelStore:
    """Read-only view of which Whisper models are cached on disk, per
    installed engine, with sizes. `engines` defaults to the detected
    engines but is injectable for tests."""

    def __init__(self, engines=None):
        self.engines = (list(engines) if engines is not None
                        else list(AVAILABLE_ENGINES))

    def payload(self):
        hf_sizes = None
        engines_out = []
        for key, name in self.engines:
            if key == "whisper":
                engines_out.append(self._openai_engine(key, name))
            elif key in ("faster", "mlx"):
                if hf_sizes is None:
                    hf_sizes = _hf_repo_sizes()
                engines_out.append(self._hf_engine(key, name, hf_sizes))
        installed_keys = {k for k, _ in self.engines}
        installable = [
            {"key": key, "name": spec["name"], "note": spec["note"],
             "approx_mb": spec.get("approx_mb", 0)}
            for key, spec in _INSTALLABLE_ENGINES.items()
            if key not in installed_keys
        ]
        # An installed engine that the app knows how to pip-install can also
        # be removed to reclaim space.
        for e in engines_out:
            e["removable"] = e["key"] in _INSTALLABLE_ENGINES
        return {
            "whisper_cache": str(_whisper_cache_dir()),
            "hf_cache": str(_hf_cache_dir()),
            "engines": engines_out,
            "installable": installable,
            "total": sum(e["total"] for e in engines_out),
        }

    def _openai_engine(self, key, name):
        raw = []
        for m in WHISPER_MODELS:
            fn = _whisper_pt_filename(m)
            size, installed = 0, False
            if fn:
                p = _whisper_cache_dir() / fn
                try:
                    if p.exists():
                        size = p.stat().st_size
                        installed = size > 0
                except OSError:
                    pass
            raw.append((m, fn or m, installed, size))
        models = _dedupe_alias_models(raw)
        total = sum(e["size"] for e in models if e["installed"])
        return {"key": key, "name": name, "supports_custom": False,
                "models": models, "total": total}

    def _hf_engine(self, key, name, hf_sizes):
        repo_of = _faster_repo_for if key == "faster" else _mlx_repo_for
        raw, known_repos = [], set()
        for m in WHISPER_MODELS:
            repo = repo_of(m)
            known_repos.add(repo)
            entry = hf_sizes.get(repo)
            raw.append((m, repo, bool(entry), entry[0] if entry else 0))
        models = _dedupe_alias_models(raw)
        total = sum(e["size"] for e in models if e["installed"])
        # Models the user fetched that aren't in the built-in list
        # (e.g. a newly released repo). Surfaced so they can be removed.
        custom = []
        for repo, (size, _hashes) in hf_sizes.items():
            if repo in known_repos or not _repo_belongs_to_engine(key, repo):
                continue
            custom.append({"model": repo, "aliases": [], "installed": True,
                           "size": size, "storage_key": repo, "custom": True})
            total += size
        models.extend(sorted(custom, key=lambda e: e["model"]))
        return {"key": key, "name": name, "supports_custom": True,
                "models": models, "total": total}


def _prefetch_model(engine, model):
    """Download a model's weights without loading them. Runs on the model
    worker thread inside a _DownloadMonitor so progress is reported."""
    if engine == "whisper":
        import whisper
        url = whisper._MODELS.get(model)
        if not url:
            raise RuntimeError(
                f"'{model}' is not a known openai-whisper model.")
        root = str(_whisper_cache_dir())
        os.makedirs(root, exist_ok=True)
        whisper._download(url, root, in_memory=False)
    elif engine == "faster":
        from faster_whisper import download_model
        download_model(model)          # accepts a size name or a HF repo id
    elif engine == "mlx":
        from huggingface_hub import snapshot_download
        snapshot_download(model if "/" in model else _mlx_repo_for(model))
    else:
        raise RuntimeError(f"Engine '{engine}' cannot download models.")


def _uninstall_model(engine, model):
    """Delete the cached weights for (engine, model). Returns bytes freed.
    Raises ApiFail when the model isn't present or a delete fails."""
    if engine == "whisper":
        fn = _whisper_pt_filename(model)
        p = (_whisper_cache_dir() / fn) if fn else None
        if p is None or not p.exists():
            raise ApiFail(404, "not_installed",
                          f"'{model}' is not downloaded for this engine.")
        try:
            size = p.stat().st_size
            p.unlink()
        except OSError as e:
            raise ApiFail(500, "delete_failed", f"Could not delete: {e}")
        return size
    if engine in ("faster", "mlx"):
        if "/" in model:
            repo = model
        else:
            repo = (_faster_repo_for(model) if engine == "faster"
                    else _mlx_repo_for(model))
        return _hf_delete_repo(repo)
    raise ApiFail(400, "bad_engine", f"Engine '{engine}' is not installed.")


def _hf_delete_repo(repo_id):
    entry = _hf_repo_sizes().get(repo_id)
    if not entry:
        raise ApiFail(404, "not_installed",
                      f"'{repo_id}' is not in the download cache.")
    size, hashes = entry
    try:
        from huggingface_hub import scan_cache_dir
        scan_cache_dir().delete_revisions(*hashes).execute()
    except ApiFail:
        raise
    except Exception as e:
        raise ApiFail(500, "delete_failed", f"Could not delete: {e}")
    return size


class ModelController:
    """Downloads and uninstalls model weights on a background thread, one
    job at a time, publishing `model_progress` / `model_done` / `models`
    SSE events. Coupled loosely to the RunController so model work and
    transcription never overlap - loading or deleting weights mid-run
    would be unsafe. `worker_fn`/`store_fn` are injectable for tests."""

    def __init__(self, broker, run_controller=None, *,
                 prefetch_fn=None, uninstall_fn=None, store_fn=None,
                 engine_op_fn=None):
        self.broker = broker
        self.run_controller = run_controller
        self.lock = threading.RLock()
        self.cancel_event = threading.Event()
        self.worker = None
        self.job = None            # {kind, ...} or None (one job at a time)
        self._prefetch = prefetch_fn or _prefetch_model
        self._uninstall = uninstall_fn or _uninstall_model
        self._store = store_fn or (lambda: ModelStore().payload())
        self._engine_op = engine_op_fn or _run_engine_op

    def _run_active(self):
        rc = self.run_controller
        return rc is not None and rc.phase in ("running", "stopping")

    def is_busy(self):
        with self.lock:
            return self.job is not None

    def list_payload(self):
        payload = self._store()
        with self.lock:
            payload["job"] = dict(self.job) if self.job else None
        payload["busy"] = self.is_busy() or self._run_active()
        return payload

    def _engine_installed(self, engine):
        return engine in {k for k, _ in AVAILABLE_ENGINES}

    def start_download(self, engine, model):
        engine = (engine or "").strip()
        model = (model or "").strip()
        with self.lock:
            if self._run_active():
                raise ApiFail(409, "busy",
                              "A transcription is running. Wait for it to "
                              "finish before downloading a model.")
            if self.job is not None:
                raise ApiFail(409, "model_busy",
                              "Another model is already downloading.")
            if not self._engine_installed(engine):
                raise ApiFail(400, "bad_engine",
                              f"Engine '{engine}' is not installed.")
            if not model:
                raise ApiFail(400, "bad_model", "No model name was given.")
            if engine == "whisper":
                # openai-whisper can only fetch its fixed catalogue.
                try:
                    import whisper
                    known = model in whisper._MODELS
                except Exception:
                    known = model in _WHISPER_PT_FALLBACK
                if not known:
                    raise ApiFail(
                        400, "unknown_model",
                        f"'{model}' is not a known openai-whisper model. "
                        "New models can only be added for the faster-whisper "
                        "and mlx-whisper engines.")
            self.cancel_event.clear()
            self.job = {"kind": "download", "engine": engine, "model": model,
                        "phase": "starting", "pct": 0.0, "downloaded": 0,
                        "total": 0, "speed": 0.0,
                        "status_text": f"Preparing to download '{model}'…"}
            self._publish_progress()
            self.worker = threading.Thread(
                target=self._download_worker, args=(engine, model),
                daemon=True, name="transcribr-model-dl")
            self.worker.start()

    def _download_worker(self, engine, model):
        def on_progress(d):
            with self.lock:
                if self.job is None:
                    return
                total = d.get("total") or 0
                downloaded = d.get("downloaded") or 0
                speed = d.get("speed") or 0.0
                self.job.update({
                    "phase": "downloading",
                    "pct": max(0.0, min(100.0,
                                        (downloaded / total * 100)
                                        if total else 0.0)),
                    "downloaded": downloaded, "total": total, "speed": speed,
                    "status_text": _download_status_text(
                        model, downloaded, total, speed),
                })
            self._publish_progress()

        try:
            with _DownloadMonitor(self.cancel_event, model,
                                  on_progress=on_progress):
                self._prefetch(engine, model)
        except _CancelledByUser:
            self._finish(model, cancelled=True)
            return
        except Exception as e:
            self._finish(model, error=f"{type(e).__name__}: {e}")
            return
        self._finish(model)

    def _finish(self, model, *, error=None, cancelled=False):
        with self.lock:
            self.job = None
        if error:
            self.broker.publish("model_done",
                                {"ok": False, "model": model, "error": error})
        elif cancelled:
            self.broker.publish(
                "model_done",
                {"ok": False, "model": model, "cancelled": True})
        else:
            self.broker.publish("model_done", {"ok": True, "model": model})
        self.broker.publish("models", self.list_payload())

    # ----- Engine install / uninstall (pip in the venv) ------------------

    def start_engine_install(self, key):
        return self._start_engine_op("install", key)

    def start_engine_uninstall(self, key):
        return self._start_engine_op("uninstall", key)

    def _start_engine_op(self, action, key):
        key = (key or "").strip()
        with self.lock:
            if self._run_active():
                raise ApiFail(409, "busy",
                              "A transcription is running. Wait for it to "
                              "finish first.")
            if self.job is not None:
                raise ApiFail(409, "model_busy",
                              "Another model or engine operation is already "
                              "in progress.")
            if key not in _INSTALLABLE_ENGINES:
                raise ApiFail(400, "bad_engine",
                              f"'{key}' is not an installable engine.")
            name = _INSTALLABLE_ENGINES[key]["name"]
            installed = self._engine_installed(key)
            if action == "install" and installed:
                raise ApiFail(409, "already_installed",
                              f"{name} is already installed.")
            if action == "uninstall" and not installed:
                raise ApiFail(404, "not_installed",
                              f"{name} isn't installed.")
            verb = "Installing" if action == "install" else "Removing"
            self.cancel_event.clear()
            self.job = {
                "kind": "engine", "action": action, "engine": key,
                "model": name,
                "phase": "installing" if action == "install" else "removing",
                "pct": 0.0,
                "status_text": f"{verb} {name}… (this can take a while)",
            }
            self._publish_progress()
            self.worker = threading.Thread(
                target=self._engine_worker, args=(action, key),
                daemon=True, name="transcribr-engine-op")
            self.worker.start()

    def _engine_worker(self, action, key):
        def on_line(line):
            line = (line or "").strip()
            if not line:
                return
            with self.lock:
                if self.job is None:
                    return
                self.job["status_text"] = line[:200]
            self._publish_progress()

        try:
            self._engine_op(action, key, on_line=on_line,
                            should_cancel=self.cancel_event.is_set)
        except _CancelledByUser:
            self._finish_engine(key, cancelled=True)
            return
        except Exception as e:
            self._finish_engine(key, error=f"{type(e).__name__}: {e}")
            return
        # Success: a newly (un)installed engine changes what's available.
        try:
            _redetect_engines()
        except Exception:
            pass
        self._finish_engine(key)

    def _finish_engine(self, key, *, error=None, cancelled=False):
        with self.lock:
            self.job = None
        name = _INSTALLABLE_ENGINES.get(key, {}).get("name", key)
        if error:
            self.broker.publish("model_done", {"ok": False, "engine": key,
                                               "model": name, "error": error})
        elif cancelled:
            self.broker.publish("model_done",
                                {"ok": False, "engine": key, "model": name,
                                 "cancelled": True})
        else:
            self.broker.publish("model_done", {"ok": True, "engine": key,
                                               "model": name})
            # The engine roster changed - tell clients to refresh /api/meta.
            self.broker.publish("engines_changed", {})
        self.broker.publish("models", self.list_payload())

    def cancel(self):
        with self.lock:
            if self.job is None:
                return False
            self.cancel_event.set()
            return True

    def uninstall(self, engine, model):
        with self.lock:
            if self._run_active():
                raise ApiFail(409, "busy",
                              "A transcription is running. Wait for it to "
                              "finish before removing a model.")
            if self.job is not None:
                raise ApiFail(409, "model_busy",
                              "A model is downloading. Wait for it to "
                              "finish before removing a model.")
        freed = self._uninstall((engine or "").strip(), (model or "").strip())
        self.broker.publish("models", self.list_payload())
        return freed

    def _publish_progress(self):
        with self.lock:
            job = dict(self.job) if self.job else {}
        self.broker.publish("model_progress", job)


# =====================================================================
# Web UI backend - audio preparation for playback
# =====================================================================
#
# The review pane's <audio> element needs a source both webviews can
# actually play and seek: mp3/m4a/aac/wav are served as-is; anything
# else (video containers, flac/ogg/opus) is extracted once to AAC in
# an .m4a, cached under the config dir keyed by (path, size, mtime).
# When the source's audio stream is already AAC the extraction is a
# near-instant remux (-c:a copy). WKWebView refuses media from servers
# without Range support, which is why /audio/current is served through
# bottle's static_file.

_AUDIO_PASSTHROUGH_EXTS = {".mp3", ".m4a", ".aac", ".wav"}
_AUDIO_CACHE_MAX_FILES = 8
_AUDIO_CACHE_MAX_BYTES = 1_000_000_000


def _audio_cache_dir():
    d = _config_dir() / "audio_cache"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _sweep_audio_cache():
    """Keep the newest few extracts within the size budget."""
    try:
        files = sorted(_audio_cache_dir().glob("*.m4a"),
                       key=lambda p: p.stat().st_mtime, reverse=True)
        total = 0
        for i, p in enumerate(files):
            total += p.stat().st_size
            if i >= _AUDIO_CACHE_MAX_FILES or total > _AUDIO_CACHE_MAX_BYTES:
                p.unlink(missing_ok=True)
    except OSError:
        pass


def _audio_cache_key(path):
    import hashlib
    st = Path(path).stat()
    raw = f"{Path(path).resolve()}|{st.st_size}|{st.st_mtime_ns}"
    return hashlib.sha1(raw.encode()).hexdigest()


def _source_audio_codec(path):
    """The first audio stream's codec name, or None."""
    try:
        import av
        with av.open(str(path)) as container:
            if container.streams.audio:
                return container.streams.audio[0].codec_context.name
            return None
    except ImportError:
        pass
    except Exception:
        return None
    ffprobe = shutil.which("ffprobe")
    if not ffprobe:
        return None
    try:
        out = subprocess.run(
            [ffprobe, "-v", "error", "-select_streams", "a:0",
             "-show_entries", "stream=codec_name", "-of", "csv=p=0",
             str(path)],
            capture_output=True, text=True, timeout=30)
        codec = (out.stdout or "").strip()
        return codec or None
    except (OSError, subprocess.TimeoutExpired):
        return None


def _extract_audio_m4a_pyav(src, target):
    """Extract the first audio stream of `src` into an AAC .m4a at
    `target` using PyAV. Remuxes without re-encoding when the source
    audio is already AAC. Raises on any failure (including PyAV being
    absent) - the caller falls back to the ffmpeg binary."""
    import av

    with av.open(str(src)) as in_container:
        if not in_container.streams.audio:
            raise ValueError("No audio stream in file.")
        in_stream = in_container.streams.audio[0]
        codec = in_stream.codec_context.name

        with av.open(str(target), mode="w", format="mp4") as out_container:
            if codec == "aac":
                # Remux: copy packets straight across.
                out_stream = out_container.add_stream(template=in_stream)
                for packet in in_container.demux(in_stream):
                    if packet.dts is None:
                        continue
                    packet.stream = out_stream
                    out_container.mux(packet)
                return

            # Transcode to AAC at the source rate; collapse >2 channels
            # to stereo (review playback doesn't need surround).
            rate = in_stream.codec_context.sample_rate or 44100
            channels = getattr(in_stream.codec_context, "channels", None)
            if channels is None:
                layout = getattr(in_stream.codec_context, "layout", None)
                channels = getattr(layout, "nb_channels", 2)
            layout = "mono" if channels == 1 else "stereo"
            out_stream = out_container.add_stream("aac", rate=rate)
            out_stream.bit_rate = 128_000
            resampler = av.AudioResampler(
                format="fltp", layout=layout, rate=rate)

            def encode_frames(frame):
                for rf in resampler.resample(frame):
                    for pkt in out_stream.encode(rf):
                        out_container.mux(pkt)

            for frame in in_container.decode(in_stream):
                encode_frames(frame)
            encode_frames(None)                      # flush resampler
            for pkt in out_stream.encode(None):      # flush encoder
                out_container.mux(pkt)


class AudioPrep:
    """Prepares one review session's audio for the <audio> element on a
    background thread, reporting progress as audio_status SSE events.
    States: probing -> extracting -> ready | unavailable."""

    def __init__(self, audio_path, broker):
        self.source = audio_path
        self.broker = broker
        self.state = "unavailable" if not audio_path else "probing"
        self.serve_path = None
        self.duration = None
        self.error = None
        if audio_path:
            threading.Thread(target=self._run, daemon=True,
                             name="transcribr-audio").start()

    def status(self):
        out = {"state": self.state}
        if self.state == "ready":
            out["url"] = "/audio/current"
            out["duration"] = self.duration
        if self.error:
            out["error"] = self.error
        return out

    def _set(self, state, error=None):
        self.state = state
        self.error = error
        self.broker.publish("audio_status", self.status())

    def _run(self):
        try:
            src = Path(self.source)
            if not src.exists():
                self._set("unavailable", "Source audio not found.")
                return
            self.duration = get_audio_duration(str(src))
            if src.suffix.lower() in _AUDIO_PASSTHROUGH_EXTS:
                self.serve_path = str(src)
                self._set("ready")
                return
            target = _audio_cache_dir() / f"{_audio_cache_key(src)}.m4a"
            if target.exists():
                self.serve_path = str(target)
                self._set("ready")
                return
            self._set("extracting")
            tmp = target.with_suffix(".part.m4a")

            # PyAV first (no external binary needed), the ffmpeg binary
            # as fallback, and a clear message when neither works.
            try:
                _extract_audio_m4a_pyav(src, tmp)
            except Exception:
                _log("PyAV playback extraction failed, trying ffmpeg:\n"
                     + traceback.format_exc())
                tmp.unlink(missing_ok=True)
                ffmpeg = shutil.which("ffmpeg")
                if not ffmpeg:
                    self._set("unavailable",
                              "Playing this file type needs the 'av' "
                              "package or an ffmpeg install.")
                    return
                codec_args = (["-c:a", "copy"]
                              if _source_audio_codec(src) == "aac"
                              else ["-c:a", "aac", "-b:a", "128k"])
                cmd = [ffmpeg, "-hide_banner", "-loglevel", "error",
                       "-y", "-i", str(src), "-vn", "-map", "0:a:0",
                       *codec_args, str(tmp)]
                proc = subprocess.run(cmd, capture_output=True,
                                      text=True, timeout=1800)
                if proc.returncode != 0 or not tmp.exists():
                    detail = (proc.stderr or "").strip().splitlines()
                    self._set("unavailable",
                              detail[-1] if detail
                              else "Extraction failed.")
                    tmp.unlink(missing_ok=True)
                    return
            os.replace(tmp, target)
            _sweep_audio_cache()
            self.serve_path = str(target)
            self._set("ready")
        except Exception as e:
            _log(f"audio prep failed: {traceback.format_exc()}")
            self._set("unavailable", f"{type(e).__name__}: {e}")


# =====================================================================
# Web UI backend - review session
# =====================================================================

def open_transcript_info(path):
    """Parse a saved transcript into the review-session info dict - the
    module-level twin of WhisperGUI._load_transcript's parsing half.
    Raises ApiFail for parse errors and the >9-speakers refusal."""
    try:
        parsed = read_paragraphs_from_file(path)
    except TranscriptParseError as e:
        raise ApiFail(400, "parse_error", str(e))
    except Exception as e:
        raise ApiFail(400, "parse_error",
                      f"Unexpected error reading {path}:\n"
                      f"{type(e).__name__}: {e}")

    max_speakers = TranscriptModel.MAX_SPEAKERS
    unique_names = []
    for name in parsed["speakers"]:
        if name and name not in unique_names:
            unique_names.append(name)
    if len(unique_names) > max_speakers:
        raise ApiFail(
            422, "too_many_speakers",
            f"This transcript contains {len(unique_names)} distinct "
            f"speakers. The review pane only supports up to "
            f"{max_speakers}. Open the file in Word and consolidate the "
            "speaker labels before loading it back into Transcribr.",
            count=len(unique_names), max=max_speakers)

    name_to_letter = {}
    preset_speaker_names = {}
    for i, name in enumerate(unique_names):
        letter = str(i + 1)
        name_to_letter[name] = letter
        preset_speaker_names[letter] = name
    preset_speakers = [name_to_letter.get(name) if name else None
                       for name in parsed["speakers"]]

    return {
        "paragraphs": parsed["paragraphs"],
        "out_path": path,
        "show_timestamp": parsed["show_timestamp"],
        "title": parsed.get("title"),
        "output_format": ("docx" if path.lower().endswith(".docx")
                          else "txt"),
        "used_partial": False,
        "result": None,
        "extra_formats": [],
        "preset_speakers": preset_speakers,
        "preset_speaker_names": preset_speaker_names,
        "loaded": True,
        "audio_path": _guess_audio_for_transcript(path),
    }


def autosave_restore_info(data):
    """Rebuild the review info dict from an autosave.json payload - the
    twin of _maybe_offer_autosave_restore's restore half."""
    paragraphs = [[tuple(seg) for seg in para]
                  for para in data["paragraphs"]]
    return {
        "paragraphs": paragraphs,
        "out_path": data.get("out_path", ""),
        "show_timestamp": data.get("show_timestamp", True),
        "title": data.get("title"),
        "output_format": data.get("output_format", "txt"),
        "used_partial": False,
        "result": None,
        "extra_formats": [],
        "loaded": bool(data.get("loaded")),
        "diarized": bool(data.get("diarized")),
        "audio_path": data.get("audio_path"),
        "preset_speakers": data.get("speakers") or [],
        "preset_speaker_names": data.get("speaker_names") or {},
    }


class ReviewSession:
    """One open review: a TranscriptModel plus document metadata and
    the save/close orchestration - the web twin of WhisperGUI's review
    host code. Mutations must present the client's last-seen model rev;
    a mismatch raises ApiFail 409 so a stale window can't clobber the
    session."""

    def __init__(self, info, broker, *, log=None):
        self.broker = broker
        self.info = info
        self.out_path = str(info.get("out_path", ""))
        self.show_timestamp = bool(info.get("show_timestamp", True))
        self.title = info.get("title")
        self.output_format = info.get("output_format", "txt")
        self.loaded = bool(info.get("loaded"))
        self.audio_path = info.get("audio_path")
        self.result = info.get("result")
        self.extra_formats = info.get("extra_formats") or []
        self.diarized = bool(info.get("diarized"))
        self._log_line = log or (lambda text: None)
        self.lock = threading.RLock()
        self.closed = False
        self.audio_status_fn = None   # set by the backend (AudioPrep)
        preset_speakers = info.get("preset_speakers") or None
        self.model = TranscriptModel(
            info["paragraphs"],
            speakers=preset_speakers,
            speaker_names=info.get("preset_speaker_names") or None,
            word_conf=info.get("word_conf"),
            on_autosave=self._write_autosave,
        )
        # Reveal enough name fields to cover preset slots (parity with
        # _enter_review_mode's set_visible_speakers call).
        names = info.get("preset_speaker_names") or {}
        slots = ([int(L) for L in names.keys()]
                 + [int(L) for L in (preset_speakers or []) if L])
        if slots:
            self.model.set_visible(max(slots))

    @classmethod
    def from_fresh(cls, info, broker, *, log=None):
        """Session for a fresh transcription: crash-safety pre-save of
        the un-labelled transcript before review opens (port of
        _enter_review_mode lines 5992-6010)."""
        session = cls(info, broker, log=log)
        if info.get("result") is not None:
            try:
                # Defence in depth (the run validators already refuse
                # this): never let the safety copy touch the recording.
                audio = info.get("audio_path")
                if audio and os.path.normcase(
                        str(Path(audio).resolve())) == os.path.normcase(
                        str(Path(info["out_path"]).resolve())):
                    raise RuntimeError(
                        "safety copy would overwrite the source audio")
                write_paragraphs_to_file(
                    info["paragraphs"], Path(info["out_path"]),
                    show_timestamp=info.get("show_timestamp", True),
                    title=info.get("title"),
                    output_format=info["output_format"],
                    speakers=None,
                )
                session._log_line(
                    f"Safety copy saved (no labels yet): "
                    f"{info['out_path']}\n")
            except Exception as e:
                _log(f"Safety save before review failed: {e}")
                session._log_line(
                    f"Warning: could not save safety copy "
                    f"({type(e).__name__}: {e}). Continuing to review.\n")
        return session

    # ----- Payload -------------------------------------------------------

    def payload(self):
        with self.lock:
            m = self.model
            conf = m.confidence_spans()
            paragraphs = []
            for i, para in enumerate(m.paragraphs):
                span = m.playback_span(i)
                if span is None:
                    play = None
                else:
                    start, dur = span
                    play = {"start": start,
                            "end": None if dur is None else start + dur}
                paragraphs.append({
                    "id": m.ids[i],
                    "start": para[0][0] if para else 0.0,
                    "end": para[-1][1] if para else 0.0,
                    "body": m.body(i),
                    "speaker": m.speakers[i],
                    "play": play,
                    "conf": [list(s) for s in conf[i]],
                })
            labelled, total = m.label_counts()
            return {
                "rev": m.rev,
                "out_path": self.out_path,
                "output_format": self.output_format,
                "show_timestamp": self.show_timestamp,
                "title": self.title,
                "loaded": self.loaded,
                "audio": (self.audio_status_fn()
                          if self.audio_status_fn is not None
                          else {"state": "unavailable"}),
                "speaker_names": dict(m.speaker_names),
                "visible_speakers": m.visible_speakers,
                "labelled": labelled,
                "total": total,
                "can_undo": m.can_undo(),
                "can_redo": m.can_redo(),
                "has_word_conf": bool(m.word_conf),
                "diarized": self.diarized,
                "paragraphs": paragraphs,
            }

    def _slim(self, **extra):
        labelled, total = self.model.label_counts()
        return {"rev": self.model.rev, "labelled": labelled,
                "total": total,
                "visible_speakers": self.model.visible_speakers, **extra}

    # ----- Mutations ------------------------------------------------------

    def _check_rev(self, rev):
        if rev != self.model.rev:
            raise ApiFail(409, "stale_rev",
                          "The document changed under you - refetch.",
                          rev=self.model.rev)

    def mutate(self, rev, action, body):
        """Apply one named mutation; returns a slim delta for hot ops
        and the full payload for structural ones."""
        with self.lock:
            self._check_rev(rev)
            m = self.model
            if action == "speaker":
                idx = int(body.get("index", -1))
                letter = body.get("slot")
                if not m.set_speaker(idx, letter):
                    raise ApiFail(400, "bad_request", "Bad index or slot.")
                result = self._slim(index=idx, speaker=letter)
            elif action == "speaker-name":
                if not m.set_speaker_name(str(body.get("slot")),
                                          str(body.get("name", ""))):
                    raise ApiFail(400, "bad_request", "Bad slot.")
                result = self._slim(speaker_names=dict(m.speaker_names))
            elif action == "visible-speakers":
                m.set_visible(int(body.get("n", m.DEFAULT_VISIBLE)))
                result = self._slim()
            elif action == "edit":
                m.commit_edit(int(body.get("index", -1)),
                              str(body.get("text", "")))
                result = self.payload()
            elif action == "split":
                new_idx = m.split(int(body.get("index", -1)),
                                  int(body.get("offset", -1)))
                result = self.payload()
                result["new_index"] = new_idx
            elif action == "merge":
                if not m.merge_with_previous(int(body.get("index", -1))):
                    raise ApiFail(400, "bad_request",
                                  "Can't merge the first paragraph.")
                result = self.payload()
            elif action == "replace-all":
                count = m.replace_all(str(body.get("find", "")),
                                      str(body.get("replace", "")),
                                      bool(body.get("match_case")))
                result = self.payload()
                result["count"] = count
            elif action == "undo":
                m.undo()
                result = self.payload()
            elif action == "redo":
                m.redo()
                result = self.payload()
            else:
                raise ApiFail(404, "unknown_action", action)
            self.broker.publish("review_changed", {"rev": m.rev})
            return result

    # ----- Autosave -------------------------------------------------------

    def _write_autosave(self, paragraphs, speakers, speaker_names):
        """Crash-recovery snapshot in the exact v0.6.0 schema (port of
        _on_review_autosave, incl. the used-or-renamed names filter)."""
        if self.closed:
            return
        defaults = {L: f"Speaker {L}" for L in TranscriptModel.LETTERS}
        used = {letter for letter in speakers if letter}
        names = {letter: name for letter, name in speaker_names.items()
                 if letter in used or name != defaults.get(letter)}
        _autosave_save({
            "out_path": self.out_path,
            "show_timestamp": self.show_timestamp,
            "title": self.title,
            "output_format": self.output_format,
            "loaded": self.loaded,
            "diarized": self.diarized,
            "audio_path": self.audio_path,
            "paragraphs": [[list(seg) for seg in para]
                           for para in paragraphs],
            "speakers": list(speakers),
            "speaker_names": names,
            "saved_at": time.time(),
        })
        self.broker.publish("autosave", {"saved_at": time.time()})

    # ----- Save / close ----------------------------------------------------

    def _write(self, out_path, speakers):
        try:
            write_paragraphs_to_file(
                self.model.paragraphs, Path(out_path),
                show_timestamp=self.show_timestamp,
                title=self.title,
                output_format=self.output_format,
                speakers=speakers,
            )
        except ImportError as e:
            raise ApiFail(500, "missing_dependency", str(e))
        except ApiFail:
            raise
        except Exception as e:
            raise ApiFail(500, "save_failed", f"{type(e).__name__}: {e}")

    def save(self, rev, mode, extra_queue=None):
        """mode: "labels" | "no_labels" | "revision". Returns the final
        out_path. Ports _on_review_save / _on_review_cancel(fresh) /
        _on_review_save_revision."""
        with self.lock:
            self._check_rev(rev)
            m = self.model
            for letter in m.LETTERS:
                m.speaker_names[letter] = m.speaker_names.get(
                    letter, "").strip() or f"Speaker {letter}"
            if mode == "revision":
                if not self.loaded:
                    raise ApiFail(400, "bad_request",
                                  "Revisions only apply to loaded "
                                  "transcripts.")
                out_path = str(_next_revision_path(Path(self.out_path)))
                self._write(out_path, m.resolved_speakers())
                self._log_line(f"Saved revision: {out_path}\n")
            elif mode == "labels":
                out_path = self.out_path
                self._write(out_path, m.resolved_speakers())
                self._log_line(
                    f"Wrote {len(m.paragraphs)} paragraphs (with speaker "
                    f"labels)\nto: {out_path}\n")
            elif mode == "no_labels":
                out_path = self.out_path
                self._write(out_path, None)
                self._log_line(f"Saved without speaker labels: "
                               f"{out_path}\n")
            else:
                raise ApiFail(400, "bad_request", f"Unknown mode {mode}.")

            # Extra formats only accompany full fresh runs (parity).
            if (mode in ("labels", "no_labels") and self.result
                    and self.extra_formats and extra_queue is not None):
                _write_extra_formats(self.result, Path(out_path),
                                     self.extra_formats, extra_queue)
            self._finish(reason="saved", out_path=out_path)
            return out_path

    def close_discard(self):
        """Close a LOADED session without saving (the original file is
        untouched). Fresh sessions use save(mode="no_labels") instead."""
        with self.lock:
            if not self.loaded:
                raise ApiFail(400, "bad_request",
                              "A fresh transcription must be saved "
                              "(with or without labels).")
            self._log_line(f"Closed without saving: {self.out_path}\n")
            self._finish(reason="discarded", out_path=None)

    def _finish(self, *, reason, out_path):
        self.closed = True
        self.model.close()
        _autosave_clear()
        if out_path:
            _recent_add(out_path)
            self.broker.publish("recents", {})
        self.broker.publish("review_closed",
                            {"reason": reason, "out_path": out_path})


# =====================================================================
# Web UI backend - HTTP server (bottle on a threading WSGI server)
# =====================================================================
#
# The web UI is served entirely from a loopback-only HTTP server: the
# static front-end build (webdist/), a JSON API, and a Server-Sent
# Events stream. Every /api/ and /audio/ request must present the
# per-session token - header X-Transcribr-Token, or ?token= for the
# two GETs that cannot set headers (EventSource and <audio>). bottle
# is imported lazily so `import transcribr` works without it.

def _webdist_dir():
    return Path(__file__).resolve().parent / "webdist"


class WebBackend:
    """Everything the web route handlers need, plus server plumbing."""

    def __init__(self, token):
        self.token = token
        self.broker = EventBroker()
        self.controller = RunController(self.broker)
        self.controller.on_paragraphs_ready = self._open_review_fresh
        self.models = ModelController(self.broker, self.controller)
        self.controller.model_busy_check = self.models.is_busy
        self.review = None        # the open ReviewSession, if any
        self.audio = None         # AudioPrep for the open session
        self.window = None        # pywebview window (web mode only)
        self.server = None        # set by serve()
        self.has_window = False   # True once a pywebview window owns us
        self._pump_stop = threading.Event()

    def _open_review_fresh(self, info):
        """paragraphs_ready arrived from the worker (review-before-save
        run): build the session and tell every client. Runs on the pump
        thread."""
        self.controller.phase = "idle"
        self.controller._append_log(
            f"\n{len(info['paragraphs'])} paragraphs ready for review.\n")
        self.controller._publish_run_state()
        session = ReviewSession.from_fresh(
            info, self.broker, log=self.controller._append_log)
        self._attach_audio(session)
        self.review = session
        self.broker.publish("review_opened", {"review": session.payload()})

    def _open_review(self, info):
        """Open a session for a loaded transcript or autosave restore."""
        session = ReviewSession(info, self.broker,
                                log=self.controller._append_log)
        self._attach_audio(session)
        self.review = session
        self.broker.publish("review_opened", {"review": session.payload()})
        return session

    def _attach_audio(self, session):
        self.audio = AudioPrep(session.audio_path, self.broker)
        session.audio_status_fn = self.audio.status

    def _live_review(self):
        session = self.review
        if session is None or session.closed:
            raise ApiFail(404, "no_review", "No review session is open.")
        return session

    def start_pump(self):
        """Drain the worker queue on a daemon thread, dispatching into
        the controller (whose handlers publish SSE) - the web
        replacement for WhisperGUI._poll_queue's after() loop."""
        def pump():
            while not self._pump_stop.is_set():
                try:
                    kind, data = self.controller.queue.get(timeout=0.5)
                except queue.Empty:
                    continue
                try:
                    self.controller.handle_message(kind, data)
                except Exception:
                    _log("web pump error:\n" + traceback.format_exc())
        threading.Thread(target=pump, daemon=True,
                         name="transcribr-pump").start()

    # -- routes ---------------------------------------------------------

    def build_app(self):
        import json
        import bottle

        app = bottle.Bottle()
        backend = self

        @app.hook("before_request")
        def _guard():
            path = bottle.request.path
            if path.startswith("/api/") or path.startswith("/audio/"):
                supplied = (bottle.request.get_header("X-Transcribr-Token")
                            or bottle.request.query.get("token"))
                if supplied != backend.token:
                    raise bottle.HTTPResponse(
                        body=json.dumps({"error": {
                            "code": "unauthorized",
                            "message": "Missing or bad token."}}),
                        status=401,
                        headers={"Content-Type": "application/json"})

        # -- static shell ------------------------------------------------

        @app.get("/")
        def index():
            root = _webdist_dir()
            if not (root / "index.html").exists():
                bottle.response.content_type = "text/html"
                return ("<h1>Transcribr</h1><p>The web interface has not "
                        "been built yet. From the repository, run:</p>"
                        "<pre>cd web && npm install && npm run build</pre>")
            return bottle.static_file("index.html", root=str(root))

        @app.get("/assets/<filepath:path>")
        def assets(filepath):
            return bottle.static_file(
                filepath, root=str(_webdist_dir() / "assets"))

        # -- meta / state ------------------------------------------------

        @app.get("/api/meta")
        def api_meta():
            import importlib.util as _ilu
            return {
                "version": __version__,
                "about_text": ABOUT_TEXT,
                "platform": sys.platform,
                "reveal_label": _REVEAL_LABEL,
                "ui_mode": "webview" if backend.has_window else "browser",
                "engines": ([{"key": "auto", "name": ENGINE_AUTO_NAME}]
                            + [{"key": k, "name": n}
                               for k, n in AVAILABLE_ENGINES]),
                "models": list(WHISPER_MODELS),
                "model_tiers": MODEL_TIERS,
                "languages": [[n, c] for n, c in LANGUAGES],
                "palettes": _PALETTES,
                "ffmpeg": bool(shutil.which("ffmpeg")),
                "pyav": _have_pyav(),
                "diarize_available": (
                    _ilu.find_spec("sherpa_onnx") is not None),
                "diarize_models": [
                    {"id": m["id"], "label": m["label"],
                     "note": m["note"], "size": m["size"]}
                    for m in DIARIZE_VOICE_MODELS],
                "annotate": annotate_enabled(),
                "readme_available": _find_readme() is not None,
            }

        def _fail(e):
            raise bottle.HTTPResponse(
                body=json.dumps({"error": {"code": e.code,
                                           "message": str(e),
                                           **e.extra}}),
                status=e.status,
                headers={"Content-Type": "application/json"})

        def _recents_payload():
            recents = []
            for p in _recent_load():
                try:
                    exists = Path(p).exists()
                except OSError:
                    exists = False
                if exists:
                    recents.append({"path": p, "name": Path(p).name,
                                    "exists": True})
            return recents

        @app.get("/api/state")
        def api_state():
            review = None
            if backend.review is not None and not backend.review.closed:
                review = backend.review.payload()
            return {
                "run": backend.controller.state(),
                "review": review,
                "autosave_pending": bool(_autosave_load()),
                "recents": _recents_payload(),
            }

        # -- review --------------------------------------------------------

        @app.get("/api/review")
        def api_review():
            try:
                return backend._live_review().payload()
            except ApiFail as e:
                _fail(e)

        @app.post("/api/transcripts/open")
        def api_transcripts_open():
            body = bottle.request.json or {}
            path = (body.get("path") or "").strip()
            try:
                if backend.review is not None and not backend.review.closed:
                    raise ApiFail(409, "review_open",
                                  "A review session is already open.")
                if backend.controller.phase in ("running", "stopping"):
                    raise ApiFail(409, "busy",
                                  "Wait for the current transcription "
                                  "to finish.")
                if not path or not Path(path).exists():
                    raise ApiFail(400, "input_not_found",
                                  f"No such file:\n{path}", path=path)
                info = open_transcript_info(path)
                session = backend._open_review(info)
            except ApiFail as e:
                _fail(e)
            backend.controller.last_output = path
            _recent_add(path)
            backend.broker.publish("recents", {})
            backend.controller._append_log(
                f"\nLoaded transcript from: {path}\n")
            return {"review": session.payload()}

        @app.post("/api/review/<action>")
        def api_review_mutate(action):
            if action in ("save", "close"):
                return _review_lifecycle(action)
            body = bottle.request.json or {}
            try:
                session = backend._live_review()
                return session.mutate(int(body.get("rev", -1)),
                                      action, body)
            except ApiFail as e:
                _fail(e)

        def _review_lifecycle(action):
            body = bottle.request.json or {}
            try:
                session = backend._live_review()
                if action == "save":
                    out_path = session.save(
                        int(body.get("rev", -1)),
                        str(body.get("mode", "labels")),
                        extra_queue=backend.controller.queue)
                    backend.controller.last_output = out_path
                    backend.review = None
                    return {"out_path": out_path}
                session.close_discard()
                backend.review = None
                return {"ok": True}
            except ApiFail as e:
                _fail(e)

        @app.get("/api/autosave")
        def api_autosave():
            data = _autosave_load()
            if not data:
                return {"pending": False}
            name = Path(data.get("out_path") or "").name or "transcript"
            return {"pending": True, "name": name,
                    "saved_at": data.get("saved_at")}

        @app.post("/api/autosave/restore")
        def api_autosave_restore():
            data = _autosave_load()
            try:
                if not data:
                    raise ApiFail(404, "no_autosave",
                                  "Nothing to restore.")
                if backend.review is not None and not backend.review.closed:
                    raise ApiFail(409, "review_open",
                                  "A review session is already open.")
                session = backend._open_review(autosave_restore_info(data))
            except ApiFail as e:
                _fail(e)
            backend.controller._append_log(
                f"Restored unsaved review session: {session.out_path}\n")
            return {"review": session.payload()}

        @app.post("/api/autosave/discard")
        def api_autosave_discard():
            _autosave_clear()
            return {"ok": True}

        # -- audio ---------------------------------------------------------

        @app.get("/audio/current")
        def audio_current():
            prep = backend.audio
            if (prep is None or prep.state != "ready"
                    or not prep.serve_path):
                _fail(ApiFail(404, "no_audio", "No audio is ready."))
            p = Path(prep.serve_path)
            mimetype = ("audio/mp4" if p.suffix.lower() == ".m4a"
                        else "auto")
            return bottle.static_file(p.name, root=str(p.parent),
                                      mimetype=mimetype)

        # -- run / batch -------------------------------------------------

        @app.post("/api/run")
        def api_run():
            body = bottle.request.json or {}
            try:
                if backend.review is not None and not backend.review.closed:
                    raise ApiFail(409, "review_open",
                                  "Finish the open review first.")
                run_id = backend.controller.start_single(
                    body.get("input"), body.get("output"),
                    current_settings(), force=bool(body.get("force")))
            except ApiFail as e:
                _fail(e)
            return {"run_id": run_id}

        @app.post("/api/batch")
        def api_batch():
            body = bottle.request.json or {}
            try:
                if backend.review is not None and not backend.review.closed:
                    raise ApiFail(409, "review_open",
                                  "Finish the open review first.")
                run_id = backend.controller.start_batch(
                    body.get("files") or [], current_settings(),
                    force=bool(body.get("force")))
            except ApiFail as e:
                _fail(e)
            return {"run_id": run_id,
                    "count": len(body.get("files") or [])}

        @app.post("/api/run/stop")
        def api_run_stop():
            stopping = backend.controller.stop()
            return {"ok": True, "stopping": stopping}

        # -- model manager ------------------------------------------------

        @app.get("/api/models")
        def api_models():
            return backend.models.list_payload()

        @app.post("/api/models/download")
        def api_models_download():
            body = bottle.request.json or {}
            try:
                backend.models.start_download(
                    str(body.get("engine", "")), str(body.get("model", "")))
                return {"ok": True}
            except ApiFail as e:
                _fail(e)

        @app.post("/api/models/download/cancel")
        def api_models_download_cancel():
            return {"ok": True, "cancelling": backend.models.cancel()}

        @app.post("/api/models/uninstall")
        def api_models_uninstall():
            body = bottle.request.json or {}
            try:
                freed = backend.models.uninstall(
                    str(body.get("engine", "")), str(body.get("model", "")))
                return {"ok": True, "freed": freed}
            except ApiFail as e:
                _fail(e)

        @app.post("/api/models/engine/install")
        def api_models_engine_install():
            body = bottle.request.json or {}
            try:
                backend.models.start_engine_install(
                    str(body.get("engine", "")))
                return {"ok": True}
            except ApiFail as e:
                _fail(e)

        @app.post("/api/models/engine/uninstall")
        def api_models_engine_uninstall():
            body = bottle.request.json or {}
            try:
                backend.models.start_engine_uninstall(
                    str(body.get("engine", "")))
                return {"ok": True}
            except ApiFail as e:
                _fail(e)

        # -- files & dialogs ----------------------------------------------

        @app.post("/api/files/inspect")
        def api_files_inspect():
            body = bottle.request.json or {}
            fmt = current_settings()["output_format"]
            out = []
            for p in (body.get("paths") or [])[:200]:
                try:
                    exists = Path(p).exists()
                except OSError:
                    exists = False
                out.append({
                    "path": p,
                    "name": Path(p).name,
                    "exists": exists,
                    "derived_output": str(Path(p).with_suffix(
                        f".transcript.{fmt}")),
                })
            return {"files": out}

        @app.post("/api/pick")
        def api_pick():
            body = bottle.request.json or {}
            try:
                result = backend.pick_files(
                    kind=body.get("kind", "media"),
                    multiple=bool(body.get("multiple")),
                    initial=body.get("initial"),
                    fmt=body.get("format"))
            except ApiFail as e:
                _fail(e)
            return result

        # -- paths, log, readme -------------------------------------------

        @app.post("/api/path/open")
        def api_path_open():
            body = bottle.request.json or {}
            p = body.get("path") or ""
            if not Path(p).exists():
                _fail(ApiFail(400, "not_found", f"No such file: {p}"))
            _open_path(p)
            return {"ok": True}

        @app.post("/api/path/reveal")
        def api_path_reveal():
            body = bottle.request.json or {}
            p = body.get("path") or ""
            if not Path(p).exists():
                _fail(ApiFail(400, "not_found", f"No such file: {p}"))
            _reveal_path(p)
            return {"ok": True}

        @app.get("/api/log")
        def api_log():
            bottle.response.content_type = "text/plain; charset=utf-8"
            try:
                lines = int(bottle.request.query.get("lines", "500"))
            except ValueError:
                lines = 500
            try:
                text = _log_file_path().read_text(errors="replace")
            except OSError:
                return ""
            return "\n".join(text.splitlines()[-lines:])

        @app.post("/api/log/open")
        def api_log_open():
            body = bottle.request.json or {}
            p = _log_file_path()
            if not p.exists():
                _fail(ApiFail(400, "not_found", "No log file yet."))
            if body.get("reveal"):
                _reveal_path(str(p))
            else:
                _open_path(str(p))
            return {"ok": True}

        @app.get("/api/readme")
        def api_readme():
            p = _find_readme()
            if p is None:
                _fail(ApiFail(404, "not_found", "README not found."))
            try:
                return {"path": str(p), "text": p.read_text(
                    encoding="utf-8", errors="replace")}
            except OSError as e:
                _fail(ApiFail(500, "read_failed", str(e)))

        @app.post("/api/client-error")
        def api_client_error():
            body = bottle.request.json or {}
            _log("web client error: "
                 f"{body.get('message')}\n{body.get('stack', '')}")
            return {"ok": True}

        # -- recents -------------------------------------------------------

        @app.get("/api/recents")
        def api_recents():
            return {"items": _recents_payload()}

        @app.post("/api/recents/clear")
        def api_recents_clear():
            _recent_save([])
            backend.broker.publish("recents", {})
            return {"ok": True}

        # -- developer annotations ----------------------------------------

        @app.get("/api/annotations")
        def api_annotations():
            return {"items": _annotations_load(),
                    "file": str(_annotations_file())}

        @app.post("/api/annotations")
        def api_annotations_add():
            body = bottle.request.json or {}
            try:
                return annotation_add(body)
            except ApiFail as e:
                _fail(e)

        @app.post("/api/annotations/delete")
        def api_annotations_delete():
            body = bottle.request.json or {}
            try:
                annotation_delete(int(body.get("id", -1)))
            except ApiFail as e:
                _fail(e)
            return {"ok": True}

        @app.post("/api/annotations/clear")
        def api_annotations_clear():
            annotations_clear()
            return {"ok": True}

        # -- settings ----------------------------------------------------

        @app.get("/api/settings")
        def api_settings_get():
            return current_settings()

        @app.put("/api/settings")
        def api_settings_put():
            try:
                incoming = bottle.request.json
            except Exception:
                incoming = None
            if not isinstance(incoming, dict):
                raise bottle.HTTPResponse(
                    body=json.dumps({"error": {
                        "code": "bad_request",
                        "message": "Body must be a JSON object."}}),
                    status=400,
                    headers={"Content-Type": "application/json"})
            merged = validate_settings(incoming, base=current_settings())
            _settings_save(merged)
            return merged

        # -- events ------------------------------------------------------

        @app.get("/api/events")
        def api_events():
            bottle.response.content_type = "text/event-stream"
            bottle.response.set_header("Cache-Control", "no-cache")
            raw_last = bottle.request.get_header("Last-Event-ID")
            try:
                last_id = int(raw_last) if raw_last else None
            except ValueError:
                last_id = None
            q, backlog = backend.broker.subscribe(last_id)

            def stream():
                try:
                    # Padding defeats intermediary buffering; retry tells
                    # EventSource how quickly to return after a drop.
                    yield ":" + (" " * 2048) + "\n"
                    yield "retry: 2000\n\n"
                    if backlog is None:
                        yield "event: resync\ndata: {}\n\n"
                    else:
                        for item in backlog:
                            yield _sse_format(item)
                    while True:
                        try:
                            item = q.get(timeout=15.0)
                        except queue.Empty:
                            yield ": ping\n\n"
                            continue
                        yield _sse_format(item)
                finally:
                    backend.broker.unsubscribe(q)

            return stream()

        # -- lifecycle ---------------------------------------------------

        @app.post("/api/shutdown")
        def api_shutdown():
            if backend.server is not None:
                threading.Thread(target=backend.server.shutdown,
                                 daemon=True).start()
            return {"ok": True}

        return app

    # -- native dialogs ----------------------------------------------------

    def pick_files(self, *, kind, multiple=False, initial=None, fmt=None):
        """Show a native file dialog and return real filesystem paths.
        With a pywebview window this is the OS dialog (callable from
        any thread; blocks until dismissed); in --serve browser mode it
        falls back to a one-shot tkinter dialog in a subprocess -
        development only."""
        if self.window is not None:
            return self._pick_files_native(kind=kind, multiple=multiple,
                                           initial=initial, fmt=fmt)
        script = _PICK_SCRIPT
        args = [sys.executable, "-c", script, kind,
                "1" if multiple else "0", initial or "", fmt or ""]
        try:
            proc = subprocess.run(args, capture_output=True, text=True,
                                  timeout=300)
        except (OSError, subprocess.TimeoutExpired) as e:
            raise ApiFail(501, "no_dialog",
                          f"No file dialog available: {e}")
        if proc.returncode != 0:
            raise ApiFail(501, "no_dialog",
                          "No file dialog available in this mode "
                          "(tkinter missing). Type the path instead.")
        import json
        try:
            paths = json.loads(proc.stdout.strip() or "[]")
        except ValueError:
            paths = []
        if not paths:
            return {"cancelled": True}
        if multiple:
            return {"paths": paths}
        return {"path": paths[0]}

    def _pick_files_native(self, *, kind, multiple, initial, fmt):
        import webview
        try:
            dialog_open = webview.FileDialog.OPEN
            dialog_save = webview.FileDialog.SAVE
        except AttributeError:      # older pywebview constants
            dialog_open = webview.OPEN_DIALOG
            dialog_save = webview.SAVE_DIALOG

        if kind == "save-output":
            ext = fmt or "docx"
            result = self.window.create_file_dialog(
                dialog_save,
                save_filename=initial or f"transcript.{ext}")
        else:
            if kind == "transcript":
                types = ("Transcripts (*.docx;*.txt)", "All files (*.*)")
            else:
                types = ("Audio and video (*.mp3;*.wav;*.m4a;*.aac;"
                         "*.flac;*.ogg;*.opus;*.mp4;*.mov;*.mkv;*.avi;"
                         "*.webm)", "All files (*.*)")
            result = self.window.create_file_dialog(
                dialog_open, allow_multiple=multiple, file_types=types)

        if not result:
            return {"cancelled": True}
        paths = ([result] if isinstance(result, str) else
                 [str(p) for p in result])
        if multiple:
            return {"paths": paths}
        return {"path": paths[0]}

    # -- plumbing ---------------------------------------------------------

    def serve(self, host="127.0.0.1", port=0):
        """Bind the server (port 0 = OS-assigned) and return it; the
        caller decides which thread runs serve_forever()."""
        from socketserver import ThreadingMixIn
        from wsgiref.simple_server import (WSGIServer, WSGIRequestHandler,
                                           make_server)

        class _Server(ThreadingMixIn, WSGIServer):
            daemon_threads = True

        class _Handler(WSGIRequestHandler):
            def log_message(self, format, *args):
                pass    # no stdout/stderr chatter (pythonw on Windows)

        self.server = make_server(host, port, self.build_app(),
                                  server_class=_Server,
                                  handler_class=_Handler)
        return self.server


# One-shot native file dialog for --serve (browser) mode, executed as
# `python -c` in a subprocess so no Tk root lingers in the server
# process. argv: kind multiple initial format; prints a JSON list.
_PICK_SCRIPT = r"""
import json, sys
import tkinter as tk
from tkinter import filedialog

kind, multiple, initial, fmt = sys.argv[1:5]
root = tk.Tk()
root.withdraw()
root.attributes("-topmost", True)

MEDIA = [("Audio/video", "*.mp3 *.wav *.m4a *.aac *.flac *.ogg *.opus "
                         "*.mp4 *.mov *.mkv *.avi *.webm"),
         ("All files", "*.*")]
TRANSCRIPT = [("Transcripts", "*.docx *.txt"), ("All files", "*.*")]

paths = []
if kind == "save-output":
    ext = fmt or "docx"
    p = filedialog.asksaveasfilename(
        defaultextension=f".{ext}",
        initialfile=initial or "",
        filetypes=[(ext, f"*.{ext}"), ("All files", "*.*")])
    if p:
        paths = [p]
else:
    types = TRANSCRIPT if kind == "transcript" else MEDIA
    if multiple == "1":
        got = filedialog.askopenfilenames(filetypes=types)
        paths = list(got or [])
    else:
        p = filedialog.askopenfilename(filetypes=types)
        if p:
            paths = [p]
root.destroy()
print(json.dumps(paths))
"""


# =====================================================================
# Entry point
# =====================================================================

def _parse_args(argv):
    import argparse
    p = argparse.ArgumentParser(
        prog="transcribr",
        description="Transcribr - local Whisper transcription GUI.")
    p.add_argument("--web", action="store_true",
                   help="open the app window (this is the default)")
    p.add_argument("--serve", action="store_true",
                   help="run the backend only and print its URL "
                        "(development, or using the app from a browser)")
    p.add_argument("--port", type=int, default=None,
                   help="port for --serve (default 8737)")
    p.add_argument("--dev-token", default=None,
                   help="fixed API token for --serve (default 'dev')")
    p.add_argument("--annotate", action="store_true",
                   help="show the developer annotation overlay (pin "
                        "notes to UI elements; saved to annotations.json "
                        "in the config dir). Also: TRANSCRIBR_ANNOTATE=1, "
                        "or `touch <config dir>/annotate.on` for the "
                        "double-clicked app")
    return p.parse_args(argv)


def main(argv=None):
    global ANNOTATE_MODE
    args = _parse_args(sys.argv[1:] if argv is None else argv)
    ANNOTATE_MODE = bool(
        args.annotate
        or os.environ.get("TRANSCRIBR_ANNOTATE", "").strip() == "1")
    if args.serve:
        return main_serve(args)
    if os.environ.get("TRANSCRIBR_UI", "").strip().lower() == "tk":
        print("Note: the classic Tk interface was retired in 0.7.0; "
              "opening the current interface.", file=sys.stderr)
    return main_web(args)


def _require_web_stack():
    """Exit with instructions rather than a traceback when the web
    interface is requested under a Python that lacks bottle (the
    installers add it to Transcribr's own environment; a bare system
    python usually won't have it)."""
    try:
        import bottle  # noqa: F401
        return
    except ImportError:
        pass
    if sys.platform == "darwin":
        venv_python = (Path.home() / "Library" / "Application Support"
                       / "Transcribr" / "venv" / "bin" / "python")
    elif sys.platform == "win32":
        venv_python = (Path(os.environ.get("LOCALAPPDATA", ""))
                       / "Transcribr" / "venv" / "Scripts" / "python.exe")
    else:
        venv_python = None
    lines = [
        "The web interface needs the 'bottle' package, which is not",
        f"installed for this Python ({sys.executable}).",
        "",
    ]
    if venv_python is not None and venv_python.exists():
        lines += [
            "Easiest fix - run Transcribr through its installed "
            "environment:",
            f'  "{venv_python}" {Path(__file__).name} --serve',
            "",
        ]
    lines += [
        "Or install bottle for this Python:",
        "  python3 -m pip install bottle",
    ]
    sys.exit("\n".join(lines))


def main_serve(args):
    """Web backend without a window: for development and as the
    escape hatch when the native window can't start."""
    _require_web_stack()
    token = args.dev_token or "dev"
    backend = WebBackend(token)
    backend.start_pump()
    server = _bind_or_exit(backend,
                           args.port if args.port is not None else 8737)
    url = f"http://127.0.0.1:{server.server_port}/?token={token}"
    _log(f"--serve listening on {server.server_port}")
    print(f"Transcribr {__version__} web backend running:\n  {url}")
    print("Open the URL in a browser. Press Ctrl+C to stop.")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()


def _bind_or_exit(backend, port):
    """Bind the web server or exit with a plain-English explanation
    (a fixed port can collide with another Transcribr session)."""
    try:
        return backend.serve(port=port)
    except OSError as e:
        sys.exit(
            f"Could not start Transcribr's local server on port {port}:"
            f"\n  {e}\n\n"
            "Another Transcribr session (or another program) is probably "
            "using that port.\nClose it, or pick a different port:\n"
            f"  python3 {Path(__file__).name} --port {port + 1}")


def _attach_drop_handler(backend, window):
    """Bridge OS drag-and-drop to the front end. pywebview reveals the
    dropped files' real filesystem paths on the PYTHON side only
    (pywebviewFullPath); we forward them as a files_dropped SSE event
    and the page applies the 1-file/many-files rule. prevent_default on
    dragover is required or the drop never fires."""
    try:
        from webview.dom import DOMEventHandler

        def on_drop(e):
            files = (e.get("dataTransfer") or {}).get("files") or []
            paths = [f.get("pywebviewFullPath") for f in files
                     if isinstance(f, dict)]
            paths = [p for p in paths if p]
            if paths:
                backend.broker.publish("files_dropped", {"paths": paths})

        window.dom.document.events.dragover += DOMEventHandler(
            lambda e: None, prevent_default=True)
        window.dom.document.events.drop += DOMEventHandler(
            on_drop, prevent_default=True)
    except Exception:
        _log("drop-handler attach failed:\n" + traceback.format_exc())


def main_web(args):
    """The desktop app: the web backend on daemon threads plus a
    native window (WKWebView / WebView2) pointed at it. webview.start()
    must own the main thread - notably on macOS."""
    _require_web_stack()
    try:
        import webview  # noqa: F401
    except ImportError:
        sys.exit(
            "The native window needs the 'pywebview' package, which is "
            f"not installed for this Python ({sys.executable}).\n"
            "Install it with:\n"
            "  python3 -m pip install pywebview\n"
            "or run without a window:  python3 transcribr.py --serve")
    import webview

    # Menu-bar/taskbar identity (ports of the Tk main()'s NSBundle
    # rename; AppUserModelID keeps the Windows taskbar icon grouped).
    if sys.platform == "darwin":
        try:
            from Foundation import NSBundle
            bundle = NSBundle.mainBundle()
            if bundle is not None:
                info = (bundle.localizedInfoDictionary()
                        or bundle.infoDictionary())
                if info is not None:
                    info["CFBundleName"] = "Transcribr"
        except Exception:
            pass
    elif sys.platform == "win32":
        try:
            import ctypes
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(
                "local.transcribr")
        except Exception:
            pass

    if args.dev_token:
        token = args.dev_token
    else:
        import secrets
        token = secrets.token_urlsafe(32)

    backend = WebBackend(token)
    backend.has_window = True
    backend.start_pump()
    server = _bind_or_exit(backend,
                           args.port if args.port is not None else 0)
    threading.Thread(target=server.serve_forever, daemon=True,
                     name="transcribr-http").start()
    url = f"http://127.0.0.1:{server.server_port}/?token={token}"
    _log(f"Transcribr {__version__} web window starting - "
         f"port {server.server_port}")

    window = webview.create_window(
        f"Transcribr {__version__}", url,
        width=1200, height=800, min_size=(900, 620))
    backend.window = window

    def on_loaded():
        _attach_drop_handler(backend, window)

    def on_closed():
        # A pending (debounced) autosave must not be lost with the
        # window; flush it, then stop the server.
        try:
            if backend.review is not None and not backend.review.closed:
                backend.review.model.flush_autosave()
        except Exception:
            _log("autosave flush on close failed:\n"
                 + traceback.format_exc())
        finally:
            threading.Thread(target=server.shutdown,
                             daemon=True).start()

    window.events.loaded += on_loaded
    window.events.closed += on_closed
    webview.start()



if __name__ == "__main__":
    main()
