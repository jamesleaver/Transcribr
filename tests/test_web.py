"""Tests for the web-backend layer: the headless TranscriptModel and
build_worker_params. Run with the whole suite:

    python3 -m unittest discover -s tests -v

Uses only the standard library. These tests are direct ports of the
editing-logic assertions that TestReviewPaneText makes through Tk
widgets; when the Tk pane is retired, this file carries that coverage.

Config-directory access is redirected to a temp directory, mirroring
test_transcribr.py.
"""

import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

try:
    import transcribr as T
except ImportError as e:
    raise unittest.SkipTest(
        f"cannot import transcribr in this interpreter: {e}")

_real_config_dir = T._config_dir
_tmp_config = None


def setUpModule():
    global _tmp_config
    _tmp_config = tempfile.TemporaryDirectory(prefix="transcribr-web-tests-")
    T._config_dir = lambda: Path(_tmp_config.name)


def tearDownModule():
    T._config_dir = _real_config_dir
    if _tmp_config is not None:
        _tmp_config.cleanup()


class _FakeTimer:
    """Stands in for threading.Timer so autosave tests are synchronous."""

    def __init__(self, delay, fn):
        self.delay = delay
        self.fn = fn
        self.started = False
        self.cancelled = False

    def start(self):
        self.started = True

    def cancel(self):
        self.cancelled = True


def _doc():
    """A small three-paragraph document.

    p0 has two segments and real end times; p1 one segment; p2 a
    placeholder-like ~1s span (as loaded transcripts have)."""
    return [
        [(0.0, 5.0, "Hello there."), (5.0, 9.0, "How are you?")],
        [(10.0, 15.5, "I am well thank you.")],
        [(16.0, 17.0, "Good.")],
    ]


_WORD_CONF = [
    (0.0, 0.4, "Hello", 0.9),
    (0.5, 0.9, "there.", 0.30),   # low  (< 0.35)
    (5.0, 5.4, "How", 0.50),      # med  (< 0.6)
    (5.5, 5.9, "are", 0.95),
    (6.0, 6.4, "you?", 0.98),
    (10.0, 10.4, "I", 0.99),
    (10.5, 10.9, "am", 0.99),
    (11.0, 11.4, "well", 0.99),
    (11.5, 11.9, "thank", 0.99),
    (12.0, 12.4, "you.", 0.99),
    (16.0, 16.4, "Good.", 0.99),
]


def _model(**kw):
    return T.TranscriptModel(_doc(), **kw)


# =====================================================================
# TranscriptModel - speakers, undo/redo
# =====================================================================

class TestModelSpeakers(unittest.TestCase):
    def test_set_speaker_and_undo(self):
        m = _model()
        rev0 = m.rev
        self.assertTrue(m.set_speaker(0, "2"))
        self.assertEqual(m.speakers[0], "2")
        self.assertGreater(m.rev, rev0)
        self.assertTrue(m.can_undo())
        self.assertTrue(m.undo())
        self.assertIsNone(m.speakers[0])
        self.assertTrue(m.can_redo())
        self.assertTrue(m.redo())
        self.assertEqual(m.speakers[0], "2")

    def test_clear_speaker(self):
        m = _model()
        m.set_speaker(1, "3")
        m.set_speaker(1, None)
        self.assertIsNone(m.speakers[1])

    def test_high_slot_reveals_name_field(self):
        m = _model()
        self.assertEqual(m.visible_speakers, m.DEFAULT_VISIBLE)
        m.set_speaker(0, "7")
        self.assertEqual(m.visible_speakers, 7)

    def test_invalid_slot_rejected(self):
        m = _model()
        self.assertFalse(m.set_speaker(0, "0"))
        self.assertFalse(m.set_speaker(0, "10"))
        self.assertFalse(m.set_speaker(99, "1"))
        self.assertFalse(m.can_undo())

    def test_visible_speakers_clamped(self):
        m = _model()
        self.assertEqual(m.set_visible(2), m.DEFAULT_VISIBLE)
        self.assertEqual(m.set_visible(99), m.MAX_SPEAKERS)

    def test_visible_initialised_from_used_slots(self):
        m = T.TranscriptModel(_doc(), speakers=["6", None, "2"])
        self.assertEqual(m.visible_speakers, 6)

    def test_rename_no_undo_step_but_revs(self):
        m = _model()
        rev0 = m.rev
        self.assertTrue(m.set_speaker_name("1", "Ms Chen"))
        self.assertEqual(m.speaker_names["1"], "Ms Chen")
        self.assertGreater(m.rev, rev0)
        self.assertFalse(m.can_undo())

    def test_label_counts_and_resolved(self):
        m = _model()
        m.set_speaker(0, "1")
        m.set_speaker(2, "2")
        m.set_speaker_name("1", "  Ms Chen  ")
        m.set_speaker_name("2", "")
        self.assertEqual(m.label_counts(), (2, 3))
        self.assertEqual(m.resolved_speakers(), ["Ms Chen", None, None])

    def test_undo_limit_caps_stack(self):
        m = _model()
        for i in range(m._UNDO_LIMIT + 20):
            m.set_speaker(0, m.LETTERS[i % 9])
        self.assertEqual(len(m._undo_stack), m._UNDO_LIMIT)

    def test_new_mutation_clears_redo(self):
        m = _model()
        m.set_speaker(0, "1")
        m.undo()
        self.assertTrue(m.can_redo())
        m.set_speaker(1, "2")
        self.assertFalse(m.can_redo())


# =====================================================================
# TranscriptModel - structure: edit / merge / split / replace
# =====================================================================

class TestModelStructure(unittest.TestCase):
    def test_commit_edit_collapses_keeping_span(self):
        m = _model()
        self.assertTrue(m.commit_edit(0, "Hello there, how are you?"))
        self.assertEqual(m.paragraphs[0],
                         [(0.0, 9.0, "Hello there, how are you?")])

    def test_commit_edit_empty_is_cancel(self):
        m = _model()
        before = [list(p) for p in m.paragraphs]
        self.assertFalse(m.commit_edit(0, "   "))
        self.assertEqual(m.paragraphs, before)
        self.assertFalse(m.can_undo())

    def test_commit_edit_unchanged_records_no_undo(self):
        m = _model()
        self.assertTrue(m.commit_edit(1, m.body(1)))
        self.assertFalse(m.can_undo())
        # ... but the paragraph still collapses to one segment.
        self.assertEqual(len(m.paragraphs[1]), 1)

    def test_merge_concatenates_and_keeps_first_speaker(self):
        m = _model()
        m.set_speaker(0, "1")
        m.set_speaker(1, "2")
        ids_before = list(m.ids)
        self.assertTrue(m.merge_with_previous(1))
        self.assertEqual(len(m.paragraphs), 2)
        self.assertEqual(len(m.paragraphs[0]), 3)
        self.assertEqual(m.speakers[0], "1")
        self.assertEqual(m.ids, [ids_before[0], ids_before[2]])

    def test_merge_first_paragraph_refused(self):
        m = _model()
        self.assertFalse(m.merge_with_previous(0))

    def test_merge_undo_restores_ids(self):
        m = _model()
        ids_before = list(m.ids)
        m.merge_with_previous(1)
        m.undo()
        self.assertEqual(m.ids, ids_before)
        self.assertEqual(len(m.paragraphs), 3)

    def test_split_interpolates_time_without_word_conf(self):
        m = _model()
        body = m.body(1)                    # "I am well thank you."
        offset = body.index("thank")
        new_idx = m.split(1, offset)
        self.assertEqual(new_idx, 2)
        self.assertEqual(m.body(1), "I am well")
        self.assertEqual(m.body(2), "thank you.")
        # Interpolated: 10.0 + 5.5 * (10/20) = 12.75
        self.assertAlmostEqual(m.paragraphs[1][0][1], 12.75, places=2)
        self.assertAlmostEqual(m.paragraphs[2][0][0], 12.75, places=2)
        self.assertAlmostEqual(m.paragraphs[2][0][1], 15.5, places=2)

    def test_split_prefers_word_timestamps(self):
        m = T.TranscriptModel(_doc(), word_conf=_WORD_CONF)
        body = m.body(1)
        offset = body.index("thank")
        m.split(1, offset)
        # The word "thank" starts at 11.5 in _WORD_CONF - not the 12.75
        # character interpolation would give.
        self.assertAlmostEqual(m.paragraphs[2][0][0], 11.5, places=2)

    def test_split_copies_speaker_and_gets_fresh_id(self):
        m = _model()
        m.set_speaker(1, "4")
        ids_before = list(m.ids)
        new_idx = m.split(1, m.body(1).index("thank"))
        self.assertEqual(m.speakers[new_idx], "4")
        self.assertEqual(m.ids[1], ids_before[1])
        self.assertNotIn(m.ids[new_idx], ids_before)

    def test_split_at_edges_refused(self):
        m = _model()
        self.assertIsNone(m.split(1, 0))
        self.assertIsNone(m.split(1, len(m.body(1))))
        self.assertFalse(m.can_undo())

    def test_split_multisegment_boundary(self):
        m = _model()
        body = m.body(0)                    # "Hello there. How are you?"
        offset = body.index("How")
        new_idx = m.split(0, offset)
        self.assertEqual(m.body(0), "Hello there.")
        self.assertEqual(m.body(new_idx), "How are you?")
        # Clean segment-boundary split keeps the original times.
        self.assertEqual(m.paragraphs[0], [(0.0, 5.0, "Hello there.")])
        self.assertEqual(m.paragraphs[1], [(5.0, 9.0, "How are you?")])

    def test_replace_all_counts_and_single_undo(self):
        m = _model()
        n = m.replace_all("you", "YOU")
        self.assertEqual(n, 2)              # "are you?" and "thank you."
        self.assertIn("YOU", m.body(0))
        self.assertIn("YOU", m.body(1))
        self.assertTrue(m.undo())
        self.assertNotIn("YOU", m.body(0))
        self.assertFalse(m.can_undo())      # exactly one step was pushed

    def test_replace_all_no_matches_records_nothing(self):
        m = _model()
        self.assertEqual(m.replace_all("zebra", "x"), 0)
        self.assertFalse(m.can_undo())

    def test_replace_all_match_case(self):
        m = _model()
        self.assertEqual(m.replace_all("hello", "x", match_case=True), 0)
        self.assertEqual(m.replace_all("Hello", "x", match_case=True), 1)


# =====================================================================
# TranscriptModel - playback spans, confidence, attention
# =====================================================================

class TestModelDerived(unittest.TestCase):
    def test_playback_span_real_ends_padded(self):
        m = _model()
        start, dur = m.playback_span(0)
        self.assertEqual(start, 0.0)
        self.assertAlmostEqual(dur, 9.3, places=2)   # 9.0 + 0.3 tail

    def test_playback_span_synthetic_runs_to_next_start(self):
        # p2's 1.0s span is a placeholder; as the LAST paragraph it
        # plays open-ended...
        m = _model()
        self.assertEqual(m.playback_span(2), (16.0, None))
        # ...while a mid-document placeholder plays to the next start.
        m2 = T.TranscriptModel([
            [(0.0, 1.0, "One.")],
            [(8.0, 9.0, "Two.")],
        ])
        start, dur = m2.playback_span(0)
        self.assertEqual(start, 0.0)
        self.assertAlmostEqual(dur, 8.3, places=2)   # to next start + 0.3

    def test_playback_span_bad_index(self):
        m = _model()
        self.assertIsNone(m.playback_span(99))

    def test_confidence_spans_thresholds(self):
        m = T.TranscriptModel(_doc(), word_conf=_WORD_CONF)
        spans = m.confidence_spans()
        body0 = m.body(0)
        lo = body0.index("there.")
        med = body0.index("How")
        self.assertIn((lo, lo + len("there."), "low"), spans[0])
        self.assertIn((med, med + len("How"), "med"), spans[0])
        self.assertEqual(spans[1], [])       # all high-confidence
        self.assertEqual(len(spans), 3)

    def test_confidence_bails_on_drift(self):
        m = T.TranscriptModel(_doc(), word_conf=_WORD_CONF)
        m.commit_edit(0, "Completely rewritten text")
        spans = m.confidence_spans()
        self.assertEqual(spans[0], [])       # never mis-highlight
        self.assertEqual(spans[1], [])

    def test_confidence_empty_without_word_conf(self):
        m = _model()
        self.assertEqual(m.confidence_spans(), [[], [], []])

    def test_attention_unlabelled(self):
        m = _model()
        m.set_speaker(0, "1")
        self.assertEqual(m.attention_flags(), [False, True, True])

    def test_attention_low_confidence_only_when_shading_on(self):
        m = T.TranscriptModel(_doc(), word_conf=_WORD_CONF)
        for i in range(3):
            m.set_speaker(i, "1")
        self.assertEqual(m.attention_flags(), [False, False, False])
        m.show_confidence = True
        # p0 contains words with prob 0.30 and 0.50 (< 0.6).
        self.assertEqual(m.attention_flags(), [True, False, False])


# =====================================================================
# TranscriptModel - autosave debounce
# =====================================================================

class TestModelAutosave(unittest.TestCase):
    def _make(self):
        fired = []
        timers = []

        def factory(delay, fn):
            t = _FakeTimer(delay, fn)
            timers.append(t)
            return t

        m = T.TranscriptModel(
            _doc(),
            on_autosave=lambda p, s, names: fired.append((p, s, names)),
            timer_factory=factory)
        return m, fired, timers

    def test_mutation_schedules_and_debounces(self):
        m, fired, timers = self._make()
        m.set_speaker(0, "1")
        self.assertEqual(len(timers), 1)
        self.assertEqual(timers[0].delay, 3.0)
        self.assertTrue(timers[0].started)
        m.set_speaker(1, "2")
        self.assertTrue(timers[0].cancelled)   # debounced
        self.assertEqual(len(timers), 2)
        self.assertEqual(fired, [])            # nothing fired yet

    def test_flush_fires_payload(self):
        m, fired, timers = self._make()
        m.set_speaker(0, "3")
        m.set_speaker_name("3", "Witness")
        m.flush_autosave()
        self.assertEqual(len(fired), 1)
        paragraphs, speakers, names = fired[0]
        self.assertEqual(speakers[0], "3")
        self.assertEqual(names["3"], "Witness")

    def test_no_callback_no_timers(self):
        m = T.TranscriptModel(_doc(), timer_factory=lambda d, f: _FakeTimer(d, f))
        m.set_speaker(0, "1")
        self.assertIsNone(m._autosave_timer)

    def test_close_cancels_pending(self):
        m, fired, timers = self._make()
        m.set_speaker(0, "1")
        m.close()
        self.assertTrue(timers[0].cancelled)
        self.assertEqual(fired, [])


# =====================================================================
# RunController (stub worker; no Whisper involved)
# =====================================================================

class TestRunController(unittest.TestCase):
    _ENGINES = [("whisper", "OpenAI Whisper (reference)")]

    def setUp(self):
        self._saved_engines = T.AVAILABLE_ENGINES
        T.AVAILABLE_ENGINES = list(self._ENGINES)
        self.settings = T.default_settings()
        self.tmp = tempfile.TemporaryDirectory(prefix="transcribr-run-")
        self.addCleanup(self.tmp.cleanup)

    def tearDown(self):
        T.AVAILABLE_ENGINES = self._saved_engines

    def _media(self, name):
        p = Path(self.tmp.name) / name
        p.write_bytes(b"\x00")
        return str(p)

    def _controller(self, worker_fn):
        broker = T.EventBroker()
        events_q, _ = broker.subscribe()
        c = T.RunController(broker, worker_fn=worker_fn)
        return c, events_q

    @staticmethod
    def _drain(c, *, until_phases=("done", "error", "cancelled"),
               timeout=5.0):
        """Feed worker-queue messages through the controller until it
        reaches a terminal phase (the pump thread's job, done
        synchronously here)."""
        import time as _time
        deadline = _time.monotonic() + timeout
        while c.phase not in until_phases:
            try:
                kind, data = c.queue.get(timeout=0.2)
            except Exception:
                if _time.monotonic() > deadline:
                    raise AssertionError(
                        f"controller stuck in phase {c.phase!r}")
                continue
            c.handle_message(kind, data)

    @staticmethod
    def _events(events_q):
        out = []
        while True:
            try:
                seq, event, payload = events_q.get_nowait()
            except Exception:
                return out
            out.append(event)

    def test_single_run_reaches_done(self):
        def worker(params, q, cancel):
            q.put(("log", "working...\n"))
            q.put(("done", params["output"]))

        c, events_q = self._controller(worker)
        in_path = self._media("a.mp3")
        run_id = c.start_single(in_path, "", self.settings)
        self.assertEqual(run_id, 1)
        self.assertEqual(c.phase, "running")
        self._drain(c)
        self.assertEqual(c.phase, "done")
        self.assertTrue(c.last_output.endswith("a.transcript.docx"))
        self.assertIn("run_state", self._events(events_q))

    def test_single_run_error_reports_first_line(self):
        def worker(params, q, cancel):
            q.put(("error", "Boom happened\ndetails follow"))

        c, events_q = self._controller(worker)
        c.start_single(self._media("a.mp3"), "", self.settings)
        self._drain(c)
        self.assertEqual(c.phase, "error")
        self.assertIn("Boom happened", "".join(c.log_lines))

    def test_validations(self):
        c, _ = self._controller(lambda p, q, e: q.put(("done", None)))
        with self.assertRaises(T.ApiFail) as cm:
            c.start_single("", "", self.settings)
        self.assertEqual(cm.exception.code, "missing_input")
        with self.assertRaises(T.ApiFail) as cm:
            c.start_single(str(Path(self.tmp.name) / "ghost.mp3"), "",
                           self.settings)
        self.assertEqual(cm.exception.code, "input_not_found")

    def test_overwrite_needs_force(self):
        in_path = self._media("a.mp3")
        out = Path(in_path).with_suffix(".transcript.docx")
        out.write_text("existing")

        done = []

        def worker(params, q, cancel):
            done.append(params["output"])
            q.put(("done", params["output"]))

        c, _ = self._controller(worker)
        with self.assertRaises(T.ApiFail) as cm:
            c.start_single(in_path, "", self.settings)
        self.assertEqual(cm.exception.code, "output_exists")
        c.start_single(in_path, "", self.settings, force=True)
        self._drain(c)
        self.assertEqual(len(done), 1)

    def test_no_engine_maps_to_apifail(self):
        T.AVAILABLE_ENGINES = []
        c, _ = self._controller(lambda p, q, e: None)
        with self.assertRaises(T.ApiFail) as cm:
            c.start_single(self._media("a.mp3"), "", self.settings)
        self.assertEqual(cm.exception.code, "no_engine")

    def test_busy_rejected_while_running(self):
        import threading as _threading
        release = _threading.Event()

        def worker(params, q, cancel):
            release.wait(5)
            q.put(("done", params["output"]))

        c, _ = self._controller(worker)
        c.start_single(self._media("a.mp3"), "", self.settings)
        with self.assertRaises(T.ApiFail) as cm:
            c.start_single(self._media("b.mp3"), "", self.settings)
        self.assertEqual(cm.exception.code, "busy")
        release.set()
        self._drain(c)

    def test_batch_sequences_and_collects_failures(self):
        def worker(params, q, cancel):
            if "bad" in params["input"]:
                q.put(("error", "codec exploded\ntrace"))
            else:
                Path(params["output"]).write_text("t")
                q.put(("done", params["output"]))

        c, events_q = self._controller(worker)
        files = [self._media("one.mp3"), self._media("bad.mp3"),
                 self._media("three.mp3")]
        c.start_batch(files, self.settings)
        self._drain(c)
        self.assertEqual(c.phase, "done")
        self.assertIn("batch_done", self._events(events_q))
        # State captured via the last batch summary in the log.
        log = "".join(c.log_lines)
        self.assertIn("Transcribed: 2", log)
        self.assertIn("Failed: 1", log)
        self.assertIn("bad.mp3: codec exploded", log)

    def test_batch_missing_inputs_rejected(self):
        c, _ = self._controller(lambda p, q, e: None)
        with self.assertRaises(T.ApiFail) as cm:
            c.start_batch([self._media("ok.mp3"),
                           str(Path(self.tmp.name) / "gone.mp3")],
                          self.settings)
        self.assertEqual(cm.exception.code, "missing_inputs")
        self.assertEqual(len(cm.exception.extra["missing"]), 1)

    def test_batch_stop_ends_after_current_file(self):
        def worker(params, q, cancel):
            q.put(("done", params["output"]))

        c, _ = self._controller(worker)
        c.start_batch([self._media("one.mp3"), self._media("two.mp3")],
                      self.settings)
        # Stop lands between files ("even if we're momentarily between
        # files" - the Tk comment); the batch must not advance.
        c.stop()
        self._drain(c)
        self.assertEqual(c.phase, "cancelled")
        log = "".join(c.log_lines)
        self.assertIn("Batch stopped", log)
        self.assertNotIn("two.mp3", log.split("Batch stopped")[0]
                         .split("---")[-1])

    def test_download_message_produces_downloading_progress(self):
        c, events_q = self._controller(lambda p, q, e: None)
        c.handle_message("download", {
            "model": "large-v3",
            "downloaded": 512 * 1024 * 1024,
            "total": 1024 * 1024 * 1024,
            "speed": 20 * 1024 * 1024,
        })
        self.assertEqual(c.progress["stage"], "downloading")
        self.assertAlmostEqual(c.progress["pct"], 50.0, places=3)
        self.assertIn("Downloading model 'large-v3'",
                      c.progress["status_text"])
        self.assertIn("first use only", c.progress["status_text"])
        self.assertIn("progress", self._events(events_q))

    def test_status_message_is_indeterminate(self):
        c, _ = self._controller(lambda p, q, e: None)
        c.handle_message("status", {"stage": "loading",
                                    "text": "Loading model 'x'..."})
        self.assertTrue(c.progress["indeterminate"])
        self.assertEqual(c.progress["stage"], "loading")
        self.assertEqual(c.progress["pct"], 0.0)

    def test_eta_message_tagged_transcribing(self):
        c, _ = self._controller(lambda p, q, e: None)
        c.handle_message("eta", {"audio_done": 30.0, "audio_total": 120.0,
                                 "wall_elapsed": 10.0, "eta_seconds": 30.0,
                                 "speed": 3.0})
        self.assertEqual(c.progress["stage"], "transcribing")
        self.assertAlmostEqual(c.progress["pct"], 25.0, places=3)


# =====================================================================
# First-run download / per-segment progress feedback
# =====================================================================

class TestDownloadFeedback(unittest.TestCase):
    """The _DownloadMonitor / _ProgressWriter machinery that surfaces
    first-run model downloads and drives the progress bar for every
    engine (mlx-whisper included)."""

    @staticmethod
    def _drain_queue(q):
        out = []
        while True:
            try:
                out.append(q.get_nowait())
            except Exception:
                return out

    def test_humanize_bytes(self):
        self.assertEqual(T._humanize_bytes(0), "0 B")
        self.assertEqual(T._humanize_bytes(512), "512 B")
        self.assertEqual(T._humanize_bytes(2 * 1024), "2 KB")
        self.assertEqual(T._humanize_bytes(5 * 1024 * 1024), "5 MB")
        self.assertEqual(T._humanize_bytes(3 * 1024 ** 3), "3.0 GB")

    def test_progress_writer_emits_eta_from_segment_lines(self):
        import queue
        q = queue.Queue()
        w = T._ProgressWriter(q, audio_duration=120.0,
                              transcribe_start=T.time.time() - 10.0)
        w.write("[00:00.000 --> 00:30.000]  Hello world\n")
        kinds, etas = [], []
        for kind, data in self._drain_queue(q):
            kinds.append(kind)
            if kind == "eta":
                etas.append(data)
        self.assertIn("log", kinds)          # verbose line still logged
        self.assertEqual(len(etas), 1)
        self.assertAlmostEqual(etas[0]["audio_done"], 30.0, places=3)
        self.assertEqual(etas[0]["audio_total"], 120.0)

    def test_progress_writer_captures_segments_when_asked(self):
        import queue
        cap = []
        w = T._ProgressWriter(queue.Queue(), None, 0.0, captured=cap)
        w.write("[00:01.000 --> 00:02.500]  One two\n")
        self.assertEqual(cap, [(1.0, 2.5, "One two")])

    def test_download_monitor_forwards_byte_bars(self):
        import io
        import queue
        import threading
        try:
            import tqdm
        except ImportError:
            self.skipTest("tqdm not installed")
        progress, logs = [], []
        cancel = threading.Event()
        orig_update = tqdm.tqdm.update
        with T._DownloadMonitor(cancel, "large-v3",
                                on_progress=progress.append,
                                on_log=logs.append):
            bar = tqdm.tqdm(total=100 * 1024 * 1024, unit="iB",
                            file=io.StringIO())
            bar.n = 50 * 1024 * 1024
            bar.update(0)   # patched update records + emits
            bar.close()
        self.assertTrue(progress, "expected a download progress callback")
        self.assertEqual(progress[-1]["model"], "large-v3")
        self.assertEqual(progress[-1]["total"], 100 * 1024 * 1024)
        self.assertTrue(any("Downloading model 'large-v3'" in t
                            for t in logs))
        # tqdm's method is restored once the context exits.
        self.assertIs(tqdm.tqdm.update, orig_update)

    def test_download_monitor_ignores_tiny_bars(self):
        import io
        import queue
        import threading
        try:
            import tqdm
        except ImportError:
            self.skipTest("tqdm not installed")
        progress = []
        with T._DownloadMonitor(threading.Event(), "tiny",
                                on_progress=progress.append):
            bar = tqdm.tqdm(total=4096, unit="iB", file=io.StringIO())
            bar.update(1024)
            bar.close()
        self.assertFalse(progress)

    def test_download_monitor_cancel_raises(self):
        import io
        import queue
        import threading
        try:
            import tqdm
        except ImportError:
            self.skipTest("tqdm not installed")
        cancel = threading.Event()
        cancel.set()
        with self.assertRaises(T._CancelledByUser):
            with T._DownloadMonitor(cancel, "m", on_progress=lambda _d: None):
                bar = tqdm.tqdm(total=10 * 1024 * 1024, unit="iB",
                                file=io.StringIO())
                bar.update(1)


# =====================================================================
# Model manager - ModelStore (cache discovery) and ModelController
# =====================================================================

class TestModelStore(unittest.TestCase):
    """Inventory shaping: presence, sizes, alias de-duplication, and
    surfacing user-downloaded 'new' models. Runs against a temp whisper
    cache dir and a stubbed huggingface_hub size map, so no engine or
    real download is needed."""

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory(prefix="transcribr-models-")
        self.addCleanup(self.tmp.cleanup)
        self._saved_whisper_dir = T._whisper_cache_dir
        self._saved_hf_sizes = T._hf_repo_sizes
        T._whisper_cache_dir = lambda: Path(self.tmp.name)
        self.addCleanup(
            lambda: setattr(T, "_whisper_cache_dir", self._saved_whisper_dir))
        self.addCleanup(
            lambda: setattr(T, "_hf_repo_sizes", self._saved_hf_sizes))

    def _pt(self, name, size):
        (Path(self.tmp.name) / name).write_bytes(b"\0" * size)

    def test_openai_models_no_alias_duplicates(self):
        self._pt("large-v3.pt", 2000)
        self._pt("small.en.pt", 500)
        T._hf_repo_sizes = lambda: {}
        store = T.ModelStore(engines=[("whisper", "OpenAI Whisper")])
        eng = store.payload()["engines"][0]
        by = {m["model"]: m for m in eng["models"]}
        # The "large"/"turbo" aliases are gone from the list entirely; only
        # the canonical names remain.
        self.assertIn("large-v3", by)
        self.assertIn("large-v3-turbo", by)
        self.assertNotIn("large", by)
        self.assertNotIn("turbo", by)
        self.assertTrue(by["large-v3"]["installed"])
        self.assertFalse(by["tiny"]["installed"])
        self.assertEqual(by["large-v3"]["size"], 2000)
        self.assertEqual(len(eng["models"]), 12)
        self.assertEqual(eng["total"], 2500)
        self.assertFalse(eng["supports_custom"])

    def test_model_alias_normalisation(self):
        self.assertEqual(T._canonical_model("large"), "large-v3")
        self.assertEqual(T._canonical_model("turbo"), "large-v3-turbo")
        self.assertEqual(T._canonical_model("medium"), "medium")
        # A stored legacy alias resolves to its canonical form, not the
        # default, when settings load.
        s = T.validate_settings({"model": "large"})
        self.assertEqual(s["model"], "large-v3")
        self.assertEqual(T.validate_settings({"model": "turbo"})["model"],
                         "large-v3-turbo")

    def test_hf_engine_standard_and_custom(self):
        T._hf_repo_sizes = lambda: {
            "Systran/faster-whisper-large-v3": (1500, ["h1"]),
            "deepdml/faster-whisper-large-v3-turbo-ct2": (900, ["h2"]),
            "mlx-community/whisper-tiny-mlx": (300, ["h3"]),
        }
        store = T.ModelStore(engines=[("faster", "faster-whisper")])
        eng = store.payload()["engines"][0]
        by = {m["model"]: m for m in eng["models"]}
        self.assertTrue(by["large-v3"]["installed"])
        self.assertEqual(by["large-v3"]["size"], 1500)
        customs = [m for m in eng["models"] if m["custom"]]
        self.assertEqual(len(customs), 1)
        self.assertEqual(customs[0]["model"],
                         "deepdml/faster-whisper-large-v3-turbo-ct2")
        # The mlx repo must not be attributed to the faster engine.
        self.assertNotIn("mlx-community/whisper-tiny-mlx",
                         [m["storage_key"] for m in eng["models"]])
        self.assertTrue(eng["supports_custom"])
        self.assertEqual(eng["total"], 2400)   # 1500 + 900

    def test_grand_total_across_engines(self):
        self._pt("tiny.pt", 1000)
        T._hf_repo_sizes = lambda: {
            "mlx-community/whisper-tiny-mlx": (300, ["h"]),
        }
        store = T.ModelStore(engines=[("whisper", "W"), ("mlx", "M")])
        self.assertEqual(store.payload()["total"], 1300)

    def test_dedupe_alias_models_helper(self):
        raw = [
            ("large-v3", "large-v3.pt", True, 3000),
            ("large", "large-v3.pt", False, 0),   # alias, same file
            ("tiny", "tiny.pt", False, 0),
        ]
        out = T._dedupe_alias_models(raw)
        self.assertEqual([e["model"] for e in out], ["large-v3", "tiny"])
        merged = out[0]
        self.assertEqual(merged["aliases"], ["large"])
        # Any member installed -> installed; size is the shared file's.
        self.assertTrue(merged["installed"])
        self.assertEqual(merged["size"], 3000)

    def test_payload_lists_installable_and_removable(self):
        T._hf_repo_sizes = lambda: {}
        # whisper NOT installed -> offered under `installable`; faster isn't
        # an installable engine so it isn't removable.
        p = T.ModelStore(engines=[("faster", "faster-whisper")]).payload()
        self.assertEqual([e["key"] for e in p["installable"]], ["whisper"])
        self.assertFalse(p["engines"][0]["removable"])
        # whisper installed -> not offered, and it's removable.
        p2 = T.ModelStore(
            engines=[("whisper", "OpenAI"), ("faster", "faster")]).payload()
        self.assertEqual(p2["installable"], [])
        wh = next(e for e in p2["engines"] if e["key"] == "whisper")
        self.assertTrue(wh["removable"])

    def test_engine_install_args_base(self):
        from unittest import mock
        # Non-Intel-mac (this test host): openai-whisper plus a numpy cap
        # to keep numba importable.
        with mock.patch.object(T.sys, "platform", "linux"):
            args = T._engine_install_args("whisper")
        self.assertIn("openai-whisper>=20250625", args)
        self.assertIn("--prefer-binary", args)
        self.assertTrue(any(a.startswith("numpy<") for a in args),
                        f"expected a numpy ceiling in {args}")

    def test_engine_install_args_intel_mac_pins_chain(self):
        from unittest import mock
        with mock.patch.object(T.sys, "platform", "darwin"), \
             mock.patch("platform.machine", return_value="x86_64"):
            args = T._engine_install_args("whisper")
        self.assertIn("torch==2.2.2", args)
        self.assertIn("numpy<2", args)
        self.assertIn("numba<0.60", args)

    def test_uninstall_openai_deletes_file(self):
        self._pt("small.en.pt", 777)
        T._hf_repo_sizes = lambda: {}
        freed = T._uninstall_model("whisper", "small.en")
        self.assertEqual(freed, 777)
        self.assertFalse((Path(self.tmp.name) / "small.en.pt").exists())
        # Second attempt: nothing to remove.
        with self.assertRaises(T.ApiFail) as cm:
            T._uninstall_model("whisper", "small.en")
        self.assertEqual(cm.exception.code, "not_installed")


class TestModelController(unittest.TestCase):
    """Download/uninstall job orchestration and its guards, with injected
    prefetch/uninstall/store so no real weights are touched."""

    def setUp(self):
        self._saved_engines = T.AVAILABLE_ENGINES
        T.AVAILABLE_ENGINES = [("faster", "faster-whisper"),
                               ("whisper", "OpenAI Whisper")]
        self.addCleanup(
            lambda: setattr(T, "AVAILABLE_ENGINES", self._saved_engines))
        self.broker = T.EventBroker()
        self.events_q, _ = self.broker.subscribe()
        self.store = {"engines": [], "total": 0,
                      "whisper_cache": "w", "hf_cache": "h"}

    def _events(self):
        out = []
        while True:
            try:
                _seq, event, _payload = self.events_q.get_nowait()
            except Exception:
                return out
            out.append(event)

    def test_download_lifecycle_publishes_and_clears(self):
        import threading
        started = threading.Event()
        release = threading.Event()

        def prefetch(engine, model):
            started.set()
            release.wait(5)

        mc = T.ModelController(self.broker, prefetch_fn=prefetch,
                               store_fn=lambda: dict(self.store))
        mc.start_download("faster", "large-v3")
        self.assertTrue(started.wait(5))
        self.assertTrue(mc.is_busy())
        # list_payload reflects the running job + busy flag.
        payload = mc.list_payload()
        self.assertTrue(payload["busy"])
        self.assertEqual(payload["job"]["model"], "large-v3")
        release.set()
        mc.worker.join(5)
        self.assertFalse(mc.is_busy())
        evs = self._events()
        self.assertIn("model_progress", evs)   # 'starting' publish
        self.assertIn("model_done", evs)
        self.assertIn("models", evs)

    def test_second_download_rejected_while_busy(self):
        import threading
        release = threading.Event()
        mc = T.ModelController(
            self.broker, prefetch_fn=lambda e, m: release.wait(5),
            store_fn=lambda: dict(self.store))
        mc.start_download("faster", "large-v3")
        with self.assertRaises(T.ApiFail) as cm:
            mc.start_download("faster", "tiny")
        self.assertEqual(cm.exception.code, "model_busy")
        with self.assertRaises(T.ApiFail) as cm:
            mc.uninstall("faster", "tiny")
        self.assertEqual(cm.exception.code, "model_busy")
        release.set()
        mc.worker.join(5)

    def test_download_blocked_during_transcription(self):
        class _RC:
            phase = "running"
        mc = T.ModelController(self.broker, run_controller=_RC(),
                               prefetch_fn=lambda e, m: None,
                               store_fn=lambda: dict(self.store))
        with self.assertRaises(T.ApiFail) as cm:
            mc.start_download("faster", "large-v3")
        self.assertEqual(cm.exception.code, "busy")

    def test_unknown_engine_and_openai_model_rejected(self):
        mc = T.ModelController(self.broker, prefetch_fn=lambda e, m: None,
                               store_fn=lambda: dict(self.store))
        with self.assertRaises(T.ApiFail) as cm:
            mc.start_download("nope", "tiny")
        self.assertEqual(cm.exception.code, "bad_engine")
        with self.assertRaises(T.ApiFail) as cm:
            mc.start_download("whisper", "totally-made-up")
        self.assertEqual(cm.exception.code, "unknown_model")

    def test_uninstall_delegates_and_publishes(self):
        calls = []
        mc = T.ModelController(
            self.broker, prefetch_fn=lambda e, m: None,
            uninstall_fn=lambda e, m: calls.append((e, m)) or 4242,
            store_fn=lambda: dict(self.store))
        freed = mc.uninstall("faster", "large-v3")
        self.assertEqual(freed, 4242)
        self.assertEqual(calls, [("faster", "large-v3")])
        self.assertIn("models", self._events())

    def test_engine_install_lifecycle_refreshes_engines(self):
        import threading
        T.AVAILABLE_ENGINES = [("faster", "faster-whisper")]
        release = threading.Event()
        ops = []

        def fake_op(action, key, on_line, should_cancel):
            ops.append((action, key))
            on_line("Collecting openai-whisper")
            release.wait(5)

        orig_detect = T._detect_engines
        T._detect_engines = lambda: [
            ("faster", "faster-whisper"),
            ("whisper", "OpenAI Whisper (reference)")]
        self.addCleanup(lambda: setattr(T, "_detect_engines", orig_detect))

        mc = T.ModelController(self.broker, engine_op_fn=fake_op,
                               store_fn=lambda: dict(self.store))
        mc.start_engine_install("whisper")
        # A second job is refused while the engine op runs.
        with self.assertRaises(T.ApiFail) as cm:
            mc.start_download("faster", "tiny")
        self.assertEqual(cm.exception.code, "model_busy")
        release.set()
        mc.worker.join(5)
        self.assertIn("whisper", [k for k, _ in T.AVAILABLE_ENGINES])
        self.assertEqual(ops, [("install", "whisper")])
        self.assertIn("engines_changed", self._events())

    def test_engine_install_guards(self):
        T.AVAILABLE_ENGINES = [("faster", "faster-whisper"),
                               ("whisper", "OpenAI Whisper (reference)")]
        mc = T.ModelController(self.broker,
                               engine_op_fn=lambda *a, **k: None,
                               store_fn=lambda: dict(self.store))
        with self.assertRaises(T.ApiFail) as cm:
            mc.start_engine_install("whisper")     # already installed
        self.assertEqual(cm.exception.code, "already_installed")
        with self.assertRaises(T.ApiFail) as cm:
            mc.start_engine_install("faster")      # not an installable engine
        self.assertEqual(cm.exception.code, "bad_engine")
        with self.assertRaises(T.ApiFail) as cm:
            mc.start_engine_uninstall("mlx")       # not installable
        self.assertEqual(cm.exception.code, "bad_engine")

    def test_download_error_reported(self):
        import json
        def boom(engine, model):
            raise RuntimeError("network down")
        mc = T.ModelController(self.broker, prefetch_fn=boom,
                               store_fn=lambda: dict(self.store))
        mc.start_download("faster", "large-v3")
        mc.worker.join(5)
        done = None
        while True:
            try:
                _s, ev, payload = self.events_q.get_nowait()
            except Exception:
                break
            if ev == "model_done":
                done = json.loads(payload)
        self.assertIsNotNone(done)
        self.assertFalse(done["ok"])
        self.assertIn("network down", done["error"])


# =====================================================================
# ReviewSession
# =====================================================================

class TestReviewSession(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory(prefix="transcribr-rev-")
        self.addCleanup(self.tmp.cleanup)
        self.broker = T.EventBroker()
        self.events_q, _ = self.broker.subscribe()

    def _events(self):
        names = []
        while True:
            try:
                _seq, event, _payload = self.events_q.get_nowait()
            except Exception:
                return names
            names.append(event)

    def _fresh_session(self, fmt="txt", with_result=True, name="doc"):
        out = str(Path(self.tmp.name) / f"{name}.transcript.{fmt}")
        info = {
            "paragraphs": _doc(),
            "out_path": out,
            "show_timestamp": True,
            "title": "Test doc",
            "output_format": fmt,
            "result": {"segments": []} if with_result else None,
            "extra_formats": [],
            "loaded": False,
            "audio_path": None,
            "word_conf": None,
        }
        return T.ReviewSession.from_fresh(info, self.broker), out

    def test_fresh_session_safety_presaves(self):
        session, out = self._fresh_session()
        self.assertTrue(Path(out).exists())   # unlabelled safety copy
        text = Path(out).read_text()
        self.assertIn("Hello there.", text)
        self.assertEqual(session.payload()["total"], 3)

    def test_save_format_override_reswaps_extension(self):
        # The review pane's Saving options: saving a .txt run as .docx
        # writes the .docx and removes the orphaned .txt safety copy.
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("needs python-docx")
        session, out = self._fresh_session(fmt="txt")
        final = session.save(session.model.rev, "no_labels",
                             out_format="docx")
        self.assertTrue(final.endswith("doc.transcript.docx"))
        self.assertTrue(Path(final).exists())
        self.assertFalse(Path(out).exists(),
                         "orphaned safety copy should be removed")

    def test_save_show_timestamp_override(self):
        session, out = self._fresh_session(fmt="txt")
        session.save(session.model.rev, "no_labels",
                     show_timestamp=False)
        self.assertNotIn("[00:00]", Path(out).read_text())

    def test_disclaimer_switches_when_verified(self):
        session, out = self._fresh_session(fmt="txt")
        session.save(session.model.rev, "no_labels",
                     verified_by="  J. Leaver  ")
        text = Path(out).read_text()
        self.assertIn("verified by J. Leaver", text)
        self.assertNotIn("may not have been checked", text)

    def test_unverified_disclaimer_warns(self):
        session, out = self._fresh_session(fmt="txt")
        session.save(session.model.rev, "no_labels")
        text = Path(out).read_text()
        self.assertIn("may not have been checked by a human", text)
        # The parser strips the disclaimer on re-open either way.
        info = T.open_transcript_info(out)
        bodies = " ".join(seg[2] for para in info["paragraphs"]
                          for seg in para)
        self.assertNotIn("Transcribed using Transcribr", bodies)

    def test_export_writes_pdf_and_keeps_session_open(self):
        try:
            import reportlab  # noqa: F401
        except ImportError:
            self.skipTest("needs reportlab")
        session, out = self._fresh_session(fmt="txt")
        path = session.export(session.model.rev, "pdf",
                              verified_by="J. Leaver")
        self.assertTrue(path.endswith(".pdf"))
        self.assertTrue(Path(path).exists())
        self.assertFalse(session.closed)     # review stays open
        self.assertTrue(Path(out).exists())  # original untouched

    def test_docx_embeds_and_recovers_audio_path(self):
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("needs python-docx")
        audio = Path(self.tmp.name) / "hearing.mp3"
        audio.write_bytes(b"\x00" * 32)
        out = Path(self.tmp.name) / "hearing.transcript.docx"
        T.write_paragraphs_to_file(
            _doc(), out, show_timestamp=True, title="T",
            output_format="docx", audio_path=str(audio))
        # Move the transcript away from its sibling so the filename
        # guess fails and only the embedded metadata can find it.
        moved = Path(self.tmp.name) / "elsewhere"
        moved.mkdir()
        target = moved / out.name
        out.rename(target)
        info = T.open_transcript_info(str(target))
        self.assertEqual(info["audio_path"], str(audio))

    @staticmethod
    def _pdf_text(path):
        """Decode reportlab's ASCII85+Flate content streams to text."""
        import base64
        import re
        import zlib
        blob = ""
        raw = Path(path).read_bytes()
        for m in re.finditer(rb"stream(.*?)endstream", raw, re.S):
            part = m.group(1).strip(b"\r\n")
            try:
                part = zlib.decompress(
                    base64.a85decode(part, adobe=True))
            except Exception:
                try:
                    part = zlib.decompress(part)
                except Exception:
                    pass
            blob += part.decode("latin-1", "ignore")
        return blob

    def test_export_pdf_carries_verified_disclaimer(self):
        # Regression: the verified-by wording must reach exported
        # PDFs exactly as it does saved docx files.
        try:
            import reportlab  # noqa: F401
        except ImportError:
            self.skipTest("needs reportlab")
        session, _ = self._fresh_session(fmt="txt")
        path = session.export(session.model.rev, "pdf",
                              verified_by="J. Leaver")
        text = self._pdf_text(path)
        self.assertIn("verified by J. Leaver", text)
        self.assertNotIn("may not have been checked", text)
        # Without a name the warning stays.
        session2, _ = self._fresh_session(fmt="txt", name="unverified")
        path2 = session2.export(session2.model.rev, "pdf")
        text2 = self._pdf_text(path2)
        self.assertIn("may not have been checked", text2)
        self.assertNotIn("verified by", text2)

    def test_verified_by_sticks_to_session(self):
        # Once set, the name applies to later writes that don't
        # mention it; an explicit empty string clears it again.
        session, _ = self._fresh_session(fmt="txt")
        session.export(session.model.rev, "pdf",
                       verified_by="J. Leaver")
        self.assertEqual(session.verified_by, "J. Leaver")
        p2 = session.export(session.model.rev, "txt")
        self.assertIn("verified by J. Leaver", Path(p2).read_text())
        p3 = session.export(session.model.rev, "txt",
                            verified_by="")
        self.assertIsNone(session.verified_by)
        self.assertIn("may not have been checked",
                      Path(p3).read_text())

    def test_docx_roundtrips_verified_by(self):
        # Saving a verified docx embeds the name in its metadata; the
        # payload of a review opened from that file pre-fills it.
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("needs python-docx")
        session, out = self._fresh_session(fmt="docx")
        session.save(session.model.rev, "no_labels",
                     verified_by="J. Leaver")
        info = T.open_transcript_info(out)
        self.assertEqual(info["verified_by"], "J. Leaver")
        reopened = T.ReviewSession(info, self.broker)
        self.assertEqual(reopened.payload()["verified_by"],
                         "J. Leaver")
        # An unverified docx seeds nothing.
        session2, out2 = self._fresh_session(fmt="docx",
                                             name="plain")
        session2.save(session2.model.rev, "no_labels")
        self.assertIsNone(
            T.open_transcript_info(out2)["verified_by"])

    def test_long_path_metadata_never_kills_save(self):
        # Regression: a long audio path plus a verified name pushed
        # the embedded metadata past python-docx's 255-character
        # property cap and the save crashed. The record must shed
        # detail (keeping the name and a relative audio path) and the
        # save must always succeed.
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("needs python-docx")
        deep = (Path(self.tmp.name) / ("d" * 110) / ("e" * 110))
        deep.mkdir(parents=True)
        audio = deep / "evidence.mp3"   # stem differs from transcript
        audio.write_bytes(b"\x00" * 32)
        out = deep / "doc.transcript.docx"
        self.assertGreater(len(str(audio)), 255)
        T.write_paragraphs_to_file(
            _doc(), out, show_timestamp=True, title="T",
            output_format="docx", verified_by="J. Leaver",
            audio_path=str(audio))
        info = T.open_transcript_info(str(out))
        self.assertEqual(info["verified_by"], "J. Leaver")
        # The relative entry still finds the audio (the filename
        # guess cannot - the stems differ).
        self.assertEqual(info["audio_path"], str(audio))

    def test_relative_audio_survives_folder_move(self):
        # Moving a whole case folder keeps playback working: the
        # transcript-relative path resolves at the new location.
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("needs python-docx")
        before = Path(self.tmp.name) / "case-a"
        before.mkdir()
        audio = before / "evidence.mp3"
        audio.write_bytes(b"\x00" * 32)
        out = before / "doc.transcript.docx"
        T.write_paragraphs_to_file(
            _doc(), out, show_timestamp=True, title="T",
            output_format="docx", audio_path=str(audio))
        after = Path(self.tmp.name) / "case-b"
        before.rename(after)
        info = T.open_transcript_info(str(after / out.name))
        self.assertEqual(info["audio_path"],
                         str(after / "evidence.mp3"))

    def test_absurd_verified_name_skips_metadata_not_save(self):
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("needs python-docx")
        out = Path(self.tmp.name) / "big.transcript.docx"
        T.write_paragraphs_to_file(
            _doc(), out, show_timestamp=True, title="T",
            output_format="docx", verified_by="N" * 300)
        info = T.open_transcript_info(str(out))
        self.assertIsNone(info["verified_by"])

    def test_extra_formats_follow_reviewed_text(self):
        # SRT/VTT/TSV reflect the user's edits; JSON keeps the raw
        # engine result.
        import json as _json
        import queue as queue_mod
        raw = {"segments": [
            {"start": 0.0, "end": 5.0, "text": "Hello there."},
            {"start": 5.0, "end": 9.0, "text": "How are you?"}]}
        out = str(Path(self.tmp.name) / "doc.transcript.txt")
        info = {
            "paragraphs": _doc(), "out_path": out,
            "show_timestamp": True, "title": "Test doc",
            "output_format": "txt", "result": raw,
            "extra_formats": ["srt", "json"], "loaded": False,
            "audio_path": None, "word_conf": None,
        }
        session = T.ReviewSession.from_fresh(info, self.broker)
        session.mutate(session.model.rev, "edit",
                       {"index": 0, "text": "Edited opening line."})
        q = queue_mod.Queue()
        session.save(session.model.rev, "no_labels", extra_queue=q)
        srt = (Path(self.tmp.name) / "doc.srt").read_text()
        self.assertIn("Edited opening line.", srt)
        self.assertNotIn("Hello there.", srt)
        raw_json = _json.loads(
            (Path(self.tmp.name) / "doc.json").read_text())
        texts = [s["text"] for s in raw_json["segments"]]
        self.assertIn("Hello there.", texts)   # raw record untouched

    def test_payload_shape(self):
        session, _ = self._fresh_session()
        p = session.payload()
        self.assertEqual(p["rev"], session.model.rev)
        self.assertEqual(len(p["paragraphs"]), 3)
        first = p["paragraphs"][0]
        self.assertEqual(first["body"], "Hello there. How are you?")
        self.assertIsNone(first["speaker"])
        self.assertAlmostEqual(first["play"]["end"], 9.3, places=2)
        self.assertFalse(p["loaded"])
        self.assertFalse(p["has_word_conf"])

    def test_mutate_rev_guard(self):
        session, _ = self._fresh_session()
        rev = session.model.rev
        session.mutate(rev, "speaker", {"index": 0, "slot": "2"})
        with self.assertRaises(T.ApiFail) as cm:
            session.mutate(rev, "speaker", {"index": 1, "slot": "1"})
        self.assertEqual(cm.exception.code, "stale_rev")
        self.assertIn("review_changed", self._events())

    def test_save_with_labels_writes_resolved_names(self):
        session, out = self._fresh_session()
        rev = session.model.rev
        rev = session.mutate(rev, "speaker", {"index": 0, "slot": "1"})["rev"]
        rev = session.mutate(rev, "speaker-name",
                             {"slot": "1", "name": "Ms Chen"})["rev"]
        result = session.save(rev, "labels",
                              extra_queue=None)
        self.assertEqual(result, out)
        text = Path(out).read_text()
        self.assertIn("Ms Chen", text)
        self.assertTrue(session.closed)
        self.assertIsNone(T._autosave_load())
        self.assertIn("review_closed", self._events())

    def test_save_no_labels_omits_names(self):
        session, out = self._fresh_session()
        rev = session.mutate(session.model.rev, "speaker",
                             {"index": 0, "slot": "1"})["rev"]
        session.save(rev, "no_labels", extra_queue=None)
        self.assertNotIn("Speaker 1", Path(out).read_text())

    def test_revision_only_for_loaded(self):
        session, _ = self._fresh_session()
        with self.assertRaises(T.ApiFail) as cm:
            session.save(session.model.rev, "revision", extra_queue=None)
        self.assertEqual(cm.exception.code, "bad_request")

    def test_loaded_roundtrip_and_revision(self):
        # Save a labelled transcript, reopen it via open_transcript_info,
        # then save a revision next to it.
        session, out = self._fresh_session(fmt="txt")
        rev = session.mutate(session.model.rev, "speaker",
                             {"index": 0, "slot": "1"})["rev"]
        rev = session.mutate(rev, "speaker-name",
                             {"slot": "1", "name": "Witness"})["rev"]
        session.save(rev, "labels", extra_queue=None)

        info = T.open_transcript_info(out)
        self.assertTrue(info["loaded"])
        self.assertEqual(info["preset_speaker_names"], {"1": "Witness"})
        self.assertEqual(info["preset_speakers"][0], "1")

        session2 = T.ReviewSession(info, self.broker)
        p = session2.payload()
        self.assertEqual(p["speaker_names"]["1"], "Witness")
        rev2 = session2.model.rev
        out2 = session2.save(rev2, "revision", extra_queue=None)
        self.assertIn(".rev1.", out2)
        self.assertTrue(Path(out2).exists())
        self.assertTrue(Path(out).exists())   # original untouched

    def test_close_discard_rules(self):
        session, _ = self._fresh_session()
        with self.assertRaises(T.ApiFail):
            session.close_discard()           # fresh must save
        info = T.open_transcript_info(
            self._make_loaded_file("Speaker 1"))
        loaded = T.ReviewSession(info, self.broker)
        loaded.close_discard()
        self.assertTrue(loaded.closed)

    def _make_loaded_file(self, speaker):
        p = Path(self.tmp.name) / "loaded.transcript.txt"
        T.write_paragraphs_to_file(
            _doc(), p, show_timestamp=True, title="x",
            output_format="txt", speakers=[speaker, None, None])
        return str(p)

    def test_too_many_speakers_refused(self):
        paras = [[(float(i), float(i) + 2.0, f"Line {i}.")]
                 for i in range(10)]
        p = Path(self.tmp.name) / "many.transcript.txt"
        T.write_paragraphs_to_file(
            paras, p, show_timestamp=True, title="x", output_format="txt",
            speakers=[f"Person {i}" for i in range(10)])
        with self.assertRaises(T.ApiFail) as cm:
            T.open_transcript_info(str(p))
        self.assertEqual(cm.exception.code, "too_many_speakers")

    def test_timestamp_marks_hide_amend_undo(self):
        session, out = self._fresh_session(fmt="txt")
        rev = session.model.rev
        rev = session.mutate(rev, "timestamp",
                            {"index": 0, "value": "hidden"})["rev"]
        session.mutate(rev, "timestamp", {"index": 1, "value": 125})
        session.save(session.model.rev, "no_labels")
        text = Path(out).read_text()
        # Paragraph 0's stamp is gone; paragraph 1 shows the amended
        # time; paragraph 2 keeps its computed stamp.
        self.assertNotIn("[00:00]", text)
        self.assertIn("[02:05]", text)
        self.assertIn("Hello there.", text)

    def test_reviewed_flag_undo_and_autosave(self):
        session, _ = self._fresh_session(fmt="txt")
        rev = session.mutate(session.model.rev, "reviewed",
                             {"index": 1, "value": True})["rev"]
        self.assertTrue(session.payload()["paragraphs"][1]["reviewed"])
        # Carried in the autosave snapshot and restored.
        session.model.flush_autosave()
        data = T._autosave_load()
        self.assertEqual(data["reviewed"], [False, True, False])
        restored = T.autosave_restore_info(data)
        s2 = T.ReviewSession(restored, self.broker)
        self.assertTrue(s2.payload()["paragraphs"][1]["reviewed"])
        # Undoable.
        session.mutate(rev, "undo", {})
        self.assertFalse(session.model.reviewed[1])

    def test_timestamp_mark_is_undoable(self):
        session, _ = self._fresh_session(fmt="txt")
        rev = session.mutate(session.model.rev, "timestamp",
                            {"index": 0, "value": "hidden"})["rev"]
        self.assertEqual(session.model.ts_marks[0], "hidden")
        session.mutate(rev, "undo", {})
        self.assertIsNone(session.model.ts_marks[0])
        session.mutate(session.model.rev, "redo", {})
        self.assertEqual(session.model.ts_marks[0], "hidden")

    def test_hidden_timestamp_survives_docx_roundtrip(self):
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("needs python-docx")
        session, out = self._fresh_session(fmt="docx")
        session.mutate(session.model.rev, "timestamp",
                       {"index": 1, "value": "hidden"})
        session.save(session.model.rev, "no_labels")
        info = T.open_transcript_info(out)
        marks = info["preset_ts_marks"]
        self.assertEqual(marks[1], "hidden")
        self.assertIsNone(marks[0])
        # The hidden paragraph carries the previous start forward so
        # playback lands nearby.
        self.assertEqual(info["paragraphs"][1][0][0],
                         info["paragraphs"][0][0][0])
        reopened = T.ReviewSession(info, self.broker)
        self.assertEqual(reopened.payload()["paragraphs"][1]["ts"],
                         "hidden")

    def test_apply_retranscribe_splices_and_undoes(self):
        session, _ = self._fresh_session(fmt="txt")
        m = session.model
        rev = session.mutate(m.rev, "speaker",
                            {"index": 2, "slot": "2"})["rev"]
        before = [m.body(i) for i in range(len(m.paragraphs))]
        segs = [(0.5, 2.0, "Corrected one."), (2.5, 4.5, "Corrected two.")]
        kept = session.apply_retranscribe(rev, [(0, 1, segs)], gap=10.0)
        self.assertEqual(kept, 1)
        payload = session.payload()
        bodies = [p["body"] for p in payload["paragraphs"]]
        # 10s gap threshold folds both segments into one paragraph.
        self.assertEqual(bodies[0], "Corrected one. Corrected two.")
        self.assertEqual(bodies[-1], before[-1])   # outside untouched
        # Replacements arrive unlabelled; the paragraph after the range
        # keeps its speaker.
        self.assertIsNone(payload["paragraphs"][0]["speaker"])
        self.assertEqual(payload["paragraphs"][-1]["speaker"], "2")
        # One undo restores the original text.
        payload = session.mutate(session.model.rev, "undo", {})
        self.assertEqual([p["body"] for p in payload["paragraphs"]],
                         before)

    def test_apply_retranscribe_refuses_stale_rev_and_empty(self):
        session, _ = self._fresh_session(fmt="txt")
        rev = session.model.rev
        session.mutate(rev, "speaker", {"index": 0, "slot": "1"})
        with self.assertRaises(T.ApiFail) as ctx:
            session.apply_retranscribe(rev, [(0, 0, [(0.0, 1.0, "x")])],
                                       gap=2.0)
        self.assertEqual(ctx.exception.code, "stale_rev")
        with self.assertRaises(T.ApiFail) as ctx:
            session.apply_retranscribe(session.model.rev,
                                       [(0, 0, [])], gap=2.0)
        self.assertEqual(ctx.exception.code, "no_speech")

    def test_retrans_preserves_reviewed_island(self):
        # A reviewed paragraph inside the selection is kept verbatim;
        # the runs on either side are re-transcribed independently.
        session, _ = self._fresh_session(fmt="txt")
        m = session.model
        # Three paragraphs at 0, 5, 10s; mark the middle one reviewed.
        session.mutate(m.rev, "reviewed", {"index": 1, "value": True})
        kept_body = m.body(1)
        events = []
        orig = self.broker.publish
        self.broker.publish = lambda k, d: (events.append((k, d)),
                                            orig(k, d))
        orig_slice = T.retranscribe_slice
        spans = []

        def fake_slice(audio, start, end, settings, q, cancel):
            spans.append((round(start, 2), end))
            return [(start + 0.1, start + 0.9, f"Redone at {int(start)}.")]

        T.retranscribe_slice = fake_slice
        backend = type("B", (), {"broker": self.broker,
                                 "retrans": None})()
        try:
            import time as _t
            job = T.RetransJob(backend, session,
                               dict(T.current_settings()),
                               m.rev, 0, 2)
            deadline = _t.monotonic() + 10
            while job.state == "running":
                if _t.monotonic() > deadline:
                    raise AssertionError("stuck")
                _t.sleep(0.02)
        finally:
            T.retranscribe_slice = orig_slice
            self.broker.publish = orig
        self.assertEqual(job.state, "done")
        # Two runs re-transcribed (para 0 and para 2); the middle stays.
        self.assertEqual(len(spans), 2)
        bodies = [m.body(i) for i in range(len(m.paragraphs))]
        self.assertEqual(bodies[0], "Redone at 0.")
        self.assertEqual(m.body(1), kept_body)
        self.assertTrue(m.reviewed[1])
        self.assertTrue(bodies[-1].startswith("Redone at "))

    def test_retrans_all_reviewed_is_refused(self):
        session, _ = self._fresh_session(fmt="txt")
        m = session.model
        for i in range(len(m.paragraphs)):
            session.mutate(m.rev, "reviewed", {"index": i, "value": True})
        events = []
        orig = self.broker.publish
        self.broker.publish = lambda k, d: (events.append((k, d)),
                                            orig(k, d))
        called = {"n": 0}
        orig_slice = T.retranscribe_slice

        def fake_slice(*a, **k):
            called["n"] += 1
            return []

        T.retranscribe_slice = fake_slice
        backend = type("B", (), {"broker": self.broker,
                                 "retrans": None})()
        try:
            import time as _t
            job = T.RetransJob(backend, session,
                               dict(T.current_settings()),
                               m.rev, 0, len(m.paragraphs) - 1)
            deadline = _t.monotonic() + 10
            while job.state == "running":
                if _t.monotonic() > deadline:
                    raise AssertionError("stuck")
                _t.sleep(0.02)
        finally:
            T.retranscribe_slice = orig_slice
            self.broker.publish = orig
        self.assertEqual(job.state, "failed")
        self.assertEqual(called["n"], 0)   # engine never ran
        self.assertTrue(any("reviewed" in d.get("message", "")
                            for k, d in events if k == "retrans"))

    def test_retrans_job_end_to_end_with_stub_engine(self):
        import time as _time
        session, _ = self._fresh_session(fmt="txt")
        events = []
        orig_publish = self.broker.publish
        self.broker.publish = lambda kind, data: (
            events.append((kind, data)), orig_publish(kind, data))
        backend = type("B", (), {"broker": self.broker,
                                 "retrans": None})()
        orig = T.retranscribe_slice
        calls = {}

        def fake_slice(audio_path, start, end, settings, q, cancel):
            calls["span"] = (start, end)
            calls["condition"] = settings["condition_on_previous_text"]
            calls["model"] = settings["model"]
            return [(start + 0.1, start + 1.5, "Fixed text.")]

        T.retranscribe_slice = fake_slice
        expected_span = (session.model.paragraphs[1][0][0],
                         session.model.paragraphs[2][0][0])
        try:
            settings = dict(T.current_settings())
            settings["condition_on_previous_text"] = False
            settings["model"] = "small"
            job = T.RetransJob(backend, session, settings,
                               session.model.rev, 1, 1)
            deadline = _time.monotonic() + 10
            while job.state == "running":
                if _time.monotonic() > deadline:
                    raise AssertionError("job stuck")
                _time.sleep(0.02)
        finally:
            T.retranscribe_slice = orig
            self.broker.publish = orig_publish
        self.assertEqual(job.state, "done")
        self.assertFalse(calls["condition"])
        self.assertEqual(calls["model"], "small")
        # Slice span: paragraph 1's start through paragraph 2's start.
        self.assertEqual(calls["span"], expected_span)
        self.assertIn("Fixed text.", session.model.body(1))
        self.assertIn(("retrans", {"state": "done",
                                   "message": "Done - the selection was "
                                   "replaced (undo reverses it)."}),
                      events)

    def test_recents_reviewed_and_verified_flags(self):
        T._recent_save([])
        T._recent_add("/tmp/a.transcript.docx")
        items = T._recent_load()
        self.assertFalse(items[0]["reviewed"])
        # A review-pane save marks it; re-opening later keeps it.
        T._recent_add("/tmp/a.transcript.docx", reviewed=True,
                      verified=True)
        T._recent_add("/tmp/a.transcript.docx")
        items = T._recent_load()
        self.assertTrue(items[0]["reviewed"])
        self.assertTrue(items[0]["verified"])
        # A fresh transcription of the same output resets the badge.
        T._recent_add("/tmp/a.transcript.docx", reviewed=False,
                      verified=False)
        self.assertFalse(T._recent_load()[0]["reviewed"])
        # Pre-0.9.5 files stored bare strings.
        import json
        T._recent_file().write_text(json.dumps(["/tmp/old.txt"]))
        items = T._recent_load()
        self.assertEqual(items[0]["path"], "/tmp/old.txt")
        self.assertFalse(items[0]["reviewed"])

    def test_autosave_schema_roundtrip(self):
        session, out = self._fresh_session()
        rev = session.mutate(session.model.rev, "speaker",
                             {"index": 0, "slot": "3"})["rev"]
        session.mutate(rev, "speaker-name", {"slot": "3", "name": "Q C"})
        session.model.flush_autosave()
        data = T._autosave_load()
        self.assertEqual(
            set(data.keys()),
            {"out_path", "show_timestamp", "title", "output_format",
             "loaded", "diarized", "verified_by", "audio_path",
             "paragraphs", "speakers", "speaker_names", "ts_marks",
             "reviewed", "saved_at"})   # v0.6.0 schema + 0.9.x additions
        self.assertEqual(data["speakers"][0], "3")
        self.assertEqual(data["speaker_names"], {"3": "Q C"})
        restored = T.autosave_restore_info(data)
        session2 = T.ReviewSession(restored, self.broker)
        p = session2.payload()
        self.assertEqual(p["paragraphs"][0]["speaker"], "3")
        self.assertEqual(p["speaker_names"]["3"], "Q C")
        self.assertEqual(p["visible_speakers"], 4)


# =====================================================================
# HTTP API (real server on an ephemeral port)
# =====================================================================

def _have_bottle():
    try:
        import bottle  # noqa: F401
        return True
    except ImportError:
        return False


@unittest.skipUnless(_have_bottle(), "needs bottle")
class TestHttpApi(unittest.TestCase):
    TOKEN = "test-token"

    @classmethod
    def setUpClass(cls):
        import threading
        cls.backend = T.WebBackend(cls.TOKEN)
        # Belt and braces: no test may ever reach the real model
        # download or pip install/uninstall machinery, whatever the
        # guards do. A regression that slips past a guard must fail
        # loudly instead of fetching real weights or uninstalling
        # packages from the developer's venv. (_uninstall stays real:
        # its own not-installed guard is part of what the tests assert,
        # and it only touches the cache dirs, which tests redirect.)
        def _blocked(*_a, **_k):
            raise AssertionError(
                "real model download / pip operation invoked from a test")
        cls.backend.models._prefetch = _blocked
        cls.backend.models._engine_op = _blocked
        cls.server = cls.backend.serve(port=0)
        cls.thread = threading.Thread(target=cls.server.serve_forever,
                                      daemon=True)
        cls.thread.start()
        cls.base = f"http://127.0.0.1:{cls.server.server_port}"

    @classmethod
    def tearDownClass(cls):
        cls.server.shutdown()
        cls.server.server_close()

    def setUp(self):
        # Each test starts with no open review and clean config files.
        self.backend.review = None
        T._autosave_clear()
        T._recent_save([])
        self.tmp = tempfile.TemporaryDirectory(prefix="transcribr-http-")
        self.addCleanup(self.tmp.cleanup)

    # -- tiny urllib client ----------------------------------------------

    def _req(self, method, path, body=None, *, token=True, raw=False):
        import json as _json
        import urllib.request
        import urllib.error
        headers = {}
        if token:
            headers["X-Transcribr-Token"] = self.TOKEN
        data = None
        if body is not None:
            data = _json.dumps(body).encode()
            headers["Content-Type"] = "application/json"
        req = urllib.request.Request(self.base + path, data=data,
                                     headers=headers, method=method)
        try:
            with urllib.request.urlopen(req, timeout=10) as resp:
                payload = resp.read()
                return resp.status, (payload if raw
                                     else _json.loads(payload or b"{}"))
        except urllib.error.HTTPError as e:
            payload = e.read()
            try:
                return e.code, _json.loads(payload or b"{}")
            except ValueError:
                return e.code, {}

    # -- tests -------------------------------------------------------------

    def test_api_requires_token(self):
        status, _ = self._req("GET", "/api/meta", token=False)
        self.assertEqual(status, 401)
        status, _ = self._req("GET", "/api/meta")
        self.assertEqual(status, 200)

    def test_meta_shape(self):
        _, meta = self._req("GET", "/api/meta")
        self.assertEqual(meta["version"], T.__version__)
        self.assertIn("light", meta["palettes"])
        self.assertIn("dark", meta["palettes"])
        self.assertIsInstance(meta["models"], list)

    def test_models_endpoint_shape_and_guards(self):
        # Point the model cache at empty temp dirs so nothing real is
        # ever read or deleted by this test.
        saved_dir, saved_sizes = T._whisper_cache_dir, T._hf_repo_sizes
        T._whisper_cache_dir = lambda: Path(self.tmp.name)
        T._hf_repo_sizes = lambda: {}
        self.addCleanup(lambda: setattr(T, "_whisper_cache_dir", saved_dir))
        self.addCleanup(lambda: setattr(T, "_hf_repo_sizes", saved_sizes))
        # Pin the installed-engine list to empty so every "isn't
        # installed" guard below holds on ANY machine. Without this, a
        # dev box that has faster-whisper installed would accept the
        # download request and spawn a real ~75MB fetch into the user's
        # HF cache, and one with openai-whisper installed would let the
        # engine-uninstall request pip-remove it from the real venv
        # (both jobs also 409 unrelated tests while in flight).
        saved_engines = T.AVAILABLE_ENGINES
        T.AVAILABLE_ENGINES = []
        self.addCleanup(lambda: setattr(T, "AVAILABLE_ENGINES",
                                        saved_engines))

        status, payload = self._req("GET", "/api/models")
        self.assertEqual(status, 200)
        for k in ("engines", "total", "whisper_cache", "hf_cache", "busy"):
            self.assertIn(k, payload)
        self.assertIsInstance(payload["engines"], list)
        # Downloading for an engine that isn't installed -> 400.
        status, err = self._req("POST", "/api/models/download",
                                {"engine": "faster", "model": "tiny"})
        self.assertEqual(status, 400)
        self.assertEqual(err["error"]["code"], "bad_engine")
        # Uninstalling a model that isn't present -> 404 (empty temp cache).
        status, err = self._req("POST", "/api/models/uninstall",
                                {"engine": "whisper", "model": "tiny"})
        self.assertEqual(status, 404)
        self.assertEqual(err["error"]["code"], "not_installed")
        # Cancel with nothing running is a harmless no-op.
        status, res = self._req("POST", "/api/models/download/cancel")
        self.assertEqual(status, 200)
        self.assertFalse(res["cancelling"])
        # Installing a non-installable engine -> 400 (never spawns pip).
        status, err = self._req("POST", "/api/models/engine/install",
                                {"engine": "faster"})
        self.assertEqual(status, 400)
        self.assertEqual(err["error"]["code"], "bad_engine")
        # Removing an engine that isn't installed -> 404.
        status, err = self._req("POST", "/api/models/engine/uninstall",
                                {"engine": "whisper"})
        self.assertEqual(status, 404)
        self.assertEqual(err["error"]["code"], "not_installed")

    def test_settings_partial_put_is_non_destructive(self):
        _, before = self._req("GET", "/api/settings")
        status, after = self._req("PUT", "/api/settings",
                                  {"model": "small.en",
                                   "beam_size": "junk"})
        self.assertEqual(status, 200)
        self.assertEqual(after["model"], "small.en")
        self.assertEqual(after["beam_size"], before["beam_size"])
        self.assertEqual(after["gap"], before["gap"])
        # Restore.
        self._req("PUT", "/api/settings", {"model": before["model"]})

    def test_run_validations_over_http(self):
        status, err = self._req("POST", "/api/run", {"input": ""})
        self.assertEqual(status, 400)
        self.assertEqual(err["error"]["code"], "missing_input")
        status, err = self._req(
            "POST", "/api/run",
            {"input": str(Path(self.tmp.name) / "ghost.mp3")})
        self.assertEqual(status, 400)
        self.assertEqual(err["error"]["code"], "input_not_found")

    def test_run_refuses_output_over_recording(self):
        # The output must never be the input, even with force - a
        # transcript once destroyed a source recording this way.
        src = Path(self.tmp.name) / "evidence.mp3"
        src.write_bytes(b"\x00" * 64)
        status, err = self._req(
            "POST", "/api/run",
            {"input": str(src), "output": str(src), "force": True})
        self.assertEqual(status, 400)
        self.assertEqual(err["error"]["code"], "output_is_input")
        # ...nor may it wear any audio/video extension.
        status, err = self._req(
            "POST", "/api/run",
            {"input": str(src),
             "output": str(Path(self.tmp.name) / "other.wav"),
             "force": True})
        self.assertEqual(status, 400)
        self.assertEqual(err["error"]["code"], "output_looks_like_media")

    def test_audio_404_when_none(self):
        status, err = self._req("GET", "/audio/current")
        self.assertEqual(status, 404)
        self.assertEqual(err["error"]["code"], "no_audio")

    def test_url_open_scheme_guard(self):
        opened = []
        saved = T._open_url
        T._open_url = opened.append
        self.addCleanup(lambda: setattr(T, "_open_url", saved))
        status, _ = self._req("POST", "/api/url/open",
                              {"url": "https://example.com"})
        self.assertEqual(status, 200)
        status, _ = self._req("POST", "/api/url/open",
                              {"url": "mailto:x@y.z"})
        self.assertEqual(status, 200)
        # Anything else (files, javascript:, empty) is refused.
        for bad in ("file:///etc/passwd", "javascript:alert(1)", ""):
            status, err = self._req("POST", "/api/url/open", {"url": bad})
            self.assertEqual(status, 400, bad)
            self.assertEqual(err["error"]["code"], "bad_url")
        self.assertEqual(opened,
                         ["https://example.com", "mailto:x@y.z"])

    def test_show_diarize_hidden_by_default(self):
        merged = T.validate_settings({})
        self.assertFalse(merged["show_diarize"])
        merged = T.validate_settings({"show_diarize": True})
        self.assertTrue(merged["show_diarize"])

    def test_annotations_roundtrip(self):
        T.annotations_clear()
        self.addCleanup(T.annotations_clear)
        # Off by default; the routes exist regardless.
        _, meta = self._req("GET", "/api/meta")
        self.assertFalse(meta["annotate"])
        # The annotate.on marker file enables the overlay live (no
        # restart), for the double-clicked app.
        marker = T._config_dir() / "annotate.on"
        marker.touch()
        try:
            _, meta = self._req("GET", "/api/meta")
            self.assertTrue(meta["annotate"])
        finally:
            marker.unlink()
        status, rec = self._req(
            "POST", "/api/annotations",
            {"view": "transcribe", "selector": "button.x",
             "element_text": "Run", "html": "<button>x</button>",
             "rect": "{}", "note": "make it green"})
        self.assertEqual(status, 200)
        self.assertEqual(rec["id"], 1)
        self.assertEqual(rec["app_version"], T.__version__)
        # An empty note is refused.
        status, err = self._req("POST", "/api/annotations", {"note": " "})
        self.assertEqual(status, 400)
        self.assertEqual(err["error"]["code"], "empty_note")
        # Oversized fields are capped, not rejected.
        status, rec2 = self._req("POST", "/api/annotations",
                                 {"note": "n", "html": "x" * 9000})
        self.assertEqual(status, 200)
        self.assertEqual(len(rec2["html"]), 4000)
        status, res = self._req("GET", "/api/annotations")
        self.assertEqual([a["id"] for a in res["items"]], [1, 2])
        status, _ = self._req("POST", "/api/annotations/delete", {"id": 1})
        self.assertEqual(status, 200)
        _, res = self._req("GET", "/api/annotations")
        self.assertEqual([a["id"] for a in res["items"]], [2])
        self._req("POST", "/api/annotations/clear")
        _, res = self._req("GET", "/api/annotations")
        self.assertEqual(res["items"], [])

    def _write_transcript(self):
        p = Path(self.tmp.name) / "doc.transcript.txt"
        T.write_paragraphs_to_file(
            _doc(), p, show_timestamp=True, title="HTTP test",
            output_format="txt", speakers=["Alice", None, None])
        return str(p)

    def test_open_mutate_save_roundtrip(self):
        path = self._write_transcript()
        status, res = self._req("POST", "/api/transcripts/open",
                                {"path": path})
        self.assertEqual(status, 200)
        doc = res["review"]
        self.assertEqual(doc["speaker_names"]["1"], "Alice")
        self.assertTrue(doc["loaded"])

        rev = doc["rev"]
        status, delta = self._req("POST", "/api/review/speaker",
                                  {"rev": rev, "index": 1, "slot": "2"})
        self.assertEqual(status, 200)
        rev = delta["rev"]
        status, delta = self._req("POST", "/api/review/speaker-name",
                                  {"rev": rev, "slot": "2",
                                   "name": "Bob"})
        rev = delta["rev"]
        status, full = self._req("POST", "/api/review/edit",
                                 {"rev": rev, "index": 2,
                                  "text": "Very good."})
        rev = full["rev"]

        status, out = self._req("POST", "/api/review/save",
                                {"rev": rev, "mode": "labels"})
        self.assertEqual(status, 200)
        text = Path(out["out_path"]).read_text()
        self.assertIn("Alice", text)
        self.assertIn("Bob", text)
        self.assertIn("Very good.", text)
        # Session is gone afterwards.
        status, _ = self._req("GET", "/api/review")
        self.assertEqual(status, 404)
        # ... and the file landed in recents (paths stored resolved).
        _, recents = self._req("GET", "/api/recents")
        self.assertIn(str(Path(out["out_path"]).resolve()),
                      [str(Path(r["path"]).resolve())
                       for r in recents["items"]])

    def test_stale_rev_conflicts(self):
        path = self._write_transcript()
        _, res = self._req("POST", "/api/transcripts/open", {"path": path})
        rev = res["review"]["rev"]
        self._req("POST", "/api/review/speaker",
                  {"rev": rev, "index": 0, "slot": "1"})
        status, err = self._req("POST", "/api/review/speaker",
                                {"rev": rev, "index": 1, "slot": "2"})
        self.assertEqual(status, 409)
        self.assertEqual(err["error"]["code"], "stale_rev")
        self._req("POST", "/api/review/close",
                  {"rev": err["error"]["rev"]})

    def test_second_open_refused_while_review_open(self):
        path = self._write_transcript()
        _, res = self._req("POST", "/api/transcripts/open", {"path": path})
        status, err = self._req("POST", "/api/transcripts/open",
                                {"path": path})
        self.assertEqual(status, 409)
        self.assertEqual(err["error"]["code"], "review_open")
        self._req("POST", "/api/review/close",
                  {"rev": res["review"]["rev"]})

    def test_sse_stream_delivers_events(self):
        import urllib.request
        req = urllib.request.Request(
            f"{self.base}/api/events?token={self.TOKEN}")
        resp = urllib.request.urlopen(req, timeout=10)
        try:
            self.assertEqual(
                resp.headers.get("Content-Type"), "text/event-stream")
            # Padding comment, then retry hint.
            first = resp.readline()
            self.assertTrue(first.startswith(b":"))
            self.assertIn(b"retry:", resp.readline())
            resp.readline()
            self.backend.broker.publish("log", {"text": "ping-test"})
            deadline_lines = []
            for _ in range(8):
                line = resp.readline()
                deadline_lines.append(line)
                if b"ping-test" in line:
                    break
            self.assertTrue(any(b"ping-test" in ln
                                for ln in deadline_lines))
        finally:
            resp.close()

    def test_autosave_matches_golden_fixture(self):
        """Schema-drift tripwire: the autosave a web session writes must
        match the v0.6.0 file byte-for-byte in structure and value
        types (dynamic fields normalised)."""
        import json as _json
        golden = _json.loads(
            (Path(__file__).parent / "fixtures"
             / "autosave-0.6.0.json").read_text())

        out = str(Path(self.tmp.name) / "golden.transcript.txt")
        info = {
            "paragraphs": _doc(), "out_path": out,
            "show_timestamp": True, "title": "Golden doc",
            "output_format": "txt", "result": None,
            "extra_formats": [], "loaded": False,
            "audio_path": None, "word_conf": None,
        }
        session = T.ReviewSession(info, T.EventBroker())
        rev = session.mutate(session.model.rev, "speaker",
                             {"index": 0, "slot": "3"})["rev"]
        session.mutate(rev, "speaker-name", {"slot": "3", "name": "Q C"})
        session.model.flush_autosave()
        produced = T._autosave_load()

        golden["out_path"] = out
        golden["saved_at"] = produced["saved_at"]
        golden["diarized"] = False       # intentional 0.9.x additions
        golden["verified_by"] = None
        golden["ts_marks"] = [None, None, None]
        golden["reviewed"] = [False, False, False]
        self.assertEqual(produced, golden)

        # A pre-0.9.0 autosave (neither new key) must still restore.
        del golden["diarized"]
        del golden["verified_by"]
        del golden["ts_marks"]
        del golden["reviewed"]
        restored = T.autosave_restore_info(golden)
        self.assertFalse(restored["diarized"])
        self.assertIsNone(restored["verified_by"])
        self.assertEqual(restored["preset_ts_marks"], [])
        self.assertEqual(restored["preset_reviewed"], [])
        session.model.close()


# =====================================================================
# AudioPrep
# =====================================================================

class TestAudioPrep(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory(prefix="transcribr-audio-")
        self.addCleanup(self.tmp.cleanup)
        self.broker = T.EventBroker()

    @staticmethod
    def _wait(prep, timeout=15.0):
        import time as _time
        deadline = _time.monotonic() + timeout
        while prep.state not in ("ready", "unavailable"):
            if _time.monotonic() > deadline:
                raise AssertionError(f"stuck in {prep.state!r}")
            _time.sleep(0.05)

    def _tiny_wav(self):
        import wave
        p = Path(self.tmp.name) / "tone.wav"
        with wave.open(str(p), "wb") as w:
            w.setnchannels(1)
            w.setsampwidth(2)
            w.setframerate(8000)
            w.writeframes(b"\x00\x00" * 4000)   # 0.5s of silence
        return p

    def test_no_source_is_unavailable(self):
        prep = T.AudioPrep(None, self.broker)
        self.assertEqual(prep.state, "unavailable")
        self.assertEqual(prep.status(), {"state": "unavailable"})

    def test_missing_file_is_unavailable(self):
        prep = T.AudioPrep(str(Path(self.tmp.name) / "ghost.mp3"),
                           self.broker)
        self._wait(prep)
        self.assertEqual(prep.state, "unavailable")

    def test_passthrough_serves_source(self):
        wav = self._tiny_wav()
        prep = T.AudioPrep(str(wav), self.broker)
        self._wait(prep)
        self.assertEqual(prep.state, "ready")
        self.assertEqual(prep.serve_path, str(wav))
        self.assertEqual(prep.status()["url"],
                         f"/audio/current?v={prep.version}")

    def test_each_session_gets_a_distinct_audio_url(self):
        # Regression: the URL used to be the same fixed string for
        # every review, so the browser kept playing the previously
        # loaded recording when the next transcript was opened.
        wav = self._tiny_wav()
        a = T.AudioPrep(str(wav), self.broker)
        b = T.AudioPrep(str(wav), self.broker)
        self._wait(a)
        self._wait(b)
        self.assertNotEqual(a.status()["url"], b.status()["url"])

    @unittest.skipUnless(T.shutil.which("ffmpeg"), "needs ffmpeg")
    def test_extracts_non_passthrough_and_caches(self):
        import subprocess as sp
        wav = self._tiny_wav()
        flac = Path(self.tmp.name) / "tone.flac"
        sp.run(["ffmpeg", "-hide_banner", "-loglevel", "error", "-y",
                "-i", str(wav), str(flac)], check=True, timeout=60)

        prep = T.AudioPrep(str(flac), self.broker)
        self._wait(prep)
        self.assertEqual(prep.state, "ready", prep.error)
        self.assertTrue(prep.serve_path.endswith(".m4a"))
        cached = Path(prep.serve_path)
        self.assertTrue(cached.exists())
        self.assertEqual(cached.parent, T._audio_cache_dir())

        # Second prep of the same source hits the cache.
        prep2 = T.AudioPrep(str(flac), self.broker)
        self._wait(prep2)
        self.assertEqual(prep2.serve_path, prep.serve_path)


# =====================================================================
# build_worker_params
# =====================================================================

class TestOutputSafety(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory(prefix="transcribr-safety-")
        self.addCleanup(self.tmp.cleanup)
        self.src = Path(self.tmp.name) / "recording.mp3"
        self.src.write_bytes(b"\x00" * 64)

    def test_identical_path_refused(self):
        with self.assertRaises(T.ApiFail) as cm:
            T.ensure_output_is_safe(str(self.src), str(self.src))
        self.assertEqual(cm.exception.code, "output_is_input")

    def test_differently_spelled_same_path_refused(self):
        sneaky = Path(self.tmp.name) / "sub" / ".." / "recording.mp3"
        with self.assertRaises(T.ApiFail) as cm:
            T.ensure_output_is_safe(str(self.src), str(sneaky))
        self.assertEqual(cm.exception.code, "output_is_input")

    def test_media_extension_refused_even_for_other_files(self):
        other = Path(self.tmp.name) / "different-name.m4a"
        with self.assertRaises(T.ApiFail) as cm:
            T.ensure_output_is_safe(str(self.src), str(other))
        self.assertEqual(cm.exception.code, "output_looks_like_media")

    def test_normal_outputs_pass(self):
        for name in ("recording.transcript.docx", "out.txt", "x.pdf"):
            T.ensure_output_is_safe(
                str(self.src), str(Path(self.tmp.name) / name))


class TestBuildWorkerParams(unittest.TestCase):
    _ENGINES = [("whisper", "OpenAI Whisper (reference)"),
                ("faster", "faster-whisper (CTranslate2)")]

    def setUp(self):
        self._saved = T.AVAILABLE_ENGINES
        T.AVAILABLE_ENGINES = list(self._ENGINES)
        self.settings = T.default_settings()
        self.settings["engine"] = "faster-whisper (CTranslate2)"

    def tearDown(self):
        T.AVAILABLE_ENGINES = self._saved

    def _params(self, **overrides):
        self.settings.update(overrides)
        return T.build_worker_params(
            self.settings, "/tmp/interview.mp3", "/tmp/out.docx",
            review_before_save=True)

    def test_engine_display_name_maps_to_key(self):
        self.assertEqual(self._params()["engine"], "faster")

    def test_unknown_engine_resolves_to_best_installed(self):
        # Stale/unknown names resolve like the Automatic entry: prefer
        # mlx, then faster, then the reference engine.
        self.assertEqual(
            self._params(engine="Something Else")["engine"], "faster")

    def test_automatic_engine_prefers_mlx_when_installed(self):
        T.AVAILABLE_ENGINES = list(self._ENGINES) + [
            ("mlx", "mlx-whisper (Apple Silicon)")]
        self.assertEqual(
            self._params(engine=T.ENGINE_AUTO_NAME)["engine"], "mlx")

    def test_automatic_engine_falls_back_to_faster(self):
        self.assertEqual(
            self._params(engine=T.ENGINE_AUTO_NAME)["engine"], "faster")

    def test_word_timestamps_always_on(self):
        # Since 0.9.0 word timings are always recorded: they sharpen
        # paragraph gaps, playback spans and confidence shading. A
        # stale settings.json key cannot turn them off.
        p = self._params(diarize=True)
        self.assertTrue(p["word_timestamps"])
        self.assertTrue(p["diarize"])
        merged = T.validate_settings({"word_timestamps": False})
        self.assertNotIn("word_timestamps", merged)

    def test_diarize_model_and_threshold_flow_through(self):
        p = self._params(diarize=True, diarize_model="campplus",
                         diarize_threshold=0.35)
        self.assertEqual(p["diarize_model"], "campplus")
        self.assertAlmostEqual(p["diarize_threshold"], 0.35)

    def test_title_falls_back_to_filename(self):
        p = self._params(title="", prompt="")
        self.assertEqual(p["title"], "interview.mp3")
        self.assertIsNone(p["initial_prompt"])

    def test_title_and_prompt_are_independent(self):
        # Title only: titles the doc, but nothing is fed to the engine.
        p = self._params(title="  Smith v Jones directions hearing ",
                         prompt="")
        self.assertEqual(p["title"], "Smith v Jones directions hearing")
        self.assertIsNone(p["initial_prompt"])
        # Prompt only: primes the engine, title falls back to the filename.
        p = self._params(title="", prompt="  Macklebum, Bloggs, DVEC ")
        self.assertEqual(p["initial_prompt"], "Macklebum, Bloggs, DVEC")
        self.assertEqual(p["title"], "interview.mp3")

    def test_legacy_prompt_migrates_to_title(self):
        # A pre-split settings.json (no "title" key) moves its prompt into
        # the title and stops priming the engine.
        s = T.validate_settings({"prompt": "Smith v Jones"})
        self.assertEqual(s["title"], "Smith v Jones")
        self.assertEqual(s["prompt"], "")
        # Post-split settings (a "title" key present) are left as-is.
        s2 = T.validate_settings({"title": "Doc", "prompt": "keywords"})
        self.assertEqual(s2["title"], "Doc")
        self.assertEqual(s2["prompt"], "keywords")

    def test_stale_review_and_highlight_keys_are_dropped(self):
        # Both became always-on / review-pane concerns in 0.9.0.
        merged = T.validate_settings({"review": False,
                                      "highlight_confidence": True})
        self.assertNotIn("review", merged)
        self.assertNotIn("highlight_confidence", merged)
        self.assertNotIn("highlight_confidence", self._params())

    def test_language_display_to_code(self):
        self.assertEqual(self._params(language="German")["language"], "de")
        self.assertIsNone(self._params(language="Auto-detect")["language"])
        self.assertEqual(
            self._params(language="Klingon")["language"], "en")

    def test_extra_formats_list(self):
        p = self._params(extra_srt=True, extra_json=True)
        self.assertEqual(p["extra_formats"], ["json", "srt"])
        self.assertEqual(self._params(extra_srt=False,
                                      extra_json=False)["extra_formats"], [])

    def test_review_flag_passthrough(self):
        p = T.build_worker_params(self.settings, "/tmp/a.mp3", "/tmp/a.txt",
                                  review_before_save=False)
        self.assertFalse(p["review_before_save"])

    def test_no_engine_raises(self):
        T.AVAILABLE_ENGINES = []
        with self.assertRaises(T._EngineNotAvailable):
            self._params()


if __name__ == "__main__":
    unittest.main()
